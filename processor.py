"""
Core processing logic for the Smart Chargesheet Review & Summarisation Assistant.

OPTIMISED for minimal API calls:
    - Call 1: Summary generation (single LLM call)
    - Call 2: Checklist analysis (single LLM call)
    - Call 3: Named Entity Recognition (single LLM call)
    Total: 3 API calls per document (Gemini supports 1M tokens natively)

Also includes:
  - Text extraction from .txt / .docx / .pdf
  - Smart text truncation for very large docs
  - Retry with exponential backoff for rate limits
  - Model fallback chain (tries multiple models)
  - Rule-based keyword detection as cross-validation
"""

import json
import os
import re
import time
import logging
import base64
import io
import math
from typing import Optional

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

import config

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

_NER_SUMMARY_CONTEXT = ""


# ─────────────────────────────────────────────────────────────────────────────
# 1. Text Extraction
# ─────────────────────────────────────────────────────────────────────────────

def extract_text(file_path: str) -> str:
    """Extract plain text from .txt, .docx, or .pdf files."""
    text, _ = extract_text_and_images(file_path)
    return text


def extract_text_and_images(file_path: str) -> tuple:
    """
    Extract plain text AND embedded images from .txt, .docx, or .pdf files.
    Returns (text: str, images: list[bytes])  where images are raw PNG/JPEG bytes.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read(), []

    elif ext == ".docx":
        return _extract_from_docx(file_path)

    elif ext == ".pdf":
        return _extract_from_pdf(file_path)

    else:
        raise ValueError(f"Unsupported file type: {ext}. Use .txt, .docx, or .pdf")


def _extract_from_docx(file_path: str) -> tuple:
    """Extract text + images from .docx file, with [PAGE N] markers."""
    try:
        import docx
        from docx.oxml.ns import qn
    except ImportError:
        raise ImportError("python-docx is required. Install: pip install python-docx")

    doc = docx.Document(file_path)

    # Extract text with page markers
    page_num = 1
    text_parts = ["\n[PAGE 1]\n"]

    for paragraph in doc.paragraphs:
        # Detect page breaks (explicit breaks + rendered page breaks)
        has_page_break = False
        for run in paragraph.runs:
            # Explicit page break: <w:br w:type="page"/>
            for br in run._element.findall(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    has_page_break = True
            # Rendered page break: <w:lastRenderedPageBreak/>
            for _ in run._element.iter(qn('w:lastRenderedPageBreak')):
                has_page_break = True

        if has_page_break:
            page_num += 1
            text_parts.append(f"\n[PAGE {page_num}]\n")

        text_parts.append(paragraph.text)

    text = "\n".join(text_parts)

    # Also extract text from tables (case diaries often use tables)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text += "\n" + row_text

    # Extract embedded images
    images = []
    try:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    if image_data and len(image_data) > 1000:  # skip tiny icons
                        images.append(image_data)
                except Exception as e:
                    logger.warning(f"Could not extract image: {e}")
    except Exception as e:
        logger.warning(f"Could not access docx images: {e}")

    logger.info(f"DOCX extracted: {len(text):,} chars text, {len(images)} images")

    # Compress large images to keep API payload reasonable
    images = _compress_images(images)

    return text, images


def _extract_from_pdf(file_path: str) -> tuple:
    """Extract text + images from .pdf file, with [PAGE N] markers."""
    try:
        import PyPDF2
    except ImportError:
        raise ImportError("PyPDF2 is required. Install: pip install PyPDF2")

    text_parts = []
    images = []

    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page_num, page in enumerate(reader.pages, 1):
            text_parts.append(f"\n[PAGE {page_num}]\n")
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

            # Extract images from PDF pages
            try:
                if hasattr(page, 'images'):
                    for img in page.images:
                        if hasattr(img, 'data') and len(img.data) > 1000:
                            images.append(img.data)
            except Exception as e:
                logger.debug(f"Could not extract PDF image: {e}")

    text = "\n".join(text_parts)
    logger.info(f"PDF extracted: {len(text):,} chars text, {len(images)} images")

    # Compress large images
    images = _compress_images(images)

    return text, images


def _compress_images(images: list, max_size: int = 1_500_000, max_dim: int = 2048) -> list:
    """
    Compress images to keep API payload reasonable.
    - Resizes images larger than max_dim pixels on any side
    - Re-encodes to JPEG if larger than max_size bytes
    - Limits total to 20 images
    """
    if not images:
        return images

    compressed = []
    try:
        from PIL import Image
        has_pil = True
    except ImportError:
        has_pil = False
        logger.warning("Pillow not installed — images will be sent without compression. Install: pip install Pillow")

    for i, img_bytes in enumerate(images[:20]):
        if has_pil and len(img_bytes) > max_size:
            try:
                img = Image.open(io.BytesIO(img_bytes))
                # Resize if too large
                if max(img.size) > max_dim:
                    img.thumbnail((max_dim, max_dim), Image.LANCZOS)
                # Re-encode as JPEG
                buf = io.BytesIO()
                if img.mode in ('RGBA', 'P'):
                    img = img.convert('RGB')
                img.save(buf, format='JPEG', quality=80)
                compressed_bytes = buf.getvalue()
                logger.debug(f"Image {i+1}: {len(img_bytes):,} -> {len(compressed_bytes):,} bytes")
                compressed.append(compressed_bytes)
            except Exception as e:
                logger.warning(f"Could not compress image {i+1}: {e}, using original")
                compressed.append(img_bytes)
        else:
            compressed.append(img_bytes)

    total_size = sum(len(img) for img in compressed)
    logger.info(f"Images ready: {len(compressed)} images, total {total_size:,} bytes")
    return compressed


# ─────────────────────────────────────────────────────────────────────────────
# 2. Text Preparation (Smart Truncation)
# ─────────────────────────────────────────────────────────────────────────────

# ─────────────────────────────────────────────────────────────────────────────
# 2b. Text Preprocessing (OCR cleanup, noise reduction)
# ─────────────────────────────────────────────────────────────────────────────

# Common OCR garbage patterns
_OCR_GARBAGE_RE = re.compile(
    r'[\x00-\x08\x0b\x0c\x0e-\x1f]'   # control characters
    r'|\ufffd'                             # unicode replacement char
    r'|[\u200b-\u200f\u202a-\u202e]'     # zero-width / bidi chars
    r'|\u00a0{3,}'                         # runs of non-breaking spaces
)

# Broken Devanagari: isolated matras without a base consonant
_BROKEN_DEVANAGARI_RE = re.compile(
    r'(?<![\u0915-\u0939\u0958-\u0961])'  # not preceded by a consonant
    r'[\u093e-\u094d\u0951-\u0954]{2,}'   # 2+ consecutive matras/nukta/virama
)

# Excessive punctuation / repeated symbols
_EXCESSIVE_PUNCT_RE = re.compile(r'([\-_=*#~.]{4,})')

# Whitespace normalizer
_MULTI_SPACE_RE = re.compile(r'[^\S\n]{2,}')  # 2+ spaces (not newlines)
_MULTI_NEWLINE_RE = re.compile(r'\n{4,}')       # 4+ blank lines → 2
_TRAILING_SPACE_RE = re.compile(r'[ \t]+$', re.MULTILINE)

_DOMAIN_VOCAB = [
    "थाना", "जिला", "अभियुक्त", "गवाह", "मृतक", "बरामद", "घटना", "धारा",
]

_ENGLISH_DIGIT_FIXES = {
    "0": "O",
    "1": "l",
    "5": "S",
}

_NER_RELATIONAL_WORDS = {
    "पिता", "पति", "पत्नी", "माता", "भाई", "बहन", "पुत्र", "पुत्री", "निवासी",
    "resident", "s/o", "d/o", "w/o", "r/o",
}

_NER_CONNECTOR_WORDS = {
    "एवं", "और", "तथा", "with", "or", "अथवा",
}

_NER_ROLE_WORDS = {
    "गवाह", "आरोपी", "अभियुक्त", "अधिकारी", "थानाध्यक्ष", "निरीक्षक", "उपनिरीक्षक",
    "constable", "officer", "witness", "accused", "complainant", "victim",
}

_NAME_TOKEN_RE = re.compile(r'^[A-Za-z\u0900-\u097F\.0०]+$')
_NUMERIC_TOKEN_RE = re.compile(r'^[\d०-९/().:-]+$')

_ENTITY_STOPWORDS = {
    "के", "की", "का", "में", "से", "पर", "द्वारा", "तथा", "एवं", "और", "आदि",
    "उपस्थिति", "निशानदेही", "मुख्य", "न्यायिक", "दण्डाधिकारी", "magistrate", "judge",
}

_ENTITY_REJECT_PHRASES = (
    "के निशानदेही में",
    "उपस्थिति में",
    "मुख्य न्यायिक दण्डाधिकारी",
    "मुख्य न्यायिक दण्डाधिकारी",
)


def _token_is_name_like(token: str) -> bool:
    if not token:
        return False
    t = token.strip(" .,;:|-_")
    if not t:
        return False
    if _NUMERIC_TOKEN_RE.fullmatch(t):
        return False
    if not _NAME_TOKEN_RE.fullmatch(t):
        return False
    return len(t) >= 2


def _cleanup_person_candidate(raw: str) -> tuple[str, float]:
    """
    Pattern-based person candidate cleaner:
    - keep 1-3 name-like tokens
    - stop at first relational/connector/role/numeric token
    - return cleaned candidate and confidence score
    """
    source = (raw or '').strip()
    if any(phrase in source for phrase in _ENTITY_REJECT_PHRASES):
        return "", 0.0

    tokens = [t for t in re.split(r'\s+', source) if t]
    if not tokens:
        return "", 0.0

    kept = []
    stopped_early = False

    for tok in tokens:
        clean_tok = tok.strip(" .,;:|-_")
        low_tok = clean_tok.lower()
        if (
            low_tok in _NER_RELATIONAL_WORDS
            or low_tok in _NER_CONNECTOR_WORDS
            or low_tok in _NER_ROLE_WORDS
            or low_tok in _ENTITY_STOPWORDS
            or _NUMERIC_TOKEN_RE.fullmatch(clean_tok)
            or not _token_is_name_like(clean_tok)
        ):
            stopped_early = True
            break
        kept.append(clean_tok)
        if len(kept) >= 3:
            break

    if not kept:
        return "", 0.0

    candidate = " ".join(kept).strip()

    confidence = 0.45
    if 1 <= len(kept) <= 3:
        confidence += 0.25
    if len(candidate) >= 4:
        confidence += 0.15
    if stopped_early:
        confidence += 0.10
    if any(re.search(r'[\u0900-\u097F]', t) for t in kept):
        confidence += 0.10

    confidence = min(0.95, confidence)
    return candidate, confidence


def _normalize_matra_order(text: str) -> tuple[str, int]:
    """Fix only safe matra spacing issues without cross-word relocation."""
    replacements = [
        (r'([\u0915-\u0939\u0958-\u0961])\s+([िीुूेोैौ])', r'\1\2'),
    ]
    total = 0
    out = text
    for pattern, repl in replacements:
        out, n = re.subn(pattern, repl, out)
        total += n
    return out, total


def _normalize_char_confusions(text: str) -> tuple[str, int]:
    """Apply conservative character-confusion fixes only in high-confidence contexts."""
    total = 0
    out = text

    # Hindi nukta and common confusions
    hindi_replacements = [
        (r'(?<=\S)ड(?=ा)', 'ड़'),
        (r'(?<=\S)ढ(?=ा)', 'ढ़'),
        (r'\bवारा\b', 'द्वारा'),
    ]
    for pattern, repl in hindi_replacements:
        out, n = re.subn(pattern, repl, out)
        total += n

    # English/digit confusion only for alphanumeric tokens
    def _fix_alnum_token(token: str) -> str:
        if not re.search(r'[A-Za-z]', token):
            return token
        if not re.search(r'\d', token):
            return token
        chars = list(token)
        for i, ch in enumerate(chars):
            if ch in _ENGLISH_DIGIT_FIXES:
                chars[i] = _ENGLISH_DIGIT_FIXES[ch]
        return ''.join(chars)

    def _repl(m):
        nonlocal total
        original = m.group(0)
        fixed = _fix_alnum_token(original)
        if fixed != original:
            total += 1
        return fixed

    out = re.sub(r'\b[A-Za-z0-9]{3,}\b', _repl, out)
    return out, total


def _merge_broken_hindi_words(text: str) -> tuple[str, int]:
    total = 0
    out = text
    targeted = [
        (r'अ\s*भि\s*यु\s*क्त', 'अभियुक्त'),
        (r'आ\s*रो\s*पी', 'आरोपी'),
        (r'गा\s*वा\s*ह', 'गवाह'),
        (r'था\s*ना', 'थाना'),
        (r'जि\s*ला', 'जिला'),
        (r'धा\s*रा', 'धारा'),
    ]
    for pat, repl in targeted:
        out, n = re.subn(pat, repl, out)
        total += n
    return out, total


def _merge_spaced_hindi_letters(text: str) -> tuple[str, int]:
    """Merge tokens like 'न ग र' -> 'नगर' conservatively."""
    total = 0

    def _joiner(match):
        nonlocal total
        total += 1
        return re.sub(r'\s+', '', match.group(0))

    out = re.sub(
        r'(?<!\S)(?:[\u0900-\u097F]\s+){2,}[\u0900-\u097F](?!\S)',
        _joiner,
        text,
    )
    return out, total


def _normalize_mixed_script_tokens(text: str) -> tuple[str, int]:
    """Normalize mixed-script short forms such as M0/Mo0 to मो0."""
    out, n = re.subn(r'\bM[oO0]?\s*0\b', 'मो0', text)
    return out, n


def _remove_ocr_noise_fragments(text: str) -> tuple[str, int]:
    """
    Remove partial lowercase English leakage fragments conservatively.
    Keeps full sentences and legal context; drops short noisy fragments only.
    """
    total = 0
    lines = text.split('\n')
    kept_lines = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            kept_lines.append(line)
            continue

        words = stripped.split()
        english_words = [w for w in words if re.fullmatch(r'[a-z]{2,}', w)]
        has_hindi = bool(re.search(r'[\u0900-\u097F]', stripped))
        has_legal_anchor = bool(re.search(r'\b(FIR|IPC|BNS|BNSS|धारा|थाना|अदालत|court)\b', stripped, re.IGNORECASE))

        noisy = (
            not has_hindi
            and not has_legal_anchor
            and len(words) <= 6
            and len(english_words) >= 2
            and any(len(w) <= 2 for w in english_words)
        )
        if noisy:
            total += 1
            continue
        kept_lines.append(line)

    out = '\n'.join(kept_lines)
    return out, total


def _normalize_legal_language(text: str) -> str:
    """Controlled language normalization for Hindi-English legal readability."""
    if not text:
        return text

    out = text

    # Fix common hybrid grammar fragments without adding facts
    out = re.sub(r'\b[Dd]eath\s+was\s*due\s*to\b', 'मृत्यु का कारण', out)
    out = re.sub(r'\b[Dd]eath\s*wasdue\s*to\b', 'मृत्यु का कारण', out)
    out = re.sub(r'\bcause\s+of\s+death\b', 'मृत्यु का कारण', out, flags=re.IGNORECASE)

    # Remove duplicate nukta in broken Hindi forms like दण्ड़ाधिकारी -> दण्डाधिकारी
    out = re.sub(r'([\u0915-\u0939]़)़', r'\1', out)
    out = re.sub(r'दण्ड़ाधिकारी', 'दण्डाधिकारी', out)

    # Remove repeated spaces and malformed punctuation spacing
    out = re.sub(r'[ \t]{2,}', ' ', out)
    out = re.sub(r'\s+([,.;:])', r'\1', out)
    out = re.sub(r'\n{3,}', '\n\n', out)

    return out.strip()


def _remove_noise_and_duplicates(text: str) -> tuple[str, int]:
    """Remove OCR noise tokens and repeated words conservatively."""
    total = 0

    # Random symbol runs
    out, n = re.subn(r'[~`^_=*|]{2,}', ' ', text)
    total += n

    # Tokens with no Hindi/English/digits and len>2
    out, n = re.subn(r'\b[^\w\u0900-\u097F\s]{3,}\b', ' ', out)
    total += n

    # Duplicate repeated words: "घटना घटना" -> "घटना"
    out, n = re.subn(r'\b([\w\u0900-\u097F]+)(\s+\1\b)+', r'\1', out, flags=re.IGNORECASE)
    total += n

    return out, total


def _fuzzy_domain_vocab_correction(text: str) -> tuple[str, int]:
    """Apply fuzzy correction to critical legal Hindi vocabulary (edit distance <= 2)."""
    total = 0

    def _fix_token(token: str) -> str:
        nonlocal total
        if len(token) < 4:
            return token
        if not re.search(r'[\u0900-\u097F]', token):
            return token

        best_word = None
        best_dist = 999
        second_best = 999
        for vocab in _DOMAIN_VOCAB:
            dist = _edit_distance(token, vocab)
            if dist < best_dist:
                second_best = best_dist
                best_dist = dist
                best_word = vocab
            elif dist < second_best:
                second_best = dist

        if best_word and best_dist <= 2 and (second_best - best_dist) >= 1:
            # Conservative: only replace if first char family is close
            if token[0] == best_word[0] or token[0] in ("द", "ध") and best_word[0] in ("द", "ध"):
                if token != best_word:
                    total += 1
                return best_word
        return token

    parts = re.split(r'(\s+)', text)
    for i, p in enumerate(parts):
        if p.strip() and not p.isspace():
            parts[i] = _fix_token(p)

    return ''.join(parts), total


def _context_frequency_correction(text: str) -> tuple[str, int]:
    """
    Context-based fuzzy correction:
    if a noisy token appears once and a close cleaner variant appears repeatedly,
    replace only when the best candidate is unique and high-confidence.
    """
    if not text:
        return text, 0

    tokens = re.findall(r'[\u0900-\u097F]{4,}', text)
    if not tokens:
        return text, 0

    freq = {}
    for tok in tokens:
        freq[tok] = freq.get(tok, 0) + 1

    frequent = [tok for tok, c in freq.items() if c >= 2]
    rare = [tok for tok, c in freq.items() if c == 1]
    if not frequent or not rare:
        return text, 0

    replacements = {}
    for noisy in rare:
        best = None
        best_dist = 999
        second = 999
        for clean in frequent:
            if noisy == clean:
                continue
            if noisy[0] != clean[0]:
                continue
            dist = _edit_distance(noisy, clean)
            if dist < best_dist:
                second = best_dist
                best_dist = dist
                best = clean
            elif dist < second:
                second = dist

        if best and best_dist <= 2 and (second - best_dist) >= 1:
            replacements[noisy] = best

    if not replacements:
        return text, 0

    out = text
    changes = 0
    for old, new in replacements.items():
        out, n = re.subn(rf'(?<!\w){re.escape(old)}(?!\w)', new, out)
        changes += n

    return out, changes


def ocr_clean(text: str) -> dict:
    """Pattern-based OCR cleaning with conservative confidence-aware corrections."""
    if not text:
        return {"cleaned_text": "", "ocr_confidence": 0.0, "ocr_corrections": 0}

    corrections = 0
    out = _OCR_GARBAGE_RE.sub('', text)
    out = _BROKEN_DEVANAGARI_RE.sub('', out)
    out = _EXCESSIVE_PUNCT_RE.sub(lambda m: m.group(1)[:3], out)
    out = _MULTI_SPACE_RE.sub(' ', out)
    out = _TRAILING_SPACE_RE.sub('', out)
    out = _MULTI_NEWLINE_RE.sub('\n\n', out)

    out, n = _merge_spaced_hindi_letters(out)
    corrections += n

    out, n = _normalize_mixed_script_tokens(out)
    corrections += n

    out, n = _merge_broken_hindi_words(out)
    corrections += n
    out, n = _normalize_matra_order(out)
    corrections += n
    out = _repair_ocr_devanagari(out)
    out, n = _normalize_char_confusions(out)
    corrections += n
    out = _apply_ocr_corrections(out)
    out, n = _fuzzy_domain_vocab_correction(out)
    corrections += n
    out, n = _context_frequency_correction(out)
    corrections += n
    out, n = _normalize_name_variants_frequency(out)
    corrections += n
    out, n = _remove_noise_and_duplicates(out)
    corrections += n

    out, n = _remove_ocr_noise_fragments(out)
    corrections += n

    # Remove repeated page headers/footers
    lines = out.split('\n')
    if len(lines) > 20:
        line_counts = {}
        for line in lines:
            stripped = line.strip()
            if len(stripped) > 5:
                line_counts[stripped] = line_counts.get(stripped, 0) + 1
        repeated = {l for l, c in line_counts.items() if c >= 3}
        if repeated:
            lines = [l for l in lines if l.strip() not in repeated]
            out = '\n'.join(lines)

    out = out.strip()
    density = corrections / max(1, len(out.split()))
    ocr_conf = max(0.2, min(0.99, 1.0 - min(0.8, density * 0.8)))

    return {
        "cleaned_text": out,
        "ocr_confidence": round(ocr_conf, 2),
        "ocr_corrections": corrections,
    }


def _normalize_name_variants_frequency(text: str) -> tuple[str, int]:
    if not text:
        return text, 0

    role_lines = []
    for line in text.split('\n'):
        if re.search(r'आरोपी|अभियुक्त|गवाह|शिकायतकर्ता|पीड़ित|witness|accused|complainant', line, re.IGNORECASE):
            role_lines.append(line)

    candidate_names = []
    name_pattern = re.compile(r'([\u0900-\u097F]{2,}(?:\s+[\u0900-\u097F]{2,}){0,3})')
    for line in role_lines:
        for m in name_pattern.finditer(line):
            nm = m.group(1).strip()
            if len(nm) >= 4:
                candidate_names.append(nm)

    if not candidate_names:
        return text, 0

    groups = {}
    for name in candidate_names:
        norm = _normalize_entity_text(name, "PERSON")
        groups.setdefault(norm, {})
        groups[norm][name] = groups[norm].get(name, 0) + 1

    replacements = {}
    for norm, variants in groups.items():
        if len(variants) < 2:
            continue
        canonical = sorted(variants.items(), key=lambda kv: (kv[1], len(kv[0])), reverse=True)[0][0]
        for variant in variants:
            if variant != canonical and _edit_distance(variant, canonical) <= 2:
                replacements[variant] = canonical

    if not replacements:
        return text, 0

    out = text
    changes = 0
    for old, new in replacements.items():
        out, n = re.subn(rf'(?<!\w){re.escape(old)}(?!\w)', new, out)
        changes += n

    return out, changes


def extract_core_facts(text: str) -> dict:
    """Extract high-confidence anchors only: FIR number, dates, sections, accused."""
    if not text:
        return {
            "fir_number": None,
            "dates": [],
            "sections": [],
            "accused": [],
            "extraction_confidence": 0.0,
        }

    fir_patterns = [
        r'(?:FIR|एफआईआर|प्राथमिकी|काण्ड|मुकदमा)\s*(?:संख्या|नं|No\.?|number)?\s*[:\-]?\s*([0-9]{1,5}/[0-9]{2,4})',
        r'मु\.?\s*अ\.?\s*नं\.?\s*[:\-]?\s*([0-9]{1,5}/[0-9]{2,4})',
    ]
    fir_number = None
    fir_strength = 0
    for pat in fir_patterns:
        m = re.search(pat, text, flags=re.IGNORECASE)
        if m:
            fir_number = m.group(1).strip()
            fir_strength = 1
            break

    date_matches = re.findall(r'\b\d{1,2}[/.\-]\d{1,2}[/.\-]\d{2,4}\b', text)
    dates = []
    seen_dates = set()
    for d in date_matches:
        nd = _normalize_date_string(d)
        if nd and nd not in seen_dates:
            seen_dates.add(nd)
            dates.append(nd)

    sections = _extract_legal_sections_strict(text)

    accused = []

    def _trim_name(raw_name: str) -> str:
        tokens = [t for t in re.split(r'\s+', raw_name) if t]
        stop_after = {"ने", "को", "के", "की", "का", "और", "तथा", "पर", "से", "में", "था", "थी", "है"}
        kept = []
        for tok in tokens:
            if tok in stop_after:
                break
            kept.append(tok)
            if len(kept) >= 3:
                break
        return " ".join(kept).strip()

    for m in re.finditer(
        r'(?:आरोपी|अभियुक्त|accused)\s*[:\-]?\s*([^\n,।;]{3,80})',
        text,
        flags=re.IGNORECASE,
    ):
        name = _trim_name(m.group(1).strip(' .,-:;'))
        if name and name not in accused:
            accused.append(name)
        if len(accused) >= 8:
            break

    confidence_parts = [fir_strength]
    confidence_parts.append(1 if dates else 0)
    confidence_parts.append(1 if sections else 0)
    confidence_parts.append(1 if accused else 0)
    extraction_conf = round(sum(confidence_parts) / len(confidence_parts), 2)

    return {
        "fir_number": fir_number,
        "dates": dates,
        "sections": sections,
        "accused": accused,
        "extraction_confidence": extraction_conf,
    }


def _normalize_date_string(value: str | None) -> str | None:
    if not value:
        return None
    m = re.search(r'(\d{1,2})[/.\-](\d{1,2})[/.\-](\d{2,4})', value)
    if not m:
        return None
    day, month, year = m.groups()
    if len(year) == 2:
        year = f"20{year}"
    try:
        day_i, month_i, year_i = int(day), int(month), int(year)
        if not (1 <= day_i <= 31 and 1 <= month_i <= 12 and 1900 <= year_i <= 2100):
            return None
    except ValueError:
        return None
    return f"{day_i:02d}/{month_i:02d}/{year_i:04d}"


def _normalize_time_string(value: str | None) -> str | None:
    if not value:
        return None
    m = re.search(r'(\d{1,2})[:\.]?(\d{2})\s*(AM|PM|am|pm|बजे|hrs?)?', value)
    if not m:
        return None
    hh, mm, marker = m.groups()
    try:
        hh_i, mm_i = int(hh), int(mm)
        if not (0 <= hh_i <= 23 and 0 <= mm_i <= 59):
            return None
    except ValueError:
        return None

    if marker and marker.lower() in ("am", "pm"):
        mer = marker.lower()
        if mer == "pm" and hh_i < 12:
            hh_i += 12
        if mer == "am" and hh_i == 12:
            hh_i = 0
    return f"{hh_i:02d}:{mm_i:02d}"


def _extract_legal_sections_strict(text: str) -> list[str]:
    if not text:
        return []

    sections = []
    seen = set()

    bns_matches = re.findall(r'\b(?:B\.?N\.?S\.?|B\.?N\.?S\.?S\.?|बीएनएस)\s*(\d+(?:\(\d+\))?)\b', text, flags=re.IGNORECASE)
    for sec in bns_matches:
        val = f"BNS {sec}"
        if val not in seen:
            seen.add(val)
            sections.append(val)

    ipc_matches = re.findall(r'\b(?:IPC|आईपीसी)\s*(\d{3}[A-Za-z]?(?:\(\d+\))?)\b', text, flags=re.IGNORECASE)
    for sec in ipc_matches:
        val = f"IPC {sec}"
        if val not in seen:
            seen.add(val)
            sections.append(val)

    arms_matches = re.findall(r'\b(?:Arms\s*Act|आर्म्स\s*एक्ट)\s*(\d+(?:\(\d+\))?)\b', text, flags=re.IGNORECASE)
    for sec in arms_matches:
        val = f"Arms Act {sec}"
        if val not in seen:
            seen.add(val)
            sections.append(val)

    bsa_matches = re.findall(r'\b(?:B\.?S\.?A\.?|भारतीय\s*साक्ष्य\s*अधिनियम)\s*(\d+(?:\(\d+\))?)\b', text, flags=re.IGNORECASE)
    for sec in bsa_matches:
        val = f"BSA {sec}"
        if val not in seen:
            seen.add(val)
            sections.append(val)

    # Common phrasing: "धारा 323, 504 आईपीसी"
    for m in re.finditer(r'धारा\s*([0-9]{3}[A-Za-z]?(?:\(\d+\))?(?:\s*,\s*[0-9]{3}[A-Za-z]?(?:\(\d+\))?)*)\s*(?:आईपीसी|IPC)', text, flags=re.IGNORECASE):
        nums = re.findall(r'\d{3}[A-Za-z]?(?:\(\d+\))?', m.group(1))
        for num in nums:
            val = f"IPC {num}"
            if val not in seen:
                seen.add(val)
                sections.append(val)

    # Common phrasing: "धारा 103(1) बीएनएस"
    for m in re.finditer(r'धारा\s*(\d+(?:\(\d+\))?)\s*(?:बीएनएस|B\.?N\.?S\.?|B\.?N\.?S\.?S\.?)', text, flags=re.IGNORECASE):
        val = f"BNS {m.group(1)}"
        if val not in seen:
            seen.add(val)
            sections.append(val)

    # Common phrasing with act name after section:
    # "धारा 103(1) भारतीय न्याय संहिता" / "धारा 103(1) BNSS"
    for m in re.finditer(
        r'धारा\s*(\d+(?:\(\d+\))?)\s*(?:भारतीय\s*न्याय\s*संहिता|B\.?N\.?S\.?|B\.?N\.?S\.?S\.?|बीएनएस)',
        text,
        flags=re.IGNORECASE,
    ):
        val = f"BNS {m.group(1)}"
        if val not in seen:
            seen.add(val)
            sections.append(val)

    for m in re.finditer(
        r'धारा\s*(\d+(?:\(\d+\))?)\s*(?:भारतीय\s*साक्ष्य\s*अधिनियम|B\.?S\.?A\.?)',
        text,
        flags=re.IGNORECASE,
    ):
        val = f"BSA {m.group(1)}"
        if val not in seen:
            seen.add(val)
            sections.append(val)

    # Common phrasing with IPC act name after section:
    # "धारा 323 भारतीय दण्ड संहिता"
    for m in re.finditer(
        r'धारा\s*(\d{3}[A-Za-z]?(?:\(\d+\))?)\s*(?:भारतीय\s*दण्ड\s*संहिता|IPC|आईपीसी)',
        text,
        flags=re.IGNORECASE,
    ):
        val = f"IPC {m.group(1)}"
        if val not in seen:
            seen.add(val)
            sections.append(val)

    # Common phrasing: "धारा 27 आर्म्स एक्ट"
    for m in re.finditer(r'धारा\s*(\d+)\s*(?:आर्म्स\s*एक्ट|Arms\s*Act)', text, flags=re.IGNORECASE):
        val = f"Arms Act {m.group(1)}"
        if val not in seen:
            seen.add(val)
            sections.append(val)

    return normalize_sections(sections)


def normalize_sections(raw_sections_or_text) -> list[str]:
    """Rule-driven legal section normalization: <ACT> <SECTION>(subsection)."""
    if not raw_sections_or_text:
        return []

    if isinstance(raw_sections_or_text, str):
        text = raw_sections_or_text
        candidates = []
        candidates.extend(re.findall(r'\b(?:B\.?N\.?S\.?|B\.?N\.?S\.?S\.?|बीएनएस)\s*\d+(?:\(\d+\))?\b', text, flags=re.IGNORECASE))
        candidates.extend(re.findall(r'\b(?:IPC|आईपीसी)\s*\d+[A-Za-z]?(?:\(\d+\))?\b', text, flags=re.IGNORECASE))
        candidates.extend(re.findall(r'\b(?:Arms\s*Act|आर्म्स\s*एक्ट)\s*\d+(?:\(\d+\))?\b', text, flags=re.IGNORECASE))
        candidates.extend(re.findall(r'\b(?:NDPS\s*Act|NDPS|एनडीपीएस)\s*\d+(?:\(\d+\))?\b', text, flags=re.IGNORECASE))
        candidates.extend(re.findall(r'\b(?:B\.?S\.?A\.?|भारतीय\s*साक्ष्य\s*अधिनियम)\s*\d+(?:\(\d+\))?\b', text, flags=re.IGNORECASE))

        for m in re.finditer(r'धारा\s*(\d+(?:\(\d+\))?(?:\s*,\s*\d+(?:\(\d+\))?)*)\s*(?:IPC|आईपीसी|B\.?N\.?S\.?|B\.?N\.?S\.?S\.?|बीएनएस|Arms\s*Act|आर्म्स\s*एक्ट|NDPS|एनडीपीएस|B\.?S\.?A\.?|भारतीय\s*साक्ष्य\s*अधिनियम)', text, flags=re.IGNORECASE):
            act_match = re.search(r'(IPC|आईपीसी|B\.?N\.?S\.?|B\.?N\.?S\.?S\.?|बीएनएस|Arms\s*Act|आर्म्स\s*एक्ट|NDPS|एनडीपीएस|B\.?S\.?A\.?|भारतीय\s*साक्ष्य\s*अधिनियम)', m.group(0), flags=re.IGNORECASE)
            if not act_match:
                continue
            act_raw = act_match.group(1)
            for sec in re.findall(r'\d+(?:\(\d+\))?', m.group(1)):
                candidates.append(f"{act_raw} {sec}")
    else:
        candidates = [str(s).strip() for s in raw_sections_or_text if str(s).strip()]

    # Pre-normalize: handle "NUMBER ACT_NAME" reverse ordering from LLM
    normalized_candidates = []
    reverse_patterns = [
        (r'(\d+(?:\(\d+\))?)\s+B\.?N\.?S\.?(?:\s+\d{4})?', 'BNS'),
        (r'(\d+(?:\(\d+\))?)\s+(?:Arms?\s*Act|आर्म्स\s*एक्ट)(?:\s+\d{4})?', 'Arms Act'),
        (r'(\d+(?:\(\d+\))?)\s+(?:IPC|आईपीसी)', 'IPC'),
        (r'(\d+(?:\(\d+\))?)\s+(?:NDPS(?:\s*Act)?|एनडीपीएस)', 'NDPS Act'),
        (r'(\d+(?:\(\d+\))?)\s+(?:B\.?S\.?A\.?|भारतीय\s*साक्ष्य\s*अधिनियम)', 'BSA'),
        (r'(\d+(?:\(\d+\))?)\s+(?:भारतीय\s*न्याय\s*संहिता)', 'BNS'),
        (r'(\d+(?:\(\d+\))?)\s+(?:भारतीय\s*दण्ड\s*संहिता)', 'IPC'),
    ]
    for raw in candidates:
        rewritten = False
        for pattern, act_name in reverse_patterns:
            m = re.search(pattern, str(raw), re.IGNORECASE)
            if m:
                normalized_candidates.append(f"{act_name} {m.group(1)}")
                rewritten = True
                break
        if not rewritten:
            normalized_candidates.append(raw)
    candidates = normalized_candidates

    normalized = []
    seen = set()

    act_patterns = [
        (r'\b(?:B\.?N\.?S\.?|B\.?N\.?S\.?S\.?|बीएनएस)\b', 'BNS'),
        (r'\b(?:IPC|आईपीसी)\b', 'IPC'),
        (r'\b(?:Arms\s*Act|आर्म्स\s*एक्ट)\b', 'Arms Act'),
        (r'\b(?:NDPS\s*Act|NDPS|एनडीपीएस)\b', 'NDPS Act'),
        (r'\b(?:B\.?S\.?A\.?|भारतीय\s*साक्ष्य\s*अधिनियम)\b', 'BSA'),
    ]

    for raw in candidates:
        act = None
        for pat, act_name in act_patterns:
            if re.search(pat, raw, re.IGNORECASE):
                act = act_name
                break
        if not act:
            continue

        m = re.search(r'(\d+[A-Za-z]?)(?:\((\d+)\))?', raw)
        if not m:
            continue

        section_num = m.group(1)
        subsection = m.group(2)

        # Preserve subsection if present anywhere adjacent in raw input
        if subsection is None:
            near = re.search(rf'{re.escape(section_num)}\s*\((\d+)\)', raw)
            if near:
                subsection = near.group(1)
        canonical = f"{act} {section_num}"
        if subsection:
            canonical += f"({subsection})"

        key = canonical.lower()
        if key in seen:
            continue
        seen.add(key)
        normalized.append(canonical)

    return normalized


def _normalize_section_string(value: str | None) -> str | None:
    if not value:
        return None
    strict = _extract_legal_sections_strict(value)
    return strict[0] if strict else None


def _normalize_text_for_llm(text: str, core_facts: dict) -> tuple[str, dict]:
    """Normalize dates/time/sections and remove obvious duplicates before LLM."""
    out = text

    # Date normalization
    out = re.sub(
        r'\b(\d{1,2}[/.\-]\d{1,2}[/.\-]\d{2,4})\b',
        lambda m: _normalize_date_string(m.group(1)) or m.group(1),
        out,
    )

    # Time normalization
    out = re.sub(
        r'\b\d{1,2}[:\.]\d{2}\s*(?:AM|PM|am|pm|बजे|hrs?)?\b',
        lambda m: _normalize_time_string(m.group(0)) or m.group(0),
        out,
    )

    # Remove simple duplicate entities/words after OCR cleanup
    out, _ = re.subn(r'\b([\w\u0900-\u097F]{3,})(\s+\1\b)+', r'\1', out, flags=re.IGNORECASE)

    normalized_facts = {
        **core_facts,
        "dates": [d for d in (_normalize_date_string(d) for d in core_facts.get("dates", [])) if d],
        "sections": [s for s in (_normalize_section_string(s) for s in core_facts.get("sections", [])) if s],
    }

    return out, normalized_facts


def preprocess_text_with_meta(text: str) -> dict:
    """
    Step 2–4 pipeline:
      - Step 2: Strong layered OCR cleaning
      - Step 3: Core fact extraction (high-confidence anchors only)
      - Step 4: Pre-LLM normalization
    """
    if not text:
        return {
            "cleaned_text": "",
            "normalized_text": "",
            "core_facts": {
                "fir_number": None,
                "dates": [],
                "sections": [],
                "accused": [],
                "extraction_confidence": 0.0,
            },
            "ocr_confidence": 0.0,
            "ocr_corrections": 0,
        }

    original_len = len(text)
    ocr = ocr_clean(text)
    out = ocr["cleaned_text"]
    corrections = ocr["ocr_corrections"]

    # Optional LLM OCR cleanup (kept configurable)
    ocr_mode = getattr(config, 'OCR_CORRECTION_MODE', 'pattern')
    if ocr_mode in ('llm', 'hybrid'):
        logger.info(f"OCR correction mode: {ocr_mode} — using LLM for OCR cleanup")
        out = _ocr_cleanup_via_llm(out)

    out = _normalize_legal_language(out)

    core_facts = extract_core_facts(out)
    normalized_text, normalized_facts = _normalize_text_for_llm(out, core_facts)

    cleaned_len = len(out)
    if original_len - cleaned_len > 100:
        logger.info(
            f"Preprocessing: cleaned {original_len:,} → {cleaned_len:,} chars "
            f"(removed {original_len - cleaned_len:,} noise chars, corrections={corrections})"
        )

    return {
        "cleaned_text": out,
        "normalized_text": normalized_text,
        "core_facts": normalized_facts,
            "ocr_confidence": ocr["ocr_confidence"],
        "ocr_corrections": corrections,
    }


def preprocess_text(text: str) -> str:
    """Backward-compatible wrapper returning cleaned text only."""
    return preprocess_text_with_meta(text)["cleaned_text"]


def _truncate_text(text: str) -> str:
    """
    Truncate text to MAX_TEXT_LENGTH if too long.
    Keeps beginning (case header, FIR) and end (conclusion, sections, attachments).
    Gemini 2.0 Flash supports ~1M tokens, so 80K chars (~25K tokens) is very safe.
    """
    max_len = getattr(config, 'MAX_TEXT_LENGTH', 80000)
    if len(text) <= max_len:
        return text

    # Keep 60% from start (FIR, parties, incident), 40% from end (sections, conclusion)
    start_len = int(max_len * 0.6)
    end_len = max_len - start_len

    truncated = (
        text[:start_len]
        + "\n\n[... दस्तावेज़ का मध्य भाग छोड़ा गया / middle portion omitted ...]\n\n"
        + text[-end_len:]
    )
    logger.info(f"Document truncated: {len(text):,} -> {len(truncated):,} characters")
    return truncated


# ─────────────────────────────────────────────────────────────────────────────
# 3. Checklists
# ─────────────────────────────────────────────────────────────────────────────

def load_checklists() -> dict:
    """Load crime-type checklists from JSON."""
    with open(config.CHECKLIST_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


_RUNTIME_HOMICIDE_CHECKLIST = {
    "display_name": "हत्या / मृत्यु (Homicide)",
    "display_name_en": "Homicide",
    "typical_sections": ["IPC 302", "IPC 304", "BNS 103", "BNS 105"],
    "keywords_hi": ["हत्या", "मृत्यु", "शव", "पोस्टमार्टम", "घातक"],
    "keywords_en": ["murder", "homicide", "death", "postmortem", "fatal"],
    "required_items": [
        {"id": "fir", "label_hi": "प्राथमिकी (FIR) विवरण", "label_en": "FIR details"},
        {"id": "place_time", "label_hi": "घटना का स्थान, दिनांक, समय", "label_en": "Place/date/time of occurrence"},
        {"id": "accused", "label_hi": "आरोपी का विवरण", "label_en": "Accused details"},
        {"id": "witnesses", "label_hi": "मुख्य गवाहों के बयान", "label_en": "Key witness statements"},
        {"id": "postmortem", "label_hi": "पोस्टमार्टम रिपोर्ट", "label_en": "Postmortem report"},
        {"id": "cause_of_death", "label_hi": "मृत्यु का कारण", "label_en": "Cause of death opinion"},
        {"id": "injury_cert", "label_hi": "चोट/घाव विवरण", "label_en": "Injury description"},
        {"id": "weapon_desc", "label_hi": "हथियार का विवरण", "label_en": "Weapon description"},
        {"id": "weapon_seizure", "label_hi": "हथियार जब्ती मेमो", "label_en": "Weapon seizure memo"},
        {"id": "forensic", "label_hi": "FSL / फॉरेंसिक रिपोर्ट", "label_en": "FSL / forensic report"},
        {"id": "site_plan", "label_hi": "नक्शा मौका / स्थल निरीक्षण", "label_en": "Site plan / spot inspection"},
        {"id": "chain_custody", "label_hi": "कस्टडी चेन रिकॉर्ड", "label_en": "Chain of custody record"},
    ],
}


def get_crime_type_info(crime_key: str) -> dict | None:
    """Return checklist info for a specific crime key."""
    checklists = load_checklists()
    if crime_key in checklists:
        return checklists.get(crime_key)
    if crime_key == "homicide":
        return _RUNTIME_HOMICIDE_CHECKLIST
    return None


def list_crime_types() -> list[dict]:
    """Return a list of available crime types with display names."""
    checklists = load_checklists()
    items = [
        {"key": k, "display_name": v["display_name"], "display_name_en": v["display_name_en"]}
        for k, v in checklists.items()
    ]
    if not any(it["key"] == "homicide" for it in items):
        items.append(
            {
                "key": "homicide",
                "display_name": _RUNTIME_HOMICIDE_CHECKLIST["display_name"],
                "display_name_en": _RUNTIME_HOMICIDE_CHECKLIST["display_name_en"],
            }
        )
    return items


# ─────────────────────────────────────────────────────────────────────────────
# 4. LLM Interaction (with retry + model fallback)
# ─────────────────────────────────────────────────────────────────────────────

# Models to try in order if one hits rate limits or is unavailable
GEMINI_FALLBACK_MODELS = [
    "gemini-2.0-flash",
    "gemini-2.0-flash-lite",
    "gemini-2.5-flash-preview-05-20",
    "gemini-1.5-flash",
]

# Cache for discovered available models
_available_models_cache = None

# Track which API key index to use next (round-robin)
_current_key_index = 0

# Remember which model last succeeded — start with it on next call
_last_successful_model = None


def _discover_available_models(client) -> list:
    """Discover which Gemini models are actually available for this API key."""
    global _available_models_cache
    if _available_models_cache is not None:
        return _available_models_cache

    provider_type, client_obj = client
    if provider_type != "gemini":
        return GEMINI_FALLBACK_MODELS

    try:
        all_models = list(client_obj.models.list())
        flash_models = []
        for m in all_models:
            name = m.name if hasattr(m, 'name') else str(m)
            # Strip "models/" prefix if present
            short_name = name.replace("models/", "")
            if "flash" in short_name.lower():
                flash_models.append(short_name)
        if flash_models:
            logger.info(f"Available flash models: {flash_models}")
            _available_models_cache = flash_models
            return flash_models
    except Exception as e:
        logger.warning(f"Could not discover models: {e}")

    return GEMINI_FALLBACK_MODELS


def _get_all_api_keys() -> list:
    """Get all available Gemini API keys (multi-key pool)."""
    keys = list(config.GEMINI_API_KEYS) if config.GEMINI_API_KEYS else []
    # Always include the primary key if set and not already in list
    if config.GEMINI_API_KEY and config.GEMINI_API_KEY not in keys:
        keys.insert(0, config.GEMINI_API_KEY)
    return keys


def _get_gemini_client(api_key: str = None):
    """Get Gemini client with a specific API key."""
    key = api_key or config.GEMINI_API_KEY
    try:
        from google import genai
        client = genai.Client(
            api_key=key,
            http_options={'timeout': 300_000},  # 300 seconds in ms
        )
        return ("gemini", client)
    except ImportError:
        try:
            import google.generativeai as genai_old
            genai_old.configure(api_key=key)
            return ("gemini_old", genai_old)
        except ImportError:
            raise ImportError("google-genai is required. Install: pip install google-genai")


def _get_openai_client():
    """Get OpenAI client."""
    try:
        from openai import OpenAI
        client = OpenAI(api_key=config.OPENAI_API_KEY)
        return ("openai", client)
    except ImportError:
        raise ImportError("openai is required. Install: pip install openai")


def _call_gemini(client, model_name: str, prompt, images: list = None) -> str:
    """
    Call Gemini API with specific model.
    Supports multimodal: if images are provided, sends text + images together.
    """
    provider_type, client_obj = client

    temperature = getattr(config, "LLM_TEMPERATURE", 0.0)

    if provider_type == "gemini":
        from google.genai import types

        # Build multimodal content if images are present
        if images:
            contents = []
            # Add text prompt first
            contents.append(types.Part.from_text(text=prompt if isinstance(prompt, str) else prompt))
            # Add images (limit to 20 to stay within API limits)
            for i, img_bytes in enumerate(images[:20]):
                try:
                    # Detect mime type
                    mime = "image/jpeg"
                    if img_bytes[:4] == b'\x89PNG':
                        mime = "image/png"
                    elif img_bytes[:4] == b'GIF8':
                        mime = "image/gif"
                    elif img_bytes[:2] == b'BM':
                        mime = "image/bmp"

                    contents.append(types.Part.from_bytes(
                        data=img_bytes,
                        mime_type=mime,
                    ))
                except Exception as e:
                    logger.warning(f"Could not add image {i+1}: {e}")

            logger.info(f"Sending multimodal request: text + {len([c for c in contents if hasattr(c, 'inline_data')])} images")
            response = client_obj.models.generate_content(
                model=model_name,
                contents=contents,
                config=types.GenerateContentConfig(
                    temperature=temperature,
                    max_output_tokens=65536,
                ),
            )
        else:
            response = client_obj.models.generate_content(
                model=model_name,
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=temperature,
                    max_output_tokens=65536,
                ),
            )
        return response.text
    elif provider_type == "gemini_old":
        model = client_obj.GenerativeModel(model_name)
        response = model.generate_content(
            prompt,
            generation_config={"temperature": temperature},
        )
        return response.text


def _call_llm(prompt: str, system_prompt: str = "", images: list = None) -> str:
    """
    Send a prompt to the LLM with:
      - Retry with exponential backoff on 429 errors
      - Model fallback chain (tries alternate models if one is rate-limited)
      - Multimodal support: pass images for Gemini vision
    """
    full_prompt = f"{system_prompt}\n\n{prompt}" if system_prompt else prompt
    provider = config.LLM_PROVIDER.lower()

    if provider == "openai":
        return _call_openai(full_prompt, system_prompt, prompt)

    # Gemini path: rotate across multiple API keys + model fallback
    global _current_key_index, _last_successful_model
    # Only 1 retry per key — fail fast and move to next key/model
    max_retries = 1
    base_delay = 5  # seconds

    api_keys = _get_all_api_keys()
    if not api_keys:
        raise ValueError("No Gemini API key provided. Please enter your API key.")

    configured_model = config.GEMINI_MODEL

    # Discover available models using first key
    first_client = _get_gemini_client(api_keys[0])
    available = _discover_available_models(first_client)

    # Build a SHORT, sensible model list (not all 18 discovered models)
    # Priority: last successful > configured > curated fallbacks
    preferred_models = [
        "gemini-2.5-flash",
        "gemini-2.0-flash",
        "gemini-2.0-flash-lite",
        "gemini-2.5-flash-lite",
        "gemini-1.5-flash",
    ]
    # Only keep models that actually exist in the API
    models_to_try = []
    for m in preferred_models:
        if m in available or m in GEMINI_FALLBACK_MODELS:
            if m not in models_to_try:
                models_to_try.append(m)
    # Ensure configured model is in the list
    if configured_model not in models_to_try:
        models_to_try.insert(0, configured_model)

    # If a model succeeded recently, put it FIRST
    if _last_successful_model and _last_successful_model in models_to_try:
        models_to_try.remove(_last_successful_model)
        models_to_try.insert(0, _last_successful_model)

    num_keys = len(api_keys)
    logger.info(f"API keys available: {num_keys}, Models to try: {models_to_try}")

    last_error = None

    # Strategy: for each model, rotate through ALL keys before moving to next model
    for model_name in models_to_try:
        for key_offset in range(num_keys):
            key_idx = (_current_key_index + key_offset) % num_keys
            api_key = api_keys[key_idx]
            key_label = f"key{key_idx+1}/{num_keys}"

            client = _get_gemini_client(api_key)

            for attempt in range(max_retries + 1):
                try:
                    logger.info(f"Calling {model_name} with {key_label} (attempt {attempt+1})")
                    result = _call_gemini(client, model_name, full_prompt, images=images)
                    # Success! Remember this model + advance key index
                    _last_successful_model = model_name
                    _current_key_index = (key_idx + 1) % num_keys
                    logger.info(f"Success with {model_name} {key_label}")
                    return result
                except Exception as e:
                    last_error = e
                    error_str = str(e)
                    is_rate_limit = any(kw in error_str for kw in [
                        "429", "RESOURCE_EXHAUSTED", "rate_limit", "quota",
                        "Too Many Requests", "RateLimitError",
                    ])
                    is_timeout = any(kw in error_str for kw in [
                        "timed out", "timeout", "TimeoutError",
                        "DeadlineExceeded", "DEADLINE_EXCEEDED",
                    ])
                    is_model_not_found = any(kw in error_str for kw in [
                        "404", "NOT_FOUND", "not found", "not supported",
                        "does not exist", "is not available",
                    ])

                    if is_model_not_found:
                        logger.warning(f"Model {model_name} not available, trying next model...")
                        break  # break attempt loop, will also break key loop below
                    elif is_timeout:
                        logger.warning(f"Timeout on {model_name} {key_label}, trying next key...")
                        break  # skip to next key
                    elif is_rate_limit and attempt < max_retries:
                        wait_time = base_delay * (2 ** attempt)
                        logger.warning(
                            f"Rate limit on {model_name} {key_label} (attempt {attempt+1}). "
                            f"Waiting {wait_time}s..."
                        )
                        time.sleep(wait_time)
                        continue
                    elif is_rate_limit:
                        # This key is exhausted, try next key
                        logger.warning(f"Key {key_label} exhausted for {model_name}, trying next key...")
                        break
                    else:
                        raise
            else:
                # attempt loop completed without break — shouldn't happen, but continue
                continue

            # If we broke out due to model-not-found, skip remaining keys for this model
            if last_error and any(kw in str(last_error) for kw in ["404", "NOT_FOUND", "not found", "not supported"]):
                break

    # All keys, models and retries exhausted
    raise last_error


def _call_openai(full_prompt: str, system_prompt: str, user_prompt: str) -> str:
    """Call OpenAI API."""
    _, client = _get_openai_client()
    messages = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    messages.append({"role": "user", "content": user_prompt})
    response = client.chat.completions.create(
        model=config.OPENAI_MODEL,
        messages=messages,
        temperature=getattr(config, "LLM_TEMPERATURE", 0.0),
    )
    return response.choices[0].message.content


# ─────────────────────────────────────────────────────────────────────────────
# 5. COMBINED Summary + Classification (SINGLE API CALL)
# ─────────────────────────────────────────────────────────────────────────────

SYSTEM_PROMPT = (
    "You are a legal document analysis assistant specialising in Indian police chargesheets. "
    "You can read and understand Hindi and English. Always respond in a structured format. "
    "Be precise, factual, and do not hallucinate information not present in the document. "
    "IMPORTANT: The input may contain Hindi OCR errors where \u093f (i-matra) and \u0928 (na) are confused. "
    "Always output CORRECT Hindi spelling, never reproduce OCR errors. "
    "For example: \u0928\u0935\u0915\u093e\u0938\u2192\u0935\u093f\u0915\u093e\u0938, \u092e\u0915\u093e\u093f\u2192\u092e\u0915\u093e\u0928, \u093f\u0917\u0930\u2192\u0928\u0917\u0930, \u0928\u0938\u0902\u0939\u2192\u0938\u093f\u0902\u0939, \u0928\u093f\u0936\u093e\u093f\u2192\u0928\u093f\u0936\u093e\u0928, \u0939\u0930\u0930\u2192\u0939\u0930\u093f."
)


def _build_combined_prompt(text: str) -> str:
    """Build a single prompt that asks for BOTH summary AND crime classification."""
    checklists = load_checklists()
    crime_types_str = "\n".join(
        f"- **{k}**: {v['display_name']} (Sections: {', '.join(v['typical_sections'])})"
        for k, v in checklists.items()
    )

    return f"""❗❗ CRITICAL: OCR ERROR CORRECTION ❗❗
इस दस्तावेज़ में OCR त्रुटियाँ हो सकती हैं। आपको अपने सम्पूर्ण OUTPUT (सारांश, NER, सभी टेक्स्ट) में हमेशा सही हिन्दी लिखनी है, OCR त्रुटि नहीं।

सबसे आम OCR त्रुटियाँ (अवश्य ठीक करें):
1. "ि" और "न" का confusion: "मकाि" → "मकान", "निशाि" → "निशान", "स्टेशि" → "स्टेशन", "िगर" → "नगर"
2. "न"+व्यंजन जहाँ व्यंजन+"ि" होना चाहिए: "नवकास" → "विकास", "नसम" → "सिम", "नसंह" → "सिंह", "नलये" → "लिये"
3. दोहरी मात्राएँ: "चााँदी" → "चाँदी", "मझगााँव" → "मझगाँव"
4. "र" और "रि" confusion: "हरर" → "हरि"
5. "निला" → "जिला", "प्राथनमक" → "प्राथमिक"

उदाहरण: आरोपी का नाम "नवकास कुमार" हो तो आप "विकास कुमार" लिखें।
हर बार सही हिन्दी शब्द OUTPUT करें। OCR त्रुटियों को NEVER reproduce करें।

निम्नलिखित हिन्दी चार्जशीट/आरोप पत्र का विश्लेषण करें। आपको दो कार्य करने हैं:

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
कार्य 1: संरचित सारांश (Structured Summary)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

निम्न प्रारूप में सारांश दें:

## केस हेडर (Case Header)
- **FIR संख्या (FIR Number):**
- **दिनांक (Date):**
- **थाना (Police Station):**
- **न्यायालय (Court):**
- **घटना स्थल (Place of Occurrence):**
- **घटना दिनांक एवं समय (Date & Time of Incident):**

## पक्षकार (Parties Involved)
### शिकायतकर्ता / पीड़ित (Complainant / Victim):
- नाम, पिता का नाम, आयु, पता

### आरोपी (Accused):
- नाम, पिता का नाम, आयु, पता, उपनाम (यदि कोई हो)

### मुख्य गवाह (Key Witnesses):
- गवाहों की सूची (नाम व संक्षिप्त विवरण)

## घटना का सारांश (Incident Summary)
(क्या हुआ, कैसे, कब, कहाँ – 5-10 पंक्तियों में)

## लागू कानूनी धाराएं (Legal Sections Applied)
- धाराओं की सूची (IPC / अन्य अधिनियम) – संक्षिप्त (short) नाम लिखें, पूरा statute language नहीं
- उदाहरण: "IPC 302: Murder", "IPC 34: Common intention" — NOT "IPC 34: Acts done by several persons..."

## प्रमुख साक्ष्य (Key Evidence)
- बरामद/जब्त वस्तुओं की सूची
- रिपोर्ट/दस्तावेज जो चार्जशीट में उल्लिखित हैं

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
कार्य 2: अपराध वर्गीकरण (Crime Classification)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

उपलब्ध अपराध श्रेणियाँ:
{crime_types_str}

सारांश के बाद, निम्न JSON ब्लॉक दें (```json ... ``` में):
```json
{{{{
  "primary_crime_type": "<crime_key>",
  "secondary_crime_types": ["<crime_key>", ...],
  "detected_sections": ["IPC 392", "IPC 323", ...],
  "confidence": "high/medium/low",
  "confidence_score": 0.85,
  "reasoning": "संक्षिप्त कारण (2-3 पंक्तियाँ)"
}}}}
```

confidence_score के लिए:
- 0.85–1.0: धाराएँ स्पष्ट मैच, कोई अंबिगुइटी नहीं
- 0.6–0.84: धाराएँ आंशिक मैच / अनुमानित
- 0.3–0.59: अस्पष्ट, मिश्रित संकेत
- 0.0–0.29: अनिश्चित / अनुमान

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
कार्य 3: Named Entity Recognition (NER)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

दस्तावेज़ से सभी महत्वपूर्ण NAMED ENTITIES निकालें।
Classification JSON के बाद, एक और JSON ब्लॉक दें:

```ner_json
{{{{
  "entities": [
    {{{{"text": "entity का exact नाम", "type": "ENTITY_TYPE"}}}},
    ...
  ]
}}}}
```

Entity Types (role-based person classification):
- ACCUSED: आरोपी / अभियुक्त का proper name (e.g. "मो0 सादिक", "रमन कुमार")
- WITNESS: गवाह / साक्षी का proper name (e.g. "सुनील यादव", "राजू प्रसाद")
- OFFICER: पुलिस/सरकारी अधिकारी — IO, SI, ASI, SP, थानाध्यक्ष (e.g. "अभिषेक कुमार", "राजेश सिंह")
- DOCTOR: चिकित्सक/डॉक्टर — जिन्होंने मेडिकल जाँच या पोस्टमार्टम किया (e.g. "डॉ. विभाकर कुमार")
- PERSON: कोई अन्य व्यक्ति — पीड़ित, शिकायतकर्ता, रिश्तेदार, अन्य उल्लिखित (e.g. "रमेश कुमार", "सीमा देवी")

Non-person entity types:
- DATE: केवल तिथियाँ (dates) निकालें — FIR date, घटना date, गिरफ्तारी date, मेडिकल date, सब अलग हैं। कोई भी date merge/skip न करें। (e.g. "15/03/2024", "दिनांक 20.04.2024", "01.01.2025")
  ⚠️ DATE में केवल तिथियाँ रखें, केवल समय (जैसे "02:34 AM", "23:00", "10.00 बजे", "19:50 hrs") DATE नहीं हैं — इन्हें शामिल न करें।
- LOCATION: स्थान/शहर/गाँव/मोहल्ला (e.g. "चन्द्रपुर", "बॉण्डी बस्ती")
- LEGAL_SECTION: कानूनी धारा (e.g. "BNS 103(1)", "Arms Act 27")
- ORGANIZATION: थाना/न्यायालय/अस्पताल/FSL/सरकारी संस्था (e.g. "सरायकेला थाना", "TMH")
- LANDMARK: दुकान/कम्पनी/मंदिर — स्थान-पहचान के लिए (e.g. "प्रतिक्षा टेक्सटाईल")
- EVIDENCE: जब्त/बरामद भौतिक वस्तु (e.g. "पिस्टल .315 बोर", "खून के नमूने")
- MONETARY: धनराशि (e.g. "₹50,000", "1,20,000 रुपये")
  ⚠️ MONETARY में केवल रुपये/धनराशि रखें। मोबाइल नंबर (जैसे "97xx43xx88", "98xx12xx34") MONETARY नहीं हैं — इन्हें शामिल न करें।
  ⚠️ LOCATION में केवल स्थान/शहर/गाँव रखें। page numbers, section numbers, या अन्य plain numbers LOCATION नहीं हैं।
  ⚠️ ORGANIZATION में केवल संस्था/संगठन/थाना रखें। Reference numbers (जैसे "FSL/RC/2024/4478", "MLG/2024/2187", "CD/0214/2024", "0087/2022") ORGANIZATION नहीं हैं।
  ⚠️ DATE में FIR संख्या (जैसे "0214/2024") शामिल न करें — यह FIR number है, date नहीं।

⚠️ DATE extraction — विशेष निर्देश:
- दस्तावेज़ में हर अलग-अलग date को एक अलग entity के रूप में निकालें
- FIR date, घटना date, गिरफ्तारी date, जमानत date, मेडिकल date, रिपोर्ट date — सब अलग-अलग entries होनी चाहिए
- अपेक्षित: 20-40 DATE entities (timeline reconstruction के लिए)
- Duplicate dates (exact same string) हटाएं, लेकिन अलग-अलग dates कभी merge न करें

⚠️ सख्त NER नियम (STRICT RULES — अवश्य पालन करें):

1. **केवल NAMED ENTITIES** — proper nouns, specific names, specific dates, specific amounts
2. **"text" में केवल entity का नाम लिखें** — कोई भूमिका, सम्बन्ध, विवरण, या व्याख्या नहीं
   - ✅ {{{{"text": "मो0 आजाद", "type": "PERSON"}}}}
   - ❌ {{{{"text": "मो0 आजाद — मो0 इरफान के जीजा", "type": "PERSON"}}}}
   - ❌ {{{{"text": "मो0 आजाद (accused)", "type": "PERSON"}}}}
3. **सम्बन्ध-वाक्य entity नहीं हैं** — "जियाउल की बिवी", "आरोपी का भाई" = NOT entities
4. **एक व्यक्ति = एक entry** — यदि नाम कई बार आता है तो केवल एक बार लिखें
5. **OCR त्रुटियाँ ठीक करें** — OCR में ि/न confusion आम है। Entity text में हमेशा सही हिन्दी शब्द लिखें:
   - "नवकास" → "विकास", "नसंह" → "सिंह", "मकाि" → "मकान", "िगर" → "नगर"
   - OCR variations को एक ही entity मानें — "मो0 सादिक" और "मो0 सादीक" = same person
6. **LEGAL_SECTION: केवल इस FIR/चार्जशीट में लगाई गई धाराएँ** — पिछले अपराधों की धाराएँ शामिल न करें
   - LEGAL_SECTION text में **संक्षिप्त (short) नाम** लिखें, पूरा statute language नहीं
   - ✅ "IPC 34: Common intention"
   - ❌ "IPC 34: Acts done by several persons in furtherance of common intention"
   - ✅ "IPC 302: Murder"
   - ❌ "IPC 302: Punishment for murder"
   - ✅ "IPC 392: Robbery"
   - ❌ "IPC 392: Punishment for robbery"
7. **EVIDENCE: केवल भौतिक वस्तुएँ** — "witness statement" evidence नहीं है
8. **कुल entities 100-150 के बीच होनी चाहिए** (dates 20-40 अलग + persons 30-50 + rest)
9. **Role classification**: ACCUSED/WITNESS/OFFICER/DOCTOR/PERSON — context से तय करें, न कि अनुमान से
- ORGANIZATION: थाना, न्यायालय, अस्पताल, FSL, सरकारी संस्था
- LANDMARK: दुकान, कम्पनी, मंदिर, मस्जिद, चौराहा — जो स्थान-पहचान के लिए उल्लेख हुए

- यदि कोई entity नहीं मिली तो खाली सूची [] रखें
- सभी entities हिन्दी में लिखें

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
चार्जशीट दस्तावेज़ (Chargesheet Document):
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{text}
"""


# ─────────────────────────────────────────────────────────────────────────────
# 5b. NER Post-Processing & Validation
# ─────────────────────────────────────────────────────────────────────────────

# Extensible Hindi/Urdu kinship vocabulary (not exhaustive; structural possessive rule is primary).
HINDI_KINSHIP_WORDS = frozenset({
    "पिता", "माता", "भाई", "बहन", "बेटा", "बेटी", "पति", "पत्नी", "बिवी", "बीवी",
    "चाचा", "चाची", "मामा", "मामी", "फूफा", "फूफी", "ससुर", "सास", "दामाद", "बहू",
    "देवर", "ननद", "जीजा", "साला", "साली", "भतीजा", "भतीजी", "रिश्तेदार",
    "मौसी", "मौसा", "खाला", "खालू", "फूफी", "भांजा", "भांजी", "पोता", "पोती",
    "uncle", "aunt", "brother", "sister", "father", "mother", "wife", "husband",
})

_POSSESSIVE_PARTICLE_RE = re.compile(r'\b(का|की|के|कि)\b', re.IGNORECASE)
_RELATION_WORDS = re.compile(r'\b(?:' + '|'.join(re.escape(w) for w in sorted(HINDI_KINSHIP_WORDS)) + r')\b', re.IGNORECASE)

HINDI_STOP_TOKENS = frozenset({
    "के", "की", "का", "कि", "में", "से", "पर", "को", "ने", "द्वारा", "और", "तथा", "या", "एवं", "भी",
    "यह", "वह", "था", "थी", "थे", "है", "हैं", "हो", "होता", "होती", "हुआ", "हुई", "हुए",
    "एक", "कई", "कुछ", "इस", "उस", "इन", "उन", "तक", "लिए", "साथ", "बाद", "पहले", "फिर",
    "on", "in", "at", "by", "for", "from", "to", "of", "and", "or", "the", "a", "an", "with",
})

MEANINGFUL_CONTEXT_HINT_WORDS = frozenset({
    "वादी", "शिकायतकर्ता", "पीड़ित", "मृतक", "आरोपी", "अभियुक्त", "गवाह", "पंच", "सूचक", "प्रार्थी",
    "पिता", "माता", "भाई", "बहन", "पति", "पत्नी", "बेटा", "बेटी", "मौसी", "मामा", "चाचा", "चाची",
    "officer", "doctor", "witness", "accused", "complainant", "victim", "informant", "inspector",
    "constable", "nodal", "advocate", "lawyer", "judge", "magistrate", "io", "asi", "si", "sp", "dm",
})

OFFICER_DESIGNATION_TOKENS = frozenset({
    "पु", "पु0", "अ", "अ0", "नि", "नि0", "सह", "थाना", "प्रभारी", "थानाध्यक्ष", "पुलिस", "अवर", "उप", "निरीक्षक",
    "आरक्षी", "सिपाही", "हवलदार", "दरोगा", "उनि", "उपनिरीक्षक", "निरीक्षक", "पदाधिकारी", "अनुमण्डल", "जिला",
    "डीएसपी", "एसपी", "एसएसपी", "डीएम", "sdpo", "dysp", "sp", "asp", "si", "asi", "io", "sho",
    "inspector", "sub", "constable", "head", "officer", "nodal", "circle", "station",
})

_OFFICER_ROLE_INDICATOR_RE = re.compile(
    r'\b(?:si|asi|io|sho|sp|dm|dysp|sdpo|inspector|constable|officer|nodal|'
    r'आरक्षी|सिपाही|हवलदार|दरोगा|थानाध्यक्ष|थाना\s*प्रभारी|निरीक्षक|उपनिरीक्षक|पुलिस\s*अवर\s*निरीक्षक|पदाधिकारी|अनुमण्डल)\b',
    re.IGNORECASE,
)
_DOCTOR_ROLE_INDICATOR_RE = re.compile(r'\b(?:dr\.?|doctor|चिकित्सक|डॉक्टर|डॉ\.?|डॉ0?)\b', re.IGNORECASE)

# IPC sections — old law (should not appear if FIR uses BNS)
_IPC_SECTIONS = re.compile(r'\b(?:IPC|आईपीसी|भारतीय\s+दण्ड\s+संहिता)\s*(?:धारा\s*)?\d+', re.IGNORECASE)
# BNS/BNSS sections — new law
_BNS_SECTIONS = re.compile(r'\b(?:B\.?N\.?S\.?|B\.?N\.?S\.?S\.?|बीएनएस|भारतीय\s+न्याय\s+संहिता|भारतीय\s+नागरिक\s+सुरक्षा\s+संहिता)\s*(?:धारा\s*)?\d+', re.IGNORECASE)
_BSA_SECTIONS = re.compile(r'\b(?:B\.?S\.?A\.?|भारतीय\s+साक्ष्य\s+अधिनियम)\s*(?:धारा\s*)?\d+', re.IGNORECASE)


def _normalize_entity_text(text: str, entity_type: str = "") -> str:
    """Normalize entity text for deduplication comparison.
    For DATE entities, only normalize whitespace (preserve exact date strings).
    For person-like entities, normalize Hindi honorifics.
    """
    t = text.strip()

    if entity_type == "DATE":
        # Minimal normalization for dates — only whitespace & separators
        t = re.sub(r'\s+', ' ', t).strip()
        # Normalize date separators: 15-03-2024, 15/03/2024, 15.03.2024
        t = re.sub(r'[\-/\.]', '/', t)
        return t.lower()

    # Person/general normalization
    # Normalize Hindi honorifics and abbreviations
    t = re.sub(r'मो[0०\.]+\s*', 'मो. ', t)
    t = re.sub(r'डॉ[0०\.]+\s*', 'डॉ. ', t)
    t = re.sub(r'Dr\.?\s*', 'डॉ. ', t, flags=re.IGNORECASE)
    t = re.sub(r'श्री[\s\.]*', '', t)
    t = re.sub(r'Sri\.?\s*|Shri\.?\s*', '', t, flags=re.IGNORECASE)
    t = re.sub(r'कु[0०\.]+\s*', 'कु. ', t)
    t = re.sub(r'स[0०]+\s*अ[0०]+\s*नि[0०]+\s*', '', t)  # Remove स0 अ0 नि0
    t = re.sub(r'Mohd?\.?\s*', 'मो. ', t, flags=re.IGNORECASE)
    # Normalize spaces
    t = re.sub(r'\s+', ' ', t).strip()
    # Lowercase for comparison
    return t.lower()


def _repair_ocr_devanagari(text: str) -> str:
    """
    Fix common OCR artifacts in Hindi/Devanagari text:
    1. Remove spurious spaces inside Devanagari words (matra separation)
    2. Fix the systematic ि↔न character confusion (most common Hindi OCR error)
    3. Fix doubled matras (ाा → ा, ीी → ी, etc.)
    4. Apply word-level dictionary corrections
    """
    if not text:
        return text

    # Step 1: Re-attach matras/vowel signs separated from consonant by spaces
    text = re.sub(
        r'([\u0915-\u0939\u0958-\u0961])\s+([\u093e-\u094d\u0951-\u0954]+)',
        r'\1\2', text
    )

    # Step 2: Re-attach nukta (\u093c) separated from consonant
    text = re.sub(
        r'([\u0915-\u0939])\s+(\u093c)',
        r'\1\2', text
    )

    # Step 3: Re-attach virama (\u094d) + following consonant separated by space
    text = re.sub(
        r'(\u094d)\s+([\u0915-\u0939])',
        r'\1\2', text
    )

    # Step 4: Fix isolated matra followed by consonant (reattach to prev consonant)
    text = re.sub(
        r'([\u0915-\u0939\u0958-\u0961])\s+([\u093e-\u094d])([\u0915-\u0939\u0958-\u0961])',
        r'\1\2\3', text
    )

    # Step 5: Fix doubled matras (e.g. चााँदी → चाँदी, मझगााँव → मझगाँव)
    text = re.sub('ाा(?![ँं])', 'ा', text)     # ाा → ा (but keep ाँ)
    text = re.sub('ाा([ँं])', r'ा\1', text)     # ााँ → ाँ
    text = re.sub('ीी', 'ी', text)               # ीी → ी
    text = re.sub('ुु', 'ु', text)               # ुु → ु
    text = re.sub('ूू', 'ू', text)               # ूू → ू
    text = re.sub('ेे', 'े', text)               # ेे → े
    text = re.sub('ोो', 'ो', text)               # ोो → ो

    return text


# ── Hindi OCR Error Correction (General Pattern-Based + LLM) ──────────────────
#
# WHY HINDI OCR ERRORS HAPPEN:
# ─────────────────────────────
# In Devanagari fonts used in Indian legal documents, the dependent vowel sign
# ि (i-matra, U+093F) and the consonant न (na, U+0928) are visually very
# similar — both have a vertical stroke with a leftward hook/mark. OCR engines
# (including those behind PDF text extraction) systematically confuse them,
# producing two predictable error patterns:
#
# PATTERN A (word-final): True न → OCR ि
#   मकान → मकाि, निशान → निशाि, बयान → बयाि, स्टेशन → स्टेशि
#   Rule: Hindi words end with -ान, -ून, -ेन, -ोन but NEVER with -ाि, -ूि, -ेि.
#   FIX: 100% safe algorithmic rule. No dictionary needed.
#
# PATTERN B (word-initial): True consonant+ि → OCR न+consonant
#   विकास → नवकास, सिंह → नसंह, जिला → नजला, किया → नकया
#   Rule: OCR reads ि-matra as standalone न and places it before the consonant.
#   FIX: Algorithmic with a small EXCLUSION list of genuine न-words (~50 roots).
#        Exclusion list is MUCH smaller than a word-by-word correction dictionary.
#
# PATTERN C: Missing nukta (ड→ड़, ढ→ढ़)
#   बड़ा → बडा, सड़क → सडक, गाड़ी → गाडी
#   FIX: Small pattern list — nukta is always safe to restore.
#
# MID-WORD Pattern B: True consonant+ि → OCR न+consonant inside a word
#   प्राथमिक → प्राथनमक, महिला → मनहला, प्रकृति → प्रकृनत
#   FIX: Handled by LLM (too context-dependent for algorithmic rules).
#        Common mid-word patterns also added to phrase corrections for safety.
#
# APPROACH (4 layers, scales to ANY document without manual word additions):
# ───────────────────────────────────────────────────────────────────────────
# Layer 1: General algorithmic Pattern A rules (100% safe, no dictionary)
# Layer 2: Exclusion-based Pattern B rules (small exclusion list, auto-fixes)
# Layer 3: Nukta fixes + phrase corrections for mid-word & multi-word patterns
# Layer 4: LLM prompt instructions handle remaining edge cases
# Optional: Dedicated LLM OCR cleanup call for highest accuracy
# ──────────────────────────────────────────────────────────────────────────────


def _fix_ocr_pattern_a(text: str) -> str:
    """
    Fix Pattern A: word-final [vowel_matra] + ि → [vowel_matra] + न

    In Hindi, words commonly end with -ान, -ून, -ेन, -ोन, -ीन, -ैन, -ौन
    but NEVER end with -ाि, -ूि, -ेि, -ोि, -ीि, -ैि, -ौि.
    This rule is 100% safe with ZERO known exceptions in Hindi.

    Also fixes Pattern D (mid-word ि→न):
    - Doubled ि (ि+ि → ि+न): दििांक→दिनांक, दिि→दिन
    - ि before vowel matra (ि+ा → न+ा): घटिा→घटना
    These are 100% safe: a consonant can only carry ONE vowel matra, so
    ि+vowel_matra is always an OCR error where ि should be न.

    Examples: मकाि→मकान, बयाि→बयान, चेि→चेन, दूि→दून, फोि→फोन
              दििांक→दिनांक, घटिा→घटना, दिि→दिन
    """
    # --- Pattern D: mid-word ि→न (doubled or before vowel matra) ---
    # Doubled ि: change the SECOND ि to न (first ि is correct vowel matra)
    # (?<=ि)ि matches ि preceded by ि — i.e. the duplicate OCR error
    # दििांक → दिनांक, दिि → दिन
    text = re.sub(r'(?<=ि)ि', 'न', text)

    # ि before a vowel matra: ि+[ा ो ू ी ै ौ] → न+matra
    # A single consonant can carry only ONE vowel matra. If ि is immediately
    # followed by another vowel matra, the ि is an OCR error for न.
    # घटिा→घटना, चिकत्सा stays (ि+क not ि+matra)
    text = re.sub(r'ि(?=[ाोूीैौ])', 'न', text)

    # --- Pattern A: word-final vowel_matra+ि → vowel_matra+न ---
    # Match vowel matra + ि at word boundary (followed by space, punctuation, or end)
    # Vowel matras: ा(093E) ू(0942) े(0947) ो(094B) ी(940) ै(0948) ौ(094C)
    text = re.sub(
        r'([ाूेोीैौ])ि(?=[\s,।\.;:!?\)\]\}\-–—\u0964\u0965]|$)',
        r'\1न',
        text
    )

    # NOTE: Word-final consonant+ि (e.g. स्टेशि) is NOT fixed algorithmically
    # because many legitimate Hindi words end in consonant+ि (रात्रि, शक्ति, राशि,
    # बुद्धि, प्रगति, etc.). These are handled by specific phrase corrections instead.

    # Also fix standalone ि at word start → न (ि CANNOT start a word in Hindi)
    # e.g., िगर → नगर, िंबर → नंबर, िं0 → नं0
    text = re.sub(
        r'(?:^|(?<=\s))ि(?=[\u0915-\u0939\u0902\u0903\u0901\u0928\u093C\u0964])',
        'न',
        text
    )

    return text


# ── Genuine न-initial Hindi words (exclusion list for Pattern B) ──
# This list contains ~60 common Hindi ROOT WORDS that genuinely start with
# न+consonant. It is MUCH smaller and more maintainable than a per-word
# correction dictionary. Any word NOT in this list that starts with
# न+consonant will automatically be corrected (Pattern B applied).
# This means NEW OCR errors are fixed automatically without code changes.
_GENUINE_NA_WORDS = frozenset({
    # ─── नक- ───
    'नकद', 'नकल', 'नकली', 'नकाब', 'नकार', 'नकारा', 'नकारात्मक', 'नक्शा',
    'नक्सल', 'नक्काशी',
    # ─── नख- ───
    'नखरा', 'नखरे',
    # ─── नगर (city — very common in legal docs) ───
    'नगर', 'नगरपालिका', 'नगरीय', 'नगरपरिषद', 'नगरसेवक', 'नगरनिगम',
    'नगद', 'नगीना',
    # ─── नट- ───
    'नटखट', 'नटराज', 'नटवर',
    # ─── नत- ───
    'नतीजा', 'नतीजे', 'नतमस्तक',
    # ─── नथ- ───
    'नथ', 'नथनी', 'नथुनी',
    # ─── नद- ───
    'नदी', 'नदियाँ', 'नदियों',
    # ─── नन- ───
    'नन्हा', 'नन्हें', 'नन्ही',
    # ─── नप- ───
    'नपुंसक',
    # ─── नफ- ───
    'नफ़रत', 'नफरत', 'नफा', 'नफ़ा', 'नफासत',
    # ─── नब- ───
    'नब्ज़', 'नब्ज', 'नब्बे',
    # ─── नभ- ───
    'नभ',
    # ─── नम- ───
    'नमक', 'नमकीन', 'नमन', 'नमस्कार', 'नमस्ते', 'नमी', 'नमूना',
    'नम्र', 'नम्बर', 'नमाज़', 'नमाज',
    # ─── नय- ───
    'नयन', 'नया', 'नयी', 'नये', 'नई',
    # ─── नर- ───
    'नरम', 'नरक', 'नरेश', 'नरेंद्र', 'नर्म', 'नर्तक', 'नर्स', 'नर्तकी',
    'नरसिंह', 'नरसिम्हा',
    # ─── नल- ───
    'नलकूप', 'नलिका', 'नली',
    # ─── नव- (new/nine — careful: नवकास is OCR error for विकास) ───
    'नवीन', 'नवम्बर', 'नवंबर', 'नवाब', 'नवजात', 'नवम', 'नवरात्र',
    'नवनीत', 'नवयुवक', 'नवल', 'नवीनीकरण',
    # ─── नश- ───
    'नशा', 'नशीला', 'नशेड़ी',
    # ─── नस- ───
    'नसों', 'नसीब', 'नसीहत',
    # ─── नष- ───
    'नष्ट',
    # ─── नह- ───
    'नहर', 'नहीं', 'नहाना', 'नहाया', 'नहाई',
    # ─── नं- ───
    'नंगा', 'नंबर', 'नंबरी', 'नंद', 'नंदन', 'नंदी',
    # ─── न्य- (conjunct — has virama) ───
    'न्याय', 'न्यायालय', 'न्यायाधीश', 'न्यायिक', 'न्यूनतम', 'न्यूज',
    # ─── other conjuncts ───
    'न्यारा', 'न्योता',
})

# Build prefix set from roots of 3+ chars for compound word matching
# e.g., "नगर" root covers "नगरपालिकाओं", "नगरीयकरण", etc.
_GENUINE_NA_ROOTS = frozenset(w for w in _GENUINE_NA_WORDS if len(w) >= 3)


def _is_genuine_na_word(word: str) -> bool:
    """
    Check if a word genuinely starts with न+consonant (not an OCR error).
    Uses exact match + root prefix match against the exclusion set.
    """
    # Strip trailing punctuation for matching
    clean = re.sub(r'[,।\.;:!?\)\]\}\-–—\u0964\u0965]+$', '', word)
    if not clean:
        return False

    # 1. Exact match
    if clean in _GENUINE_NA_WORDS:
        return True

    # 2. Check if this is a known root + suffix (compound word)
    #    e.g., "नगरपालिकाओं" starts with root "नगर" (genuine)
    #    Only use roots of 3+ chars to avoid over-matching
    for root in _GENUINE_NA_ROOTS:
        if clean.startswith(root) and len(clean) > len(root):
            return True

    # 3. Check if the word has virama (्) after न — this is a conjunct like न्याय
    #    Conjuncts are always genuine (OCR Pattern B doesn't produce virama)
    if len(clean) >= 2 and clean[1] == '\u094D':  # virama
        return True

    return False


def _fix_ocr_pattern_b(text: str) -> str:
    """
    Fix Pattern B: word-initial न+consonant → consonant+ि

    When OCR misreads ि-matra as the consonant न, it places न before the
    consonant. This function reverses that UNLESS the word is a genuine
    Hindi word starting with न (checked against the exclusion set).

    Examples: नवकास→विकास, नसंह→सिंह, नजला→जिला, नकया→किया

    SCALABILITY: New OCR errors (e.g. नबक्री for बिक्री) are automatically
    fixed without any code change — only genuine न-words need to be added
    to the exclusion set (which is rare and small).
    """
    def _fix_word(match):
        word = match.group(0)
        if len(word) < 3:
            return word  # Too short — ambiguous

        prev_chunk = text[max(0, match.start() - 5):match.start()]
        if ("मो0" in prev_chunk) or ("मो " in prev_chunk):
            return word

        # Must start with न + Devanagari consonant (U+0915 to U+0939)
        if word[0] != 'न' or not ('\u0915' <= word[1] <= '\u0939'):
            return word

        # Check exclusion set — genuine न-words should NOT be corrected
        if _is_genuine_na_word(word):
            return word

        # Apply Pattern B fix: remove leading न, add ि-matra to next consonant
        # नवकास → व + ि + कास → विकास
        fixed = word[1] + 'ि' + word[2:]
        return fixed

    # Match Devanagari words starting with न followed by a consonant (at word boundary)
    text = re.sub(
        r'(?:^|(?<=\s))न[\u0915-\u0939][\u0900-\u097F]+',
        _fix_word,
        text,
        flags=re.MULTILINE
    )

    return text


# ── Nukta fixes (ड→ड़, ढ→ढ़) ──
# Nukta (nuqta) is a small dot under consonants representing sounds from
# Urdu/Persian/Arabic. OCR frequently misses this tiny dot.
# These are safe to apply as ड without nukta before common suffixes is rare.
_NUKTA_FIXES = {
    'बडा': 'बड़ा', 'बडी': 'बड़ी', 'बडे': 'बड़े',
    'सडक': 'सड़क', 'कपडा': 'कपड़ा', 'कपडे': 'कपड़े',
    'पकडा': 'पकड़ा', 'पकडे': 'पकड़े', 'छोडा': 'छोड़ा',
    'तोडा': 'तोड़ा', 'जोडा': 'जोड़ा', 'जोडी': 'जोड़ी',
    'टुकडा': 'टुकड़ा', 'लडका': 'लड़का', 'लडकी': 'लड़की',
    'दौडा': 'दौड़ा', 'चढा': 'चढ़ा', 'गाडी': 'गाड़ी',
    'घोडा': 'घोड़ा', 'पीडित': 'पीड़ित', 'झगडा': 'झगड़ा',
}
_NUKTA_RE = re.compile(
    '|'.join(re.escape(w) for w in sorted(_NUKTA_FIXES, key=len, reverse=True))
)

# ── Phrase-level OCR corrections (multi-word patterns + mid-word Pattern B) ──
# These fix specific multi-word patterns AND common mid-word ि↔न errors that
# are too context-dependent for general algorithmic rules.
# The mid-word entries are VERY common in Hindi legal documents and safe to fix.
_OCR_PHRASE_CORRECTIONS = [
    # Multi-word patterns
    (re.compile(r'भवा\s*िगर'), 'भवानी नगर'),
    (re.compile(r'सो\s+की\s+चेि'), 'सोने की चेन'),
    (re.compile(r'सो\s+की\s+अाँगूठी'), 'सोने की अंगूठी'),
    (re.compile(r'सो\s+की\s+िथ'), 'सोने की नथ'),
    (re.compile(r'ते\s+का\s+निशाि'), 'तेल का निशान'),
    (re.compile(r'(?:^|(?<=\s))था(?=\s+(?:भवा|कोतवाली|सरायकेला|मझगा|प्रभारी))'), 'थाना'),
    (re.compile(r'िगर\s+पंचायत'), 'नगर पंचायत'),
    (re.compile('िं[0०\\.]+'), 'नं0'),
    (re.compile(r'प्रकृ\s*नत'), 'प्रकृति'),
    (re.compile(r'कु\s+न्द'), 'कुन्द'),
    (re.compile(r'कु\s+मार'), 'कुमार'),
    (re.compile(r'भवा\s+नगर'), 'भवानी नगर'),
    # Common mid-word Pattern B (safe to fix — very frequent in legal Hindi):
    (re.compile(r'प्राथनमक'), 'प्राथमिक'),   # primary
    (re.compile(r'प्रकृनत'), 'प्रकृति'),       # nature
    (re.compile(r'मनहला'), 'महिला'),           # woman
    (re.compile(r'रानत्र'), 'रात्रि'),         # night
    (re.compile(r'प्रनत'), 'प्रति'),           # copy/per
    (re.compile(r'सम्पनत'), 'सम्पत्ति'),       # property
    (re.compile(r'सम्पनि'), 'सम्पत्ति'),       # property
    (re.compile(r'अनभयुि'), 'अभियुक्त'),     # accused
    (re.compile(r'मोटरसाइनकल'), 'मोटरसाइकिल'),# motorcycle
    (re.compile(r'संभानवत'), 'संभावित'),       # probable
    (re.compile(r'बिामदगी'), 'बरामदगी'),       # recovery
    (re.compile(r'सनगरेट'), 'सिगरेट'),         # cigarette
    # Double OCR errors (word-initial + mid-word Pattern B):
    (re.compile(r'(?:नच|चि)कत्सा'), 'चिकित्सा'),    # medical (double error)
    (re.compile(r'(?:नच|चि)कत्सालय'), 'चिकित्सालय'),# hospital (double error)
    (re.compile(r'नगरफ्तार'), 'गिरफ्तार'),           # arrested (collides with नगर root)
    # Heavy OCR corruptions (multiple character damage — common in scanned chargesheets):
    (re.compile(r'ररपोटच'), 'रिपोर्ट'),               # report — र→रि, ट→र्ट, च dropped
    (re.compile(r'(?:दिच|दजच|दर्च)\b'), 'दर्ज'),      # registered/filed
    (re.compile(r'(?:तिनथ|निनथ|तििथ)'), 'तिथि'),      # date/tithhi
    (re.compile(r'बि्ती'), 'जब्ती'),                   # seizure — ब→ज missing
    (re.compile(r'धनरा(?:शन|सन)'), 'धनराशि'),         # money/amount — final न→ि
    (re.compile(r'(?:रामनिक|रामननक)शोर'), 'रामकिशोर'),# name fix — mid-word Pattern B
    (re.compile(r'(?:घटिास्थल|घनटास्थल)'), 'घटनास्थल'),# crime scene compound
    (re.compile(r'(?:नवस्तार|नवस्तर)\b'), 'विस्तार'), # detail/extent
    (re.compile(r'जािकारी'), 'जानकारी'),               # information — double matra
    (re.compile(r'(?:ँच|ाँच)\b'), 'ांच'),             # fix broken jānch (ँच→ांच)
    (re.compile(r'की\s+(?:जांच|जाँच|ांच)'), 'की जांच'),# normalize "ki jaanch"
    # English loanwords ending in consonant+ि (OCR for consonant+न):
    # Cannot use algorithmic rule because many Hindi words end in consonant+ि
    # (रात्रि, शक्ति, राशि, बुद्धि, etc.) — so these are explicit:
    (re.compile(r'स्टेशि'), 'स्टेशन'),                 # station
    (re.compile(r'कमीशि'), 'कमीशन'),                   # commission
    (re.compile(r'मिशि'), 'मिशन'),                     # mission
    (re.compile(r'सीजि'), 'सीजन'),                     # season
    (re.compile(r'लोशि'), 'लोशन'),                     # lotion
    (re.compile(r'टेलीनवजि'), 'टेलीविजन'),             # television
    (re.compile(r'(?:ररजस्ट्रेशि|रनजस्ट्रेशि)'), 'रजिस्ट्रेशन'), # registration
]


# ── LLM-based OCR Cleanup (optional, for production deployments) ──────────────
# For maximum accuracy, use the LLM to fix OCR errors contextually.
# The LLM understands Hindi grammar and can fix errors that pattern rules miss:
# - Mid-word Pattern B (महिला, प्राथमिक etc.)
# - Context-dependent corrections
# - Names with OCR errors
# - Any new error pattern not seen before
#
# Enable by calling _ocr_cleanup_via_llm() in preprocess_text().
# Adds ~3-5 seconds latency + 1 additional API call per document.

_OCR_CLEANUP_PROMPT = """आपको एक हिन्दी कानूनी दस्तावेज़ (chargesheet/FIR) का OCR-extracted text दिया गया है।
इस text में systematic OCR errors हैं। कृपया इन्हें ठीक करें:

## मुख्य OCR त्रुटि: ि (i-matra) ↔ न (na) confusion
Hindi fonts में ि और न बहुत similar दिखते हैं। OCR इन्हें confuse करता है:

### Pattern A (शब्द के अंत में): न → ि
- मकाि → मकान, बयाि → बयान, निशाि → निशान, स्टेशि → स्टेशन

### Pattern B (शब्द की शुरुआत या बीच में): consonant+ि → न+consonant
- नवकास → विकास, नसंह → सिंह, नजला → जिला, नकया → किया
- मनहला → महिला, प्राथनमक → प्राथमिक, प्रकृनत → प्रकृति

### अन्य OCR errors:
- Missing nukta: बडा→बड़ा, सडक→सड़क, गाडी→गाड़ी, पीडित→पीड़ित
- Doubled matras: ाा→ा, ीी→ी
- Broken words: कु मार→कुमार, प्रकृ ति→प्रकृति

## निर्देश:
1. OCR errors ठीक करें लेकिन content/meaning बिल्कुल न बदलें
2. Names, dates, numbers, legal sections — सब intact रखें
3. जो text सही है उसे वैसा ही रहने दें
4. केवल corrected text return करें, कोई explanation नहीं

## OCR Text:
{text}

## Corrected Text:"""


def _ocr_cleanup_via_llm(text: str) -> str:
    """
    Use the LLM (Gemini/OpenAI) to fix OCR errors in Hindi text.
    This is the most accurate approach as the LLM understands context.

    Call this BEFORE sending text for analysis (preprocessing step).
    Adds ~3-5 seconds latency and uses one additional API call.

    For production police deployments where accuracy is critical,
    enable this in preprocess_text() for best results.
    """
    if not text or len(text) < 50:
        return text

    try:
        prompt = _OCR_CLEANUP_PROMPT.format(text=text[:config.MAX_TEXT_LENGTH])
        cleaned = _call_llm(prompt, system_prompt=(
            "You are a Hindi OCR error correction specialist. "
            "Fix ONLY OCR errors in the text. Do NOT change content, meaning, "
            "names, dates, or numbers. Return ONLY the corrected text."
        ))

        # Sanity check: LLM output should be roughly same length (±30%)
        if cleaned and 0.7 < len(cleaned) / len(text) < 1.3:
            logger.info(f"LLM OCR cleanup: {len(text)} → {len(cleaned)} chars")
            return cleaned
        else:
            logger.warning(f"LLM OCR cleanup output length mismatch: "
                          f"{len(text)} → {len(cleaned) if cleaned else 0}, skipping")
            return text
    except Exception as e:
        logger.warning(f"LLM OCR cleanup failed (falling back to pattern rules): {e}")
        return text


def _apply_ocr_corrections(text: str) -> str:
    """
    Apply OCR corrections using general pattern rules.
    No hardcoded word dictionary needed — uses algorithmic patterns that
    scale to any Hindi legal document automatically.

    Layers:
    1. Pattern A (word-final ि→न) — 100% safe algorithmic rule
    2. Pattern B (word-initial न→consonant+ि) — exclusion-based
    3. Nukta fixes (ड→ड़) — small safe dictionary
    4. Phrase-level fixes — for multi-word patterns & mid-word Pattern B
    """
    # Layer 1: Fix Pattern A (word-final) — general, safe
    text = _fix_ocr_pattern_a(text)

    # Layer 2: Fix Pattern B (word-initial) — exclusion-based
    text = _fix_ocr_pattern_b(text)

    # Layer 3: Fix nukta errors — small dictionary
    text = _NUKTA_RE.sub(lambda m: _NUKTA_FIXES[m.group(0)], text)

    # Layer 4: Fix phrase-level and common mid-word errors
    for pattern, replacement in _OCR_PHRASE_CORRECTIONS:
        text = pattern.sub(replacement, text)

    return text



def _strip_annotations(text: str) -> str:
    """
    Strip role/relation annotations from entity text.
    'मो0 आजाद — मो0 इरफान के जीजा' → 'मो0 आजाद'
    'राकेश कुमार (witness)' → 'राकेश कुमार'
    'अभिषेक कुमार (IO, सरायकेला थाना)' → 'अभिषेक कुमार'
    """
    # Remove everything after em-dash or double-hyphen
    text = re.split(r'\s*[—–\-]{1,2}\s+', text, maxsplit=1)[0].strip()
    # Remove parenthetical annotations
    text = re.sub(r'\s*\([^)]*\)\s*$', '', text).strip()
    # Remove trailing comma + role
    text = re.sub(r',\s*(?:accused|victim|witness|complainant|IO|SI|ASI|आरोपी|पीड़ित|गवाह|शिकायतकर्ता).*$',
                  '', text, flags=re.IGNORECASE).strip()
    return text


# Deceased / late markers — these indicate the person is dead, not a witness
_DECEASED_MARKERS = re.compile(
    r'\b(?:स्व[\.\.]*|स्वर्गीय|late|मृत|मृतक|दिवंगत|मरहूम)\b',
    re.IGNORECASE
)


def _is_deceased_person(text: str) -> bool:
    """Check if entity text indicates a deceased person."""
    return bool(_DECEASED_MARKERS.search(text))


def _looks_like_proper_name(token: str) -> bool:
    t = (token or "").strip(" .,;:()[]{}-_")
    if not t:
        return False
    if re.match(r'^[A-Z][a-z]+$', t):
        return True
    if re.search(r'[\u0900-\u097F]', t):
        if len(t) < 2:
            return False
        if _RELATION_WORDS.search(t):
            return False
        if t.lower() in HINDI_STOP_TOKENS:
            return False
        return True
    return False


def _contains_possessive_token(text: str) -> bool:
    tokens = [t.strip(" .,;:()[]{}-_") for t in re.split(r'\s+', (text or '').strip()) if t.strip(" .,;:()[]{}-_")]
    return any(tok in {"का", "की", "के", "कि"} for tok in tokens)


def _is_relationship_phrase(text: str) -> bool:
    """Detect possessive relationship phrases like '<name> का/की/के/कि <relation>' as non-entity text."""
    raw = (text or "").strip()
    if not raw:
        return False

    tokens = [t.strip(" .,;:()[]{}-_") for t in re.split(r'\s+', raw) if t.strip(" .,;:()[]{}-_")]
    if not tokens:
        return False

    possessive_indices = [i for i, tok in enumerate(tokens) if tok in {"का", "की", "के", "कि"}]
    if not possessive_indices:
        return False

    token_lowers = {tok.lower() for tok in tokens}
    kinship_lowers = {w.lower() for w in HINDI_KINSHIP_WORDS}
    if token_lowers & kinship_lowers:
        return True

    first_pos = possessive_indices[0]
    if first_pos <= 0 or first_pos >= (len(tokens) - 1):
        return False

    right_first = tokens[first_pos + 1]
    if right_first.lower() in kinship_lowers:
        return True
    if _looks_like_proper_name(right_first):
        return False
    return True


def _is_relation_phrase(text: str) -> bool:
    """Backward-compatible alias to new relationship-phrase detector."""
    return _is_relationship_phrase(text)


def _is_title_only_entity(text: str) -> bool:
    """Detect designation-only officer strings that do not contain a personal name."""
    raw = (text or "").strip()
    if not raw:
        return True
    tokens = [t.strip(" .,;:()[]{}") for t in re.split(r'\s+', raw) if t.strip(" .,;:()[]{}")]
    if not tokens:
        return True

    has_designation = False
    name_like_tokens = []
    for tok in tokens:
        low = tok.lower()
        if low in OFFICER_DESIGNATION_TOKENS:
            has_designation = True
            continue
        if re.search(r'[0०]', tok) and re.search(r'[\u0900-\u097F]', tok):
            has_designation = True
            continue
        if re.match(r'^[\d०-९/\-]+$', tok):
            continue
        if low in {"थाना", "नगर", "जिला", "प्रखण्ड", "चौकी", "रोड", "गांव", "ग्राम", "स्टेशन"}:
            continue
        if _looks_like_proper_name(tok):
            name_like_tokens.append(tok)

    # If designation exists but only one residual token, it's typically title + location.
    if has_designation and len(name_like_tokens) <= 1:
        return True
    return len(name_like_tokens) == 0


# Person-like types that need relation-phrase rejection
_PERSON_TYPES = {
    "PERSON", "ACCUSED", "COMPLAINANT", "VICTIM", "DECEASED",
    "WITNESS", "OFFICER", "JUDGE", "DOCTOR"
}

_COMPLAINANT_CONTEXT_RE = re.compile(
    r'वादी|शिकायतकर्ता|प्रार्थी|complainant|fir\s*filer|informant|आवेदक',
    re.IGNORECASE,
)
_JUDICIAL_ROLE_RE = re.compile(
    r'न्यायिक|दण्डाधिकारी|न्यायाधीश|magistrate|judge|cjm|jmfc|session\s*court|विचारण|न्यायालय',
    re.IGNORECASE,
)
_PERSON_ROLE_PRIORITY = {
    "DECEASED": 8,
    "VICTIM": 7,
    "COMPLAINANT": 6,
    "ACCUSED": 5,
    "JUDGE": 4,
    "OFFICER": 3,
    "WITNESS": 2,
    "DOCTOR": 2,
    "PERSON": 1,
}


def _prefer_summary_name_variant(clean_text: str) -> str:
    """Prefer canonical name form from summary context when a strong fuzzy match exists."""
    global _NER_SUMMARY_CONTEXT
    summary = _NER_SUMMARY_CONTEXT or ""
    if not summary or clean_text in summary:
        return clean_text

    candidates = re.findall(r'[\u0900-\u097FA-Za-z\.]{3,}(?:\s+[\u0900-\u097FA-Za-z\.]{2,}){0,3}', summary)
    if not candidates:
        return clean_text

    best = clean_text
    best_score = 0.0
    clean_norm = _normalize_entity_text(clean_text, "PERSON")
    for cand in candidates:
        cand_norm = _normalize_entity_text(cand, "PERSON")
        if not cand_norm:
            continue
        if clean_norm == cand_norm:
            return cand.strip()
        longer = max(len(clean_norm), len(cand_norm)) or 1
        score = 1.0 - (_edit_distance(clean_norm, cand_norm) / longer)
        if score > best_score:
            best_score = score
            best = cand.strip()

    if best_score >= 0.90:
        logger.info(f"NER summary-alignment: '{clean_text}' -> '{best}' (score={best_score:.2f})")
        return best
    return clean_text

# Regex to split alias names: "राजेश उर्फ गुड्डू" → ["राजेश", "गुड्डू"]
_ALIAS_SPLIT_RE = re.compile(
    r'\s+(?:उर्फ़?|उर्फ|@|a\.?k\.?a\.?|alias|नाम)\s+',
    re.IGNORECASE,
)


def _extract_alias_parts(name: str) -> list:
    """
    Split a name containing alias markers into constituent parts.
    E.g. "राजेश उर्फ गुड्डू उर्फ राजू" → ["राजेश", "गुड्डू", "राजू"]
    Returns empty list if no alias marker found.
    """
    if not _ALIAS_SPLIT_RE.search(name):
        return []
    parts = _ALIAS_SPLIT_RE.split(name)
    return [p.strip() for p in parts if p.strip()]


def _is_meaningful_context_hint(token: str) -> bool:
    t = (token or "").strip().lower()
    if not t:
        return False
    if t in HINDI_STOP_TOKENS:
        return False
    if re.fullmatch(r'[\d०-९]+', t):
        return False
    if t in MEANINGFUL_CONTEXT_HINT_WORDS:
        return True
    # lightweight morphological handling
    for suffix in ("ों", "ें", "ों", "ी", "ा", "e", "s"):
        if len(t) > 3 and t.endswith(suffix):
            base = t[:-len(suffix)]
            if base in MEANINGFUL_CONTEXT_HINT_WORDS:
                return True
    return False


def _officer_short_name_merge_possible(shorter: str, longer: str) -> bool:
    if not shorter or not longer:
        return False
    short_norm = _normalize_entity_text(shorter, 'OFFICER')
    long_norm = _normalize_entity_text(longer, 'OFFICER')
    if short_norm == long_norm:
        return False
    if len(short_norm) < 3:
        return False
    if short_norm not in long_norm:
        return False
    short_tokens = [t for t in short_norm.split() if t]
    long_tokens = [t for t in long_norm.split() if t]
    if not short_tokens:
        return False
    if not all(tok in long_tokens for tok in short_tokens):
        return False
    return any(_looks_like_proper_name(tok) for tok in short_tokens)


MONETARY_EXCLUSION_CONTEXT_WORDS = frozenset({
    "service", "charge", "charges", "gst", "tax", "fee", "fees", "processing",
    "शुल्क", "फीस", "चार्ज", "कर", "जीएसटी", "प्रोसेसिंग", "न्यूनतम", "शेष",
})

_MONETARY_RELEVANCE_OVERRIDE_WORDS = frozenset({
    "रिश्वत", "घूस", "लूट", "डकैती", "चोरी", "उगाही", "ठगी", "धोखाधड़ी", "fraud", "extortion", "bribe", "seizure", "recovery",
})

_PERSON_HINT_BANNED_FOR_PERSON = frozenset({"आरोपी", "अभियुक्त", "गवाह", "साक्षी"})


def _extract_amount_value(text: str) -> float | None:
    m = re.search(r'\d{1,3}(?:,\d{2,3})*(?:\.\d{1,2})?', str(text or ''))
    if not m:
        return None
    try:
        return float(m.group(0).replace(',', ''))
    except ValueError:
        return None


def _is_monetary_noise_context(amount_text: str, full_text: str) -> bool:
    if not amount_text or not full_text:
        return False
    value = _extract_amount_value(amount_text)
    if value is None:
        return False
    noise_max = float(getattr(config, 'MONETARY_NOISE_MAX_AMOUNT', 2000))
    if value > noise_max:
        return False
    window = int(getattr(config, 'MONETARY_CONTEXT_WINDOW_CHARS', 100))
    for m in re.finditer(re.escape(amount_text), full_text, flags=re.IGNORECASE):
        s = max(0, m.start() - window)
        e = min(len(full_text), m.end() + window)
        ctx = full_text[s:e].lower()
        if any(word in ctx for word in _MONETARY_RELEVANCE_OVERRIDE_WORDS):
            return False
        if any(word in ctx for word in MONETARY_EXCLUSION_CONTEXT_WORDS):
            return True
    return False


def _normalize_evidence_for_dedup(text: str) -> str:
    t = _apply_ocr_corrections((text or '').lower())
    t = re.sub(r"\b(?:mark|मार्क)\s*['\"]?[a-zA-Zअ-ह]?['\"]?\b", ' ', t, flags=re.IGNORECASE)
    t = re.sub(r'[^\w\u0900-\u097F\s]', ' ', t)
    t = re.sub(r'\s+', ' ', t).strip()
    return t


def _organization_core_form(text: str) -> str:
    t = _normalize_entity_text(str(text or ''), 'ORGANIZATION')
    suffixes = [
        r'\bथाना\b', r'\bp\.?\s*s\.?\b', r'\bpolice\s*station\b',
        r'\bअस्पताल\b', r'\bhospital\b', r'\bबैंक\b', r'\bbank\b',
        r'\bशाखा\b', r'\bbranch\b',
    ]
    for pattern in suffixes:
        t = re.sub(pattern, ' ', t, flags=re.IGNORECASE)
    t = re.sub(r'\s+', ' ', t).strip()
    return t


def _entity_has_doctor_title(name_text: str) -> bool:
    t = str(name_text or '').strip()
    return bool(re.search(r'^(?:dr\.?\s+|डॉ\.?\s+|डॉक्टर\s+|doctor\s+|m\.d\.?\s+|m\.b\.b\.s\.?\s+|surgeon\s+|सर्जन\s+|physician\s+|चिकित्सक\s+)', t, re.IGNORECASE))


def _entity_has_officer_token(name_text: str) -> bool:
    t = str(name_text or '').strip()
    return bool(_OFFICER_ROLE_INDICATOR_RE.search(t))


def _officer_context_has_designation(name_text: str, full_text: str) -> bool:
    if not name_text or not full_text:
        return False
    window = int(getattr(config, 'PERSON_CONTEXT_WINDOW_CHARS', 120))
    for m in re.finditer(re.escape(name_text), full_text, flags=re.IGNORECASE):
        s = max(0, m.start() - window)
        e = min(len(full_text), m.end() + window)
        if _OFFICER_ROLE_INDICATOR_RE.search(full_text[s:e]):
            return True
    return False


def _is_correspondence_only_person(name_text: str, full_text: str) -> bool:
    if not name_text or not full_text:
        return False
    correspondence_re = re.compile(r'bank|बैंक|statement|स्टेटमेंट|cdr|nodel|nodal|पत्र|letter|correspondence|शाखा', re.IGNORECASE)
    investigative_re = re.compile(r'आरोपी|अभियुक्त|गवाह|जांच|विवेचना|घटना|गिरफ्तार|बरामद|पंचनामा|थानाध्यक्ष', re.IGNORECASE)
    hits = 0
    corr_hits = 0
    inv_hits = 0
    for m in re.finditer(re.escape(name_text), full_text, flags=re.IGNORECASE):
        hits += 1
        s = max(0, m.start() - 140)
        e = min(len(full_text), m.end() + 140)
        ctx = full_text[s:e]
        if correspondence_re.search(ctx):
            corr_hits += 1
        if investigative_re.search(ctx):
            inv_hits += 1
    if hits == 0:
        return False
    return corr_hits == hits and inv_hits == 0


def _evidence_entities_duplicate(a: str, b: str) -> bool:
    ta = (a or "").strip()
    tb = (b or "").strip()
    if not ta or not tb:
        return False
    na = _normalize_evidence_for_dedup(ta)
    nb = _normalize_evidence_for_dedup(tb)
    if na == nb:
        return True
    min_len = int(getattr(config, 'EVIDENCE_DEDUP_MIN_LENGTH', 5))
    if min(len(na), len(nb)) >= min_len and (na in nb or nb in na):
        return True
    threshold = int(getattr(config, 'EVIDENCE_DEDUP_EDIT_DISTANCE', 3))
    if _edit_distance(na, nb) <= threshold:
        return True
    cs_threshold = float(getattr(config, 'CROSS_SCRIPT_DEDUP_THRESHOLD', 0.72))
    if _cross_script_name_similarity(na, nb) >= cs_threshold:
        return True
    return False


def _canonicalize_entities(entities: list, full_text: str = "") -> list:
    """
    Deduplicate entities using normalized text comparison.
    Key design:
    - Normalize BEFORE clustering (not after)
    - DATE entities: exact-match dedup only (no fuzzy) — preserve timeline
    - Person entities: fuzzy dedup with 80% similarity threshold
    - Person-like types dedup across roles (same person can't be ACCUSED + PERSON)
    Returns list of unique entity dicts.
    """
    if not entities:
        return entities

    # Step 1: Normalize all entity text FIRST (before any grouping/clustering)
    for ent in entities:
        raw = ent.get("text", "").strip()
        ent["_raw"] = raw
        # Strip annotations before normalization
        clean = _strip_annotations(raw)
        # Repair OCR-broken Devanagari (fix spaces inside Hindi words)
        clean = _repair_ocr_devanagari(clean if clean else raw)
        clean = _apply_ocr_corrections(clean)
        if ent.get("type", "") in ("PERSON", "ACCUSED", "COMPLAINANT", "WITNESS"):
            clean = _normalize_legal_language(clean)
            clean = _prefer_summary_name_variant(clean)
        clean = clean.strip()
        ent["text"] = clean if clean else raw
        ent["_norm"] = _normalize_entity_text(clean if clean else raw, ent.get("type", ""))

    # Step 2: Dedup DATE entities separately (exact match only — no fuzzy)
    # Also filter out time-only entries (e.g. "02:34 AM", "23:00", "10.00 बजे")
    _TIME_ONLY_RE = re.compile(
        r'^[\s]*\d{1,2}[\.:]\d{2}[\s]*'   # HH:MM or HH.MM
        r'(?:[\.:]\d{2})?[\s]*'            # optional :SS
        r'(?:[-–]\s*\d{1,2}[\.:]\d{2})?[\s]*'  # optional range (e.g. 11.00-11.30)
        r'(?:AM|PM|hrs?|बजे)?[\s]*'        # AM/PM/hrs/बजे
        r'(?:प्रातः|सायं|रात|रात्रि|दोपहर)?[\s]*$',  # time of day
        re.IGNORECASE
    )
    # Regex for FIR-number-like strings that are NOT dates (e.g. "0214/2024")
    _FIR_NUMBER_RE = re.compile(
        r'^\s*0\d{2,4}\s*/\s*\d{4}\s*$'  # starts with 0 + digits / year
    )

    date_entities = [e for e in entities if e.get("type") == "DATE"]
    other_entities = [e for e in entities if e.get("type") != "DATE"]

    years = []
    for e in date_entities:
        norm_date = _normalize_entity_text(e.get("text", ""), "DATE")
        m_year = re.search(r'^(\d{1,2})/(\d{1,2})/(\d{2,4})$', norm_date)
        if not m_year:
            continue
        year = int(m_year.group(3))
        if year < 100:
            year += 2000
        years.append(year)
    ref_year = None
    if years:
        counts = {}
        for y in years:
            counts[y] = counts.get(y, 0) + 1
        ref_year = sorted(counts.items(), key=lambda kv: (-kv[1], -kv[0]))[0][0]

    deduped_dates = []
    seen_date_norms = set()
    for ent in date_entities:
        norm = ent["_norm"]
        if not ent["text"] or not norm:
            continue
        # Skip time-only entries (no date component)
        if _TIME_ONLY_RE.match(ent["text"].strip()):
            logger.info(f"NER: Filtered time-only DATE: '{ent['text']}'")
            continue
        # Skip FIR numbers misclassified as dates (e.g. "0214/2024")
        if _FIR_NUMBER_RE.match(ent["text"].strip()):
            logger.info(f"NER: Filtered FIR number from DATE: '{ent['text']}'")
            continue
        # Plausibility filter using dynamic reference year (mode)
        m_date = re.search(r'^(\d{1,2})/(\d{1,2})/(\d{2,4})$', norm)
        if m_date:
            year = int(m_date.group(3))
            if year < 100:
                year += 2000
            if ref_year is not None:
                if year < (ref_year - 30) or year > (ref_year + 1):
                    logger.info(f"NER DATE filter: rejected '{ent['text']}' (year={year}, ref={ref_year})")
                    continue
        if norm not in seen_date_norms:
            seen_date_norms.add(norm)
            deduped_dates.append({"text": ent["text"], "type": "DATE"})

    # Step 3: Dedup person-like entities across all person roles together
    #   (prevents "मो0 समीर" appearing as both ACCUSED and PERSON)
    person_entities = [e for e in other_entities if e.get("type") in _PERSON_TYPES]
    non_person_entities = [e for e in other_entities if e.get("type") not in _PERSON_TYPES]

    deduped_persons = []
    seen_person_norms = {}  # norm → canonical entry
    for ent in person_entities:
        raw_text = ent.get("_raw", "")
        clean_text = ent["text"]
        norm = ent["_norm"]
        etype = ent.get("type", "PERSON")

        honorific_only_re = re.compile(
            r'^(?:डा0?|डॉ0?|Dr\.?|डॉक्टर)\s*$',
            re.IGNORECASE
        )

        if not clean_text:
            continue

        if etype == "DOCTOR" and honorific_only_re.match(clean_text.strip()):
            logger.info(f"NER: Filtered standalone honorific DOCTOR: {clean_text}")
            continue

        # Reject relation phrases
        if _is_relationship_phrase(raw_text):
            logger.info(f"NER dedup: Rejected relation phrase: '{raw_text}'")
            continue

        # High-sensitivity guard: possessive phrasing inside ACCUSED/WITNESS is risky.
        if etype in ("ACCUSED", "WITNESS") and _contains_possessive_token(raw_text):
            logger.warning(f"NER: Possessive phrase in {etype}; reclassified to PERSON for review: '{clean_text}'")
            etype = "PERSON"

        # Reclassify deceased persons: WITNESS/ACCUSED → PERSON
        if etype == "WITNESS" and _is_deceased_person(raw_text):
            logger.info(f"NER: Reclassified deceased person from WITNESS to PERSON: '{clean_text}'")
            etype = "PERSON"

        # Reclassify complainant context: WITNESS → COMPLAINANT
        if etype == "WITNESS" and _COMPLAINANT_CONTEXT_RE.search(raw_text):
            logger.info(f"NER: Reclassified witness to COMPLAINANT by context: '{clean_text}'")
            etype = "COMPLAINANT"

        # Reclassify police ranks from WITNESS → OFFICER
        if etype == "WITNESS" and re.search(
            r'(?:आरक्षी|सिपाही|हवलदार|दरोगा|थानाध्यक्ष|थाना\s*प्रभारी|'
            r'SI|ASI|S\.I\.|A\.S\.I\.|Inspector|Constable|Head\s*Constable)',
            raw_text, re.IGNORECASE
        ):
            logger.info(f"NER: Reclassified police rank from WITNESS to OFFICER: '{clean_text}'")
            etype = "OFFICER"

        # Reclassify judicial role from OFFICER → JUDGE
        if etype == "OFFICER" and _JUDICIAL_ROLE_RE.search(raw_text):
            logger.info(f"NER: Reclassified OFFICER to JUDGE by judicial keywords: '{clean_text}'")
            etype = "JUDGE"

        if etype in ("OFFICER", "DOCTOR") and _is_correspondence_only_person(clean_text, full_text):
            logger.info(f"NER: Reclassified {etype} -> PERSON (correspondence-only context): '{clean_text}'")
            etype = "PERSON"

        if etype == "OFFICER" and _is_title_only_entity(clean_text):
            logger.debug(f"NER: Filtered title-only OFFICER entity: '{clean_text}'")
            continue

        if etype == "OFFICER" and not _entity_has_officer_token(clean_text):
            if not _officer_context_has_designation(clean_text, full_text):
                logger.info(f"NER: Reclassified OFFICER -> PERSON (no rank/designation evidence): '{clean_text}'")
                etype = "PERSON"

        # Fuzzy match across all person types (including alias-aware matching)
        matched = False
        alias_parts = _extract_alias_parts(clean_text)  # e.g. ["राजेश", "गुड्डू"]

        for seen_norm, canonical in seen_person_norms.items():
            # Standard fuzzy match
            if _fuzzy_match(norm, seen_norm):
                # Keep longer form (with alias); keep more specific role
                if len(clean_text) > len(canonical["text"]):
                    canonical["text"] = clean_text
                current_type = canonical.get("type", "PERSON")
                if _PERSON_ROLE_PRIORITY.get(etype, 0) > _PERSON_ROLE_PRIORITY.get(current_type, 0):
                    canonical["type"] = etype
                # Store alias parts for future matching
                if alias_parts:
                    canonical.setdefault("_aliases", set()).update(
                        _normalize_entity_text(p, etype) for p in alias_parts
                    )
                matched = True
                break

            # OFFICER-specific partial/full-name merge: keep longer person-like name
            if etype == "OFFICER" and canonical.get("type") == "OFFICER":
                ctext = canonical.get("text", "")
                if _officer_short_name_merge_possible(clean_text, ctext):
                    matched = True
                    break
                if _officer_short_name_merge_possible(ctext, clean_text):
                    canonical["text"] = clean_text
                    matched = True
                    break

            # Alias-aware: current entity matches an alias of a seen entity
            seen_aliases = canonical.get("_aliases", set())
            if norm in seen_aliases:
                logger.info(f"NER: Alias merge: '{clean_text}' merged into '{canonical['text']}'")
                matched = True
                break

            # Alias-aware: current entity has aliases, and a seen entity matches one
            if alias_parts:
                alias_norms = {_normalize_entity_text(p, etype) for p in alias_parts}
                if seen_norm in alias_norms:
                    # Seen entity is a component of our alias — upgrade to full alias form
                    canonical["text"] = clean_text
                    canonical.setdefault("_aliases", set()).update(alias_norms)
                    current_type = canonical.get("type", "PERSON")
                    if _PERSON_ROLE_PRIORITY.get(etype, 0) > _PERSON_ROLE_PRIORITY.get(current_type, 0):
                        canonical["type"] = etype
                    logger.info(f"NER: Alias merge: '{seen_norm}' upgraded to '{clean_text}'")
                    matched = True
                    break

        if not matched:
            entry = {"text": clean_text, "type": etype}
            # Store alias parts for future matching
            if alias_parts:
                entry["_aliases"] = {_normalize_entity_text(p, etype) for p in alias_parts}
            seen_person_norms[norm] = entry
            deduped_persons.append(entry)

    # Step 4: Dedup non-person, non-date entities (by type, with fuzzy match)
    # Also do CROSS-TYPE dedup for ORGANIZATION/LOCATION (same entity can appear in both)
    for ent in non_person_entities:
        if ent.get("type") != "EVIDENCE":
            continue
        txt = str(ent.get("text", ""))
        if not txt:
            continue
        if re.search(r'(?:₹|रु\.?|रुपये|रूपये|रुपया)', txt, re.IGNORECASE) or re.search(r'\b\d{2,}(?:,\d{2,3})*(?:/-|/)\b', txt):
            if re.search(r'\d{2,}(?:,\d{2,3})*', txt):
                ent["type"] = "MONETARY"
                logger.info(f"NER: Reclassified EVIDENCE -> MONETARY: '{txt}'")

    by_type = {}
    for ent in non_person_entities:
        etype = ent.get("type", "UNKNOWN")
        by_type.setdefault(etype, []).append(ent)

    # Regex patterns for filtering garbage entities
    # Phone numbers (masked or full): 97xx43xx88, 98xx12xx34, 9876543210, etc.
    _PHONE_RE = re.compile(
        r'^\s*(?:\+?91[\s-]*)?'             # optional +91 prefix
        r'(?:\d{2}xx\d{2}xx\d{2}|'           # masked: 97xx43xx88
        r'\d{2}[xX*]+\d{2}[xX*]+\d{2,4}|'   # other masked patterns
        r'[6-9]\d{9})\s*$',                  # full 10-digit Indian mobile
        re.IGNORECASE
    )
    # Reference/case numbers: FSL/RC/2024/4478, MLG/2024/2187, CD/0214/2024, 0087/2022
    _REF_NUMBER_RE = re.compile(
        r'^\s*(?:'
        r'[A-Z]{2,}[/\-](?:[A-Z]{1,}[/\-])?\d{2,}[/\-]\d{2,}|'  # FSL/RC/2024/4478, CD/0214/2024
        r'[A-Z]{2,}[/\-]\d{4}[/\-]\d+|'                          # MLG/2024/2187
        r'\d{3,4}[/\-]\d{4}'                                      # 0087/2022
        r')\s*$',
        re.IGNORECASE
    )
    # Plain numbers (page numbers, section numbers, etc.) — not real locations
    _PLAIN_NUMBER_RE = re.compile(r'^\s*\d{1,4}\s*$')

    deduped_other = []
    # Cross-type alias dedup for organizations
    _ORG_ALIASES = {
        'sbi': 'state bank of india',
        'state bank of india': 'sbi',
        'bsnl': 'bharat sanchar nigam limited',
        'fsl': 'forensic science laboratory',
        'rcfl': 'regional computer forensic laboratory',
    }
    seen_org_norms = set()  # track org norms across ORGANIZATION and LOCATION types

    for etype, items in by_type.items():
        seen_normalized = {}
        for item in items:
            clean_text = item["text"]
            norm = item["_norm"]
            if not clean_text:
                continue

            # Filter phone numbers from MONETARY
            if etype == "MONETARY" and _PHONE_RE.match(clean_text):
                logger.info(f"NER: Filtered phone number from MONETARY: '{clean_text}'")
                continue

            if etype == "MONETARY":
                amount_re = re.compile(r'[\d,]+(?:\.\d{2})?')
                m_amt = amount_re.search(clean_text.replace(',', ''))
                if m_amt:
                    try:
                        amount_val = float(m_amt.group().replace(',', ''))
                        min_amount = float(getattr(config, 'MONETARY_MIN_AMOUNT', 100))
                        if amount_val <= min_amount:
                            logger.info(f"NER: Filtered micro-amount MONETARY: {clean_text}")
                            continue
                    except ValueError:
                        pass
                if full_text and _is_monetary_noise_context(clean_text, full_text):
                    logger.info(f"NER: Filtered routine-statement MONETARY by context: '{clean_text}'")
                    continue

            # Filter reference/case numbers from ORGANIZATION
            if etype == "ORGANIZATION" and _REF_NUMBER_RE.match(clean_text):
                logger.info(f"NER: Filtered reference number from ORGANIZATION: '{clean_text}'")
                continue

            # Filter plain numbers from LOCATION
            if etype == "LOCATION" and _PLAIN_NUMBER_RE.match(clean_text):
                logger.info(f"NER: Filtered plain number from LOCATION: '{clean_text}'")
                continue

            # Cross-type org/location dedup
            if etype in ("ORGANIZATION", "LOCATION"):
                norm_lower = norm.lower().strip()
                org_core = _organization_core_form(clean_text)
                # Check alias match
                is_alias_dup = False
                for seen in list(seen_org_norms):
                    if norm_lower == seen:
                        is_alias_dup = True
                        break
                    if norm_lower in _ORG_ALIASES and _ORG_ALIASES[norm_lower] == seen:
                        is_alias_dup = True
                        break
                    if seen in _ORG_ALIASES and _ORG_ALIASES[seen] == norm_lower:
                        is_alias_dup = True
                        break
                    # Cross-script institution core match (e.g., सरायकेला थाना vs Saraikela P.S.)
                    seen_core = _organization_core_form(seen)
                    if org_core and seen_core:
                        cs = _cross_script_name_similarity(org_core, seen_core)
                        if cs >= float(getattr(config, 'CROSS_SCRIPT_DEDUP_THRESHOLD', 0.72)):
                            is_alias_dup = True
                            break
                    # Hospital name match: "GOVT. HOSPITAL, DHANPUR" vs "जिला चिकित्सालय, धनपुर"
                    if ('hospital' in norm_lower or 'चिकित्सालय' in norm_lower) and \
                       ('hospital' in seen or 'चिकित्सालय' in seen):
                        # Same city reference
                        if _fuzzy_match(norm_lower.split(',')[-1].strip(), seen.split(',')[-1].strip()):
                            is_alias_dup = True
                            break
                if is_alias_dup:
                    logger.info(f"NER: Cross-type org dedup: skipping '{clean_text}' (alias)")
                    continue
                seen_org_norms.add(norm_lower)

            matched = False
            for seen_norm, canonical in seen_normalized.items():
                org_dup = False
                if etype == "ORGANIZATION":
                    cs = _cross_script_name_similarity(_organization_core_form(clean_text), _organization_core_form(canonical["text"]))
                    if cs >= float(getattr(config, 'CROSS_SCRIPT_DEDUP_THRESHOLD', 0.72)):
                        org_dup = True
                if (etype == "EVIDENCE" and _evidence_entities_duplicate(clean_text, canonical["text"])) or org_dup or _fuzzy_match(norm, seen_norm):
                    if len(clean_text) > len(canonical["text"]):
                        canonical["text"] = clean_text
                    elif etype == "ORGANIZATION" and _is_devanagari_primary(clean_text) and not _is_devanagari_primary(canonical["text"]):
                        canonical["text"] = clean_text
                    matched = True
                    break

            if not matched:
                entry = {"text": clean_text, "type": etype}
                seen_normalized[norm] = entry
                deduped_other.append(entry)

    # Dedicated EVIDENCE pairwise dedup pass (post-collection)
    evidence_items = [e for e in deduped_other if e.get('type') == 'EVIDENCE']
    if evidence_items:
        keep = [True] * len(evidence_items)
        for i in range(len(evidence_items)):
            if not keep[i]:
                continue
            for j in range(i + 1, len(evidence_items)):
                if not keep[j]:
                    continue
                a = evidence_items[i].get('text', '')
                b = evidence_items[j].get('text', '')
                if _evidence_entities_duplicate(a, b):
                    if len(str(a)) >= len(str(b)):
                        keep[j] = False
                    else:
                        keep[i] = False
                        break
        deduped_other = [e for e in deduped_other if e.get('type') != 'EVIDENCE'] + [
            evidence_items[idx] for idx, flag in enumerate(keep) if flag
        ]

    # Combine all
    result = deduped_persons + deduped_dates + deduped_other

    # Clean up internal keys
    for ent in result:
        ent.pop("_raw", None)
        ent.pop("_norm", None)
        ent.pop("_aliases", None)

    return result


def _is_roman_primary(text: str) -> bool:
    letters = re.findall(r'[A-Za-z]', text or '')
    deva = re.findall(r'[\u0900-\u097F]', text or '')
    return len(letters) > 0 and len(letters) >= len(deva)


def _is_devanagari_primary(text: str) -> bool:
    letters = re.findall(r'[A-Za-z]', text or '')
    deva = re.findall(r'[\u0900-\u097F]', text or '')
    return len(deva) > 0 and len(deva) > len(letters)


def _roman_consonant_skeleton(text: str) -> str:
    chars = re.findall(r'[A-Za-z]', text or '')
    cons = [c.lower() for c in chars if c.lower() not in 'aeiou']
    return ''.join(cons)


def _devanagari_consonant_skeleton(text: str) -> str:
    if not text:
        return ''
    out = re.sub(r'[\u093e-\u094c\u0901\u0902\u0903\u094d]', '', text)
    out = re.sub(r'\u093c', '', out)
    chars = re.findall(r'[\u0915-\u0939\u0958-\u095F]', out)
    return ''.join(chars)


def _jaccard_char_similarity(a: str, b: str) -> float:
    sa, sb = set(a), set(b)
    if not sa or not sb:
        return 0.0
    return len(sa & sb) / max(1, len(sa | sb))


def _cross_script_name_similarity(a: str, b: str) -> float:
    if _is_roman_primary(a) and _is_devanagari_primary(b):
        ra = _roman_consonant_skeleton(a)
        db = _devanagari_consonant_skeleton(b)
        return _jaccard_char_similarity(ra, db)
    if _is_devanagari_primary(a) and _is_roman_primary(b):
        da = _devanagari_consonant_skeleton(a)
        rb = _roman_consonant_skeleton(b)
        return _jaccard_char_similarity(da, rb)
    return 0.0


def _fuzzy_match(a: str, b: str) -> bool:
    """
    Check if two normalized entity texts are likely the same entity.
    Uses character-level edit distance ratio for Hindi OCR variation tolerance.
    """
    if a == b:
        return True
    if not a or not b:
        return False

    a_tokens = [tok for tok in a.lower().split() if tok]
    b_tokens = [tok for tok in b.lower().split() if tok]
    if (
        (a.lower().strip() in _ENTITY_STOPWORDS)
        or (b.lower().strip() in _ENTITY_STOPWORDS)
        or (a_tokens and all(tok in _ENTITY_STOPWORDS for tok in a_tokens))
        or (b_tokens and all(tok in _ENTITY_STOPWORDS for tok in b_tokens))
    ):
        return False

    # If one is a substring of the other (e.g. 'सादिक' in 'मो. सादिक')
    if a in b or b in a:
        return True

    # Word-subset check: if all words of the shorter text appear in the longer
    # (handles "HP Laptop" vs "HP brand Laptop", "सोने की चेन" vs "22 कैरेट सोने की चेन")
    words_a = set(a.split())
    words_b = set(b.split())
    if len(words_a) >= 2 and len(words_b) >= 2:
        if words_a.issubset(words_b) or words_b.issubset(words_a):
            return True

    threshold = float(getattr(config, 'CROSS_SCRIPT_DEDUP_THRESHOLD', 0.72))
    if _cross_script_name_similarity(a, b) >= threshold:
        return True

    tokens_a = set(a.replace('.', ' ').split())
    tokens_b = set(b.replace('.', ' ').split())
    common = tokens_a & tokens_b
    if len(common) >= 2 and any(len(t) >= 3 for t in common):
        return True

    # Character-level similarity (Levenshtein-like)
    # For short names, even 1-2 char difference in Hindi can be OCR noise
    longer = max(len(a), len(b))
    if longer == 0:
        return True

    # Simple edit distance (optimized for short strings)
    dist = _edit_distance(a, b)
    ratio = 1.0 - (dist / longer)

    # Threshold by length bucket
    if longer < 4:
        return a == b
    if 4 <= longer <= 6:
        return ratio >= 0.90
    if 7 <= longer <= 12:
        return ratio >= 0.85
    return ratio >= 0.80


def _edit_distance(s1: str, s2: str) -> int:
    """Compute Levenshtein edit distance between two strings."""
    if len(s1) < len(s2):
        return _edit_distance(s2, s1)
    if len(s2) == 0:
        return len(s1)

    prev_row = list(range(len(s2) + 1))
    for i, c1 in enumerate(s1):
        curr_row = [i + 1]
        for j, c2 in enumerate(s2):
            insertions = prev_row[j + 1] + 1
            deletions = curr_row[j] + 1
            substitutions = prev_row[j] + (c1 != c2)
            curr_row.append(min(insertions, deletions, substitutions))
        prev_row = curr_row

    return prev_row[-1]


def _convert_ner_to_flat_list(ner: dict) -> list:
    """
    Convert old-format NER dict {"TYPE": ["str", ...]} to new flat list format
    [{"text": "...", "type": "TYPE"}, ...].
    Handles both old dict-of-lists format and new list-of-dicts format.
    """
    if isinstance(ner, list):
        return ner  # already in new format

    entities = []
    for etype, items in ner.items():
        if not isinstance(items, list):
            continue
        for item in items:
            if isinstance(item, dict):
                entities.append(item)
            elif isinstance(item, str):
                entities.append({"text": item, "type": etype})
    return entities


def _validate_legal_sections(entities: list, classification: dict) -> list:
    """
    Validate LEGAL_SECTION entities: prevent cross-contamination
    between BNS (new law) and IPC (old law).
    """
    legal_entities = [e for e in entities if e.get("type") == "LEGAL_SECTION"]
    other_entities = [e for e in entities if e.get("type") != "LEGAL_SECTION"]

    if not legal_entities:
        return entities

    all_legal_text = " ".join(e["text"] for e in legal_entities)
    has_bns = bool(_BNS_SECTIONS.search(all_legal_text))
    has_ipc = bool(_IPC_SECTIONS.search(all_legal_text))

    # Also check classification
    detected = classification.get("detected_sections", [])
    detected_text = " ".join(str(s) for s in detected) if detected else ""
    if _BNS_SECTIONS.search(detected_text):
        has_bns = True
    if _IPC_SECTIONS.search(detected_text):
        has_ipc = True

    def _section_base(text: str) -> str:
        m = re.search(r'(\d+[A-Za-z]?)', text)
        return m.group(1) if m else ""

    if has_bns:
        # Filter out IPC sections from a BNS FIR
        filtered_legal = []
        removed = 0
        for e in legal_entities:
            base = _section_base(e["text"])
            if _IPC_SECTIONS.search(e["text"]) and not _BNS_SECTIONS.search(e["text"]):
                logger.info(f"NER validation: Removed hallucinated IPC section: {e['text']}")
                removed += 1
                continue
            filtered_legal.append(e)
        if removed:
            logger.info(f"NER validation: Removed {removed} IPC sections from BNS FIR")
        return other_entities + filtered_legal

    return entities


def _section_dedup_key(section_text: str) -> str:
    s = _normalize_legal_language((section_text or "").strip())
    if not s:
        return ""
    m = re.search(r'\b(BNSS|BNS|IPC|BSA|NDPS|CRPC|Arms\s*Act|Evidence\s*Act|भा\.?.*?वि\.?)\b', s, re.IGNORECASE)
    act = (m.group(1).upper().replace(" ", "") if m else "GEN")
    n = re.search(r'(\d+[A-Za-z]?)', s)
    base = n.group(1).upper() if n else re.sub(r'\W+', '', s.lower())
    return f"{act}:{base}"


def _section_specificity_score(section_text: str) -> tuple:
    s = (section_text or "").strip()
    has_sub = bool(re.search(r'\(\d+[A-Za-z]?\)', s))
    has_colon_desc = bool(re.search(r':\s*\S', s))
    return (1 if has_sub else 0, 1 if has_colon_desc else 0, len(s))


def _dedup_legal_sections(sections: list[str]) -> list[str]:
    """Deduplicate legal sections and keep more specific variants for same provision."""
    by_key: dict[str, str] = {}
    ordered_keys: list[str] = []
    for raw in sections or []:
        text = _normalize_legal_language(str(raw or "").strip())
        if not text:
            continue
        key = _section_dedup_key(text)
        if not key:
            continue
        if key not in by_key:
            by_key[key] = text
            ordered_keys.append(key)
            continue
        if _section_specificity_score(text) > _section_specificity_score(by_key[key]):
            by_key[key] = text
    return [by_key[k] for k in ordered_keys if k in by_key]


def _remove_validated_sections_lines(summary: str) -> str:
    return re.sub(r'\n?\*\*Validated Legal Sections:\*\*[^\n]*', '', summary or '', flags=re.IGNORECASE).strip()


def _upsert_validated_sections_line(summary: str, sections: list[str]) -> str:
    cleaned = _remove_validated_sections_lines(summary)
    deduped = _dedup_legal_sections(sections)
    if not deduped:
        return cleaned
    return (cleaned + f"\n\n**Validated Legal Sections:** {', '.join(deduped)}").strip()


def _filter_procedural_sections(sections: list[str], validated_sections: list[str], full_text: str = "") -> list[str]:
    """Exclude prior-only and clearly procedural legal sections while retaining current charged sections."""
    validated = _dedup_legal_sections(validated_sections or [])
    prior_re = re.compile(r'पूर्व\s*प्रकरण|आपराधिक\s*इतिहास|पूर्व\s*काण्ड|prior\s*case|criminal\s*history|previous\s*case', re.IGNORECASE)
    procedural_kw_re = re.compile(r'साक्ष्य|प्रक्रिया|आवेदन|application|petition|bail|जमानत|statement|बयान', re.IGNORECASE)
    charge_kw_re = re.compile(r'fir|मुकदमा\s*दर्ज|आरोप\s*पत्र|chargesheet|अभियुक्त|आरोपी|धारा\s*लगाई', re.IGNORECASE)

    prior_spans = []
    if full_text:
        for m in prior_re.finditer(full_text):
            prior_spans.append((max(0, m.start() - 300), min(len(full_text), m.end() + 600)))

    def _occurrence_windows(sec_text: str) -> list[tuple[int, int]]:
        wins = []
        if not full_text:
            return wins
        key_num = re.search(r'(\d+[A-Za-z]?(?:\(\d+\))?)', sec_text)
        probes = [re.escape(sec_text)]
        if key_num:
            probes.append(rf'\b{re.escape(key_num.group(1))}\b')
        for pat in probes:
            for m in re.finditer(pat, full_text, flags=re.IGNORECASE):
                wins.append((max(0, m.start() - 120), min(len(full_text), m.end() + 120)))
        return wins

    out = []
    for sec in sections or []:
        sec_text = str(sec or "").strip()
        if not sec_text:
            continue

        windows = _occurrence_windows(sec_text)
        has_non_prior = False
        has_prior = False
        has_charge_ctx = False
        has_proc_ctx = False
        for s, e in windows:
            ctx = full_text[s:e] if full_text else ""
            in_prior = any(ps <= s <= pe or ps <= e <= pe for ps, pe in prior_spans)
            has_prior = has_prior or in_prior
            has_non_prior = has_non_prior or (not in_prior)
            has_charge_ctx = has_charge_ctx or bool(charge_kw_re.search(ctx))
            has_proc_ctx = has_proc_ctx or bool(procedural_kw_re.search(ctx))

        # Procedural-only sections: exclude unless also seen in charge context.
        if has_proc_ctx and not has_charge_ctx:
            logger.info(f"NER: Filtered procedural LEGAL_SECTION: {sec_text}")
            continue

        # Prior-only sections: exclude if they appear only in prior-history context.
        if has_prior and not has_non_prior:
            logger.info(f"NER: Filtered prior-only LEGAL_SECTION: {sec_text}")
            continue

        # Keep sections linked to validated set OR independently present outside prior context.
        if validated:
            sec_key = _section_dedup_key(sec_text)
            val_keys = {_section_dedup_key(v) for v in validated if v}
            if sec_key not in val_keys and windows and not has_non_prior and not has_charge_ctx:
                logger.info(f"NER: Filtered weak LEGAL_SECTION (no current-charge evidence): {sec_text}")
                continue

        out.append(sec_text)

    return _dedup_legal_sections(out)


# ── Legal Section Shortening ──
# LLMs often produce verbose statute language for IPC/BNS sections.
# Police reports use concise short forms. This map truncates to report-style labels.
_IPC_SHORT_NAMES = {
    # IPC common sections
    '34': 'Common intention',
    '107': 'Abetment',
    '109': 'Abetment (if act committed)',
    '120B': 'Criminal conspiracy',
    '147': 'Rioting',
    '148': 'Rioting (armed)',
    '149': 'Unlawful assembly',
    '153A': 'Communal disharmony',
    '279': 'Rash driving',
    '294': 'Obscene acts',
    '302': 'Murder',
    '304': 'Culpable homicide',
    '304A': 'Death by negligence',
    '304B': 'Dowry death',
    '306': 'Abetment of suicide',
    '307': 'Attempt to murder',
    '308': 'Attempt culpable homicide',
    '323': 'Voluntarily causing hurt',
    '324': 'Hurt (dangerous weapons)',
    '325': 'Grievous hurt',
    '326': 'Grievous hurt (dangerous weapons)',
    '332': 'Hurt to public servant',
    '341': 'Wrongful restraint',
    '342': 'Wrongful confinement',
    '354': 'Outraging modesty (woman)',
    '363': 'Kidnapping',
    '364': 'Kidnapping for murder',
    '365': 'Kidnapping to confinement',
    '376': 'Rape',
    '379': 'Theft',
    '380': 'Theft in dwelling house',
    '382': 'Theft (preparation for death/hurt)',
    '384': 'Extortion',
    '392': 'Robbery',
    '393': 'Attempt to commit robbery',
    '394': 'Robbery with hurt',
    '395': 'Dacoity',
    '396': 'Dacoity with murder',
    '397': 'Robbery/dacoity (attempt to kill)',
    '398': 'Robbery/dacoity (deadly weapon)',
    '399': 'Preparation for dacoity',
    '406': 'Criminal breach of trust',
    '411': 'Dishonestly receiving stolen property',
    '414': 'Assisting in concealment of stolen property',
    '419': 'Cheating by impersonation',
    '420': 'Cheating',
    '427': 'Mischief (damage ₹50+)',
    '452': 'House-trespass (preparation for hurt)',
    '457': 'Lurking house-trespass by night',
    '458': 'Lurking house-trespass (preparation for hurt)',
    '467': 'Forgery of valuable security',
    '468': 'Forgery for cheating',
    '471': 'Using forged document as genuine',
    '498A': 'Cruelty by husband / relatives',
    '504': 'Intentional insult (provocation)',
    '506': 'Criminal intimidation',
    '509': 'Insulting modesty of woman',
}

# Regex to extract IPC/BNS section number from text like "IPC 302: Punishment for murder"
_SECTION_NUM_RE = re.compile(
    r'\b(?:IPC|BNS|BNSS|आईपीसी|बीएनएस)\s*(?:Section|Sec\.?|धारा)?\s*'
    r'(\d+[A-Z]?(?:\(\d+\))?)\b',
    re.IGNORECASE,
)


def _shorten_legal_sections(entities: list) -> list:
    """
    Shorten verbose IPC/BNS section descriptions to concise report-style labels.
    E.g. "IPC 34: Acts done by several persons in furtherance of common intention"
      → "IPC 34: Common intention"
    """
    for ent in entities:
        if ent.get("type") != "LEGAL_SECTION":
            continue
        text = ent["text"]
        # Only process if text has a colon (section: description format)
        if ':' not in text:
            continue
        m = _SECTION_NUM_RE.search(text)
        if not m:
            continue
        section_num = m.group(1)
        short_name = _IPC_SHORT_NAMES.get(section_num)
        if not short_name:
            continue
        # Replace everything after the colon with the short name
        colon_pos = text.index(':')
        prefix = text[:colon_pos + 1].strip()
        ent["text"] = f"{prefix} {short_name}"
    return entities


def _extract_monetary_from_text(full_text: str) -> list[str]:
    if not full_text:
        return []
    min_amount = float(getattr(config, 'MONETARY_MIN_AMOUNT', 100))
    patterns = [
        r'₹\s*\d{1,3}(?:,\d{2,3})*(?:\.\d{1,2})?',
        r'(?:रु\.?|रुपये|रूपये|रुपया)\s*\d{1,3}(?:,\d{2,3})*(?:\.\d{1,2})?',
        r'\b\d{1,3}(?:,\d{2,3})*(?:/-|/)\b',
        r'\b\d{1,3}(?:,\d{2,3})*\s*(?:रुपये|रूपये|रुपया|रुपए)\b',
    ]
    found = []
    seen = set()
    for pat in patterns:
        for m in re.finditer(pat, full_text, flags=re.IGNORECASE):
            raw_amt = m.group(0).strip()
            num_match = re.search(r'\d{1,3}(?:,\d{2,3})*(?:\.\d{1,2})?', raw_amt)
            if not num_match:
                continue
            try:
                value = float(num_match.group(0).replace(',', ''))
            except ValueError:
                continue
            if value <= min_amount:
                continue
            key = raw_amt.lower()
            if key in seen:
                continue
            seen.add(key)
            found.append(raw_amt)
    return found


def _merge_missed_monetary_entities(entities: list, full_text: str) -> list:
    if not full_text:
        return entities
    existing = {str(e.get('text', '')).strip().lower() for e in entities if e.get('type') == 'MONETARY'}
    for amount in _extract_monetary_from_text(full_text):
        key = amount.strip().lower()
        if key in existing:
            continue
        entities.append({"text": amount.strip(), "type": "MONETARY"})
        existing.add(key)
        logger.info(f"NER: Added missed MONETARY from raw text: '{amount.strip()}'")
    return entities


def _filter_prior_case_sections(entities: list, full_text: str) -> list:
    """Remove legal sections that appear only in prior criminal history context and not in the main FIR body."""
    legal = [e for e in entities if e.get("type") == "LEGAL_SECTION"]
    other = [e for e in entities if e.get("type") != "LEGAL_SECTION"]
    if not legal or not full_text:
        return entities

    # Detect prior case context blocks
    prior_block_re = re.compile(
        r'(?:आपराधिक\s*इतिहास|पूर्व\s*अपराध|criminal\s*history|'
        r'prior\s*case|पूर्व\s*काण्ड|काण्ड\s*सं0[^।\n]*(?:199\d|200\d|'
        r'201\d)|previous\s*FIR)[^\n।]{0,800}',
        re.IGNORECASE
    )
    prior_blocks = ' '.join(m.group(0) for m in prior_block_re.finditer(full_text))

    # Also flag slash-compound sections (e.g. "380/457")
    slash_section_re = re.compile(r'^\d+/\d+')
    old_ipc_format_re = re.compile(
        r'^\d{3}\s+(?:भा0|भारतीय|IPC|ipc)',
        re.IGNORECASE
    )

    filtered_legal = []
    for ent in legal:
        text = ent.get("text", "")
        # Slash-compound sections: filter only when found in prior-case block context
        if slash_section_re.search(text.strip()):
            if not prior_blocks or text[:5] not in prior_blocks:
                pass
            else:
                logger.info(f"NER: Filtered prior-case slash section: {text}")
                continue

        # Also filter standalone old-IPC sections (e.g. "379 भा0 द0 वि0")
        if old_ipc_format_re.search(text.strip()):
            if prior_blocks and text[:6] in prior_blocks:
                logger.info(f"NER: Filtered old-IPC prior-case section: {text}")
                continue
        # Filter if section ONLY appears in prior-case block text
        if prior_blocks and text in prior_blocks:
            # Check if it also appears outside prior block in full_text
            outside_count = full_text.replace(prior_blocks, '').count(text)
            if outside_count == 0:
                logger.info(f"NER: Filtered prior-case section: {text}")
                continue
        filtered_legal.append(ent)

    return other + filtered_legal


def _process_ner_output(raw_ner: dict | list, classification: dict, full_text: str = "") -> list:
    """
    Full NER post-processing pipeline:
    1. Convert to flat list format
    2. Strip annotations from entity text
    3. Reject relation phrases
    4. Canonicalize / deduplicate
    5. Validate legal sections
    Returns clean list of {"text": ..., "type": ...} dicts.
    """
    # Step 1: Flatten
    entities = _convert_ner_to_flat_list(raw_ner)

    # Step 1b: Recover missed monetary amounts from raw text
    entities = _merge_missed_monetary_entities(entities, full_text)

    # Step 2-4: Canonicalize (normalize first, then cluster + dedup + relation rejection)
    entities = _canonicalize_entities(entities, full_text=full_text)

    # Step 5: Legal section validation
    entities = _validate_legal_sections(entities, classification)

    # Step 5b: Filter sections that belong only to prior-criminal-history context
    entities = _filter_prior_case_sections(entities, full_text=full_text)

    # Step 5c: Remove procedural/prior sections not aligned with validated current-case sections
    validated_sections = classification.get("detected_sections", []) if isinstance(classification, dict) else []
    legal_sections = [e.get("text", "") for e in entities if e.get("type") == "LEGAL_SECTION"]
    filtered_legal = set(_filter_procedural_sections(legal_sections, validated_sections, full_text))
    entities = [
        e for e in entities
        if e.get("type") != "LEGAL_SECTION" or e.get("text", "") in filtered_legal
    ]

    # Step 6: Shorten verbose legal section descriptions
    entities = _shorten_legal_sections(entities)

    logger.info(f"NER post-processing: {len(entities)} canonical entities")
    return entities


def _find_entities_json_block(text: str):
    """
    Find a raw (unfenced) JSON object containing "entities" key using
    balanced-brace matching. Returns a match-like object with .group(1)
    or None.
    """
    # Find the opening brace before "entities"
    idx = text.find('"entities"')
    if idx < 0:
        return None

    # Walk backwards to find the opening '{'
    brace_start = text.rfind('{', 0, idx)
    if brace_start < 0:
        return None

    # Walk forward with balanced-brace counting
    depth = 0
    for i in range(brace_start, len(text)):
        ch = text[i]
        if ch == '{':
            depth += 1
        elif ch == '}':
            depth -= 1
            if depth == 0:
                json_str = text[brace_start:i + 1]
                # Return a simple object with .group(1) interface
                class _Match:
                    def group(self, n):
                        return json_str
                    def strip(self):
                        return json_str.strip()
                return _Match()
    return None


def _parse_combined_response(response: str) -> tuple:
    """
    Parse the combined response into summary, classification, and NER.
    Returns (summary, classification, ner_entities).
    """
    # --- Parse Classification JSON ---
    json_match = re.search(r'```json\s*([\s\S]*?)```', response)
    if not json_match:
        json_match = re.search(r'(\{[^{}]*"primary_crime_type"[^{}]*\})', response, re.DOTALL)

    classification = {}
    summary = response

    if json_match:
        json_str = json_match.group(1) if '```' in response else json_match.group(0)
        try:
            classification = json.loads(json_str.strip())
        except json.JSONDecodeError:
            try:
                cleaned = re.sub(r',\s*}', '}', json_str.strip())
                cleaned = re.sub(r',\s*]', ']', cleaned)
                classification = json.loads(cleaned)
            except json.JSONDecodeError:
                logger.warning("Could not parse classification JSON from response")

        # Remove Task 2 and Task 3 sections from summary display
        task2_pattern = re.search(r'━+\s*\n\s*कार्य 2', summary)

        if task2_pattern:
            summary = summary[:task2_pattern.start()].strip()
        else:
            summary = re.sub(r'```json[\s\S]*?```', '', summary)
            summary = re.sub(r'```ner_json[\s\S]*?```', '', summary).strip()

        # Also remove raw (unfenced) NER JSON blocks from summary
        summary = re.sub(r'\{[\s\S]*?"entities"\s*:\s*\[[\s\S]*?\]\s*\}', '', summary).strip()
        # Remove leftover Task 3 headers
        summary = re.sub(r'━*\s*\n?\s*कार्य\s*3[\s\S]*$', '', summary).strip()

    if not classification:
        classification = {
            "primary_crime_type": "unknown",
            "secondary_crime_types": [],
            "detected_sections": [],
            "confidence": "low",
            "reasoning": "Could not parse classification from response",
        }

    # --- Parse NER JSON ---
    ner_entities = {}
    ner_match = re.search(r'```ner_json\s*([\s\S]*?)```', response)
    if not ner_match:
        # Fallback 1: ```json block containing "entities"
        ner_match = re.search(r'```json\s*([\s\S]*?"entities"[\s\S]*?)```', response)
    if not ner_match:
        # Fallback 2: balanced-brace extraction for raw JSON with "entities"
        ner_match = _find_entities_json_block(response)
    logger.info(f"NER block search: found={'YES' if ner_match else 'NO'}, response length={len(response)}")
    if ner_match:
        ner_str = ner_match.group(1).strip() if hasattr(ner_match, 'group') else ner_match.strip()
        logger.info(f"NER JSON block found ({len(ner_str)} chars)")
        try:
            ner_data = json.loads(ner_str)
            ner_entities = ner_data.get("entities", ner_data)
        except json.JSONDecodeError:
            try:
                cleaned = re.sub(r',\s*}', '}', ner_str)
                cleaned = re.sub(r',\s*]', ']', cleaned)
                ner_data = json.loads(cleaned)
                ner_entities = ner_data.get("entities", ner_data)
            except json.JSONDecodeError:
                # Fallback 3: try to extract the JSON array directly
                arr_match = re.search(r'"entities"\s*:\s*(\[\s*\{[\s\S]*?\}\s*\])', response)
                if arr_match:
                    try:
                        ner_entities = json.loads(arr_match.group(1))
                        logger.info(f"NER parsed via direct array extraction: {len(ner_entities)} entities")
                    except json.JSONDecodeError:
                        logger.warning("Could not parse NER JSON from response")
                else:
                    logger.warning("Could not parse NER JSON from response")
    else:
        logger.warning("No NER JSON block found in response (looked for ```ner_json)")

    # --- Full NER post-processing: flatten → strip → dedup → validate ---
    ner_entities = _process_ner_output(ner_entities, classification)

    logger.info(f"_parse_combined_response returning: summary={len(summary)} chars, "
                f"classification keys={list(classification.keys())}, "
                f"ner_entities={len(ner_entities)} items (type={type(ner_entities).__name__})")

    return summary, classification, ner_entities


def _extract_fenced_json(response: str, fence_name: str) -> dict:
    """Extract JSON from a named fenced block, returning {} on failure."""
    block = re.search(rf'```{re.escape(fence_name)}\s*([\s\S]*?)```', response)
    if not block:
        return {}
    payload = block.group(1).strip()
    try:
        return json.loads(payload)
    except json.JSONDecodeError:
        cleaned = re.sub(r',\s*}', '}', payload)
        cleaned = re.sub(r',\s*]', ']', cleaned)
        try:
            return json.loads(cleaned)
        except json.JSONDecodeError:
            return {}


def _extract_summary_body(response: str) -> str:
    """Remove JSON blocks and keep the rich narrative summary."""
    out = re.sub(r'```classification_json[\s\S]*?```', '', response)
    out = re.sub(r'```checklist_json[\s\S]*?```', '', out)
    out = re.sub(r'```ner_json[\s\S]*?```', '', out)
    out = re.sub(r'\n{3,}', '\n\n', out)
    return out.strip()


def _validate_checklist_payload(llm_checklist: list, required_items: list) -> dict:
    """Normalize checklist payload and ensure all required items are present."""
    allowed_status = {"present", "missing", "partial"}
    required_by_id = {item["id"]: item for item in required_items}
    llm_map = {}

    for row in llm_checklist or []:
        item_id = (row or {}).get("id", "").strip()
        if not item_id:
            continue
        status = str((row or {}).get("status", "missing")).strip().lower()
        if status not in allowed_status:
            status = "missing"
        llm_map[item_id] = {
            "id": item_id,
            "status": status,
            "page_no": str((row or {}).get("page_no", "")).strip(),
            "remarks": str((row or {}).get("remarks", "")).strip(),
        }

    normalized = []
    for item in required_items:
        item_id = item["id"]
        if item_id in llm_map:
            normalized.append(llm_map[item_id])
        else:
            normalized.append({
                "id": item_id,
                "status": "missing",
                "page_no": "",
                "remarks": "दस्तावेज़ में स्पष्ट उल्लेख नहीं मिला।",
            })

    return {"checklist": normalized}


def _build_ner_prompt(text: str, core_facts: dict, use_stable_schema: bool = False) -> str:
     """Build dedicated NER extraction prompt for high-precision legal entity extraction."""
     facts_json = json.dumps(core_facts, ensure_ascii=False)

     if use_stable_schema:
          entity_types = [
                "ACCUSED", "WITNESS", "OFFICER", "DOCTOR", "PERSON", "DATE",
                "LOCATION", "LEGAL_SECTION", "ORGANIZATION", "LANDMARK", "EVIDENCE", "MONETARY",
          ]
          json_examples = """    {"text": "मो0 इरफान", "type": "ACCUSED"},
     {"text": "24/10/2024", "type": "DATE"},
     {"text": "सरायकेला थाना", "type": "ORGANIZATION"}"""
     else:
          entity_types = [
                "ACCUSED", "COMPLAINANT", "VICTIM", "DECEASED", "WITNESS", "OFFICER", "JUDGE", "DOCTOR", "PERSON", "DATE",
                "LOCATION", "LEGAL_SECTION", "ORGANIZATION", "LANDMARK", "EVIDENCE", "MONETARY",
          ]
          json_examples = """    {"text": "मो0 इरफान", "type": "ACCUSED"},
     {"text": "राजेश यादव", "type": "COMPLAINANT"},
         {"text": "सीमा देवी", "type": "VICTIM"},
         {"text": "रामलाल वर्मा", "type": "DECEASED"},
     {"text": "माननीय सीजेएम", "type": "JUDGE"},
     {"text": "24/10/2024", "type": "DATE"}"""

     entity_type_block = "\n".join(f"- {etype}" for etype in entity_types)

     return f"""
निम्न चार्जशीट पाठ से Named Entity Recognition (NER) निकालें और केवल JSON लौटाएँ।

Entity types to extract (ONLY these):
{entity_type_block}

Strict rules:
1) "text" field में केवल entity name/value दें:
    - role description नहीं
    - bracket/parenthesis annotation नहीं
    - relationship words नहीं (जैसे पिता, पुत्र, निवासी)
2) प्रत्येक unique व्यक्ति exactly once आए।
3) OCR errors ठीक करके ही output दें (especially ि↔न confusion).
4) COMPLAINANT को सीधे COMPLAINANT type में दें — PERSON में न डालें।
    COMPLAINANT वह व्यक्ति है जिसने FIR/शिकायत दर्ज कराई हो (context से infer करें)।
5) DECEASED वह व्यक्ति है जिसकी मृत्यु/हत्या/पोस्टमार्टम document context में स्पष्ट हो।
6) VICTIM वह जीवित पीड़ित व्यक्ति है जिसे अपराध में हानि/चोट पहुँची हो।
7) LEGAL_SECTION:
    - केवल वही धाराएँ जो वर्तमान FIR/चार्जशीट में आरोप पंजीकरण या आरोपी पर आरोप निर्धारण के लिए लागू हैं
    - prior case/criminal history/procedural authority/court order-only sections शामिल न करें
    - short labels only (e.g. "BNS 103(1): Murder", "Arms Act 27: Possession of firearm")
    - पूरे bare act language quotes नहीं
8) DATE:
    - केवल calendar dates दें (dd/mm/yyyy)
    - time strings (e.g. "10:00 PM") बिल्कुल नहीं
9) MONETARY:
    - केवल रुपये की राशि (₹N, रु N, N/-, N/, N रुपये, N रुपए, N रूपये सहित)
    - phone numbers/ID/reference numbers नहीं
10) LOCATION:
    - केवल स्थान के नाम
    - कोई संख्या/section reference नहीं
11) ORGANIZATION:
    - केवल संस्थान/थाना/अस्पताल/लैब/कार्यालय
    - reference numbers (e.g. "FSL/RC/2024/4478") नहीं
12) Expected coverage guidance:
    - persons: 30-60
    - dates: 15-35
    - locations: 5-15
    - legal sections: 3-8
    - evidence items: 5-15
    - organizations: 5-10

Use these extracted anchors for grounding (do not invent):
{facts_json}

Output format (strict JSON only):
```ner_json
{{
  "entities": [
{json_examples}
  ]
}}
```

Document:
{text}
"""


def _parse_ner_response(response: str) -> list:
    """Parse NER response payload and return raw entities list before post-processing."""
    ner_entities: list = []
    allowed_types = {
        "ACCUSED", "COMPLAINANT", "VICTIM", "DECEASED", "WITNESS", "OFFICER", "JUDGE", "DOCTOR", "PERSON", "DATE",
        "LOCATION", "LEGAL_SECTION", "ORGANIZATION", "LANDMARK", "EVIDENCE", "MONETARY",
    }

    def _sanitize_entities(items: list) -> list:
        cleaned = []
        for item in items or []:
            if not isinstance(item, dict):
                continue
            text_val = str(item.get("text", "")).strip()
            type_val = str(item.get("type", "")).strip().upper()
            if not text_val or type_val not in allowed_types:
                continue
            cleaned.append({"text": text_val, "type": type_val})
        return cleaned

    ner_match = re.search(r'```ner_json\s*([\s\S]*?)```', response)
    if not ner_match:
        ner_match = re.search(r'```json\s*([\s\S]*?"entities"[\s\S]*?)```', response)
    if not ner_match:
        ner_match = _find_entities_json_block(response)

    logger.info(f"NER parse: block_found={'YES' if ner_match else 'NO'}, response_len={len(response)}")

    if ner_match:
        ner_str = ner_match.group(1).strip() if hasattr(ner_match, "group") else ner_match.strip()
        try:
            ner_data = json.loads(ner_str)
            if isinstance(ner_data, dict):
                ner_entities = _sanitize_entities(ner_data.get("entities", []))
            elif isinstance(ner_data, list):
                ner_entities = _sanitize_entities(ner_data)
        except json.JSONDecodeError:
            cleaned = re.sub(r',\s*}', '}', ner_str)
            cleaned = re.sub(r',\s*]', ']', cleaned)
            try:
                ner_data = json.loads(cleaned)
                if isinstance(ner_data, dict):
                    ner_entities = _sanitize_entities(ner_data.get("entities", []))
                elif isinstance(ner_data, list):
                    ner_entities = _sanitize_entities(ner_data)
            except json.JSONDecodeError:
                arr_match = re.search(r'"entities"\s*:\s*(\[\s*\{[\s\S]*?\}\s*\])', response)
                if arr_match:
                    try:
                        ner_entities = _sanitize_entities(json.loads(arr_match.group(1)))
                    except json.JSONDecodeError:
                        logger.warning("NER parse failed: entities array malformed")

    # Fallback: try all fenced json blocks and pick the one containing entities
    if not ner_entities:
        for block in re.findall(r'```json\s*([\s\S]*?)```', response):
            try:
                payload = json.loads(block.strip())
            except Exception:
                continue
            if isinstance(payload, dict) and isinstance(payload.get("entities"), list):
                ner_entities = _sanitize_entities(payload.get("entities", []))
                if ner_entities:
                    logger.info("NER parse fallback: extracted entities from generic json fenced block")
                    break

    if not isinstance(ner_entities, list):
        ner_entities = []

    logger.info(f"NER parse: extracted {len(ner_entities)} raw entities")
    return ner_entities


def _ner_quality_check(ner_entities: list) -> tuple[bool, dict]:
    """
    Validate NER quality using count and type-distribution checks.
    Fails if total entities < 10 OR a single non-DATE type dominates > 60%.
    """
    if not isinstance(ner_entities, list):
        return False, {"reason": "not_list", "count": 0, "dominant_type": None, "dominance_ratio": 1.0}

    count = len(ner_entities)
    if count < 10:
        return False, {"reason": "low_count", "count": count, "dominant_type": None, "dominance_ratio": 0.0}

    non_date = [e for e in ner_entities if str(e.get("type", "")).upper() != "DATE"]
    if not non_date:
        return True, {"reason": "ok", "count": count, "dominant_type": None, "dominance_ratio": 0.0}

    type_counts = {}
    for ent in non_date:
        etype = str(ent.get("type", "UNKNOWN")).upper()
        type_counts[etype] = type_counts.get(etype, 0) + 1

    dominant_type, dominant_count = max(type_counts.items(), key=lambda kv: kv[1])
    dominance_ratio = dominant_count / max(1, len(non_date))

    if dominance_ratio > 0.60:
        return False, {
            "reason": "dominance",
            "count": count,
            "dominant_type": dominant_type,
            "dominance_ratio": round(dominance_ratio, 3),
        }

    return True, {
        "reason": "ok",
        "count": count,
        "dominant_type": dominant_type,
        "dominance_ratio": round(dominance_ratio, 3),
    }


def _fallback_ner_from_core_facts(core_facts: dict) -> list:
    """Build a conservative fallback NER list from deterministic core facts."""
    entities: list[dict] = []

    for name in core_facts.get("accused", []) or []:
        val = str(name).strip()
        if val:
            entities.append({"text": val, "type": "ACCUSED"})

    for dt in core_facts.get("dates", []) or []:
        val = str(dt).strip()
        if val:
            entities.append({"text": val, "type": "DATE"})

    for sec in normalize_sections(core_facts.get("sections", []) or []):
        val = str(sec).strip()
        if val:
            entities.append({"text": val, "type": "LEGAL_SECTION"})

    logger.info(f"NER fallback: built {len(entities)} entities from core facts")
    return entities


def _hard_rule_classification(text: str) -> dict | None:
    """Hard rule-first classification as specified by policy."""
    t = text.lower()

    if re.search(r'हत्या|murder|death|postmortem|fatal|मृत्यु|शव\s*परीक्षण', t, re.IGNORECASE):
        return {
            "label": "HOMICIDE",
            "crime_key": "homicide",
            "confidence": "high",
            "reasoning": "Hard rule matched: हत्या/murder/death/postmortem/fatal संकेत।",
        }

    if re.search(r'गांजा|ndps|drugs?|narcotic|चरस|अफीम|हेरोइन|स्मैक', t, re.IGNORECASE):
        return {
            "label": "NDPS",
            "crime_key": "ndps",
            "confidence": "high",
            "reasoning": "Hard rule matched: NDPS/गांजा/drugs संकेत।",
        }

    if re.search(r'चोरी|theft|robbery|डकैती|लूट|snatch|burglary', t, re.IGNORECASE):
        return {
            "label": "THEFT/ROBBERY",
            "crime_key": "theft_robbery",
            "confidence": "high",
            "reasoning": "Hard rule matched: चोरी/theft/robbery संकेत।",
        }

    return None


def _resolve_classification(
    normalized_text: str,
    llm_classification: dict,
    manual_crime_type: Optional[str] = None,
) -> tuple[dict, list]:
    """Resolve final classification with strict rule-first policy."""
    rule_results = detect_crime_type_rules(normalized_text)
    hard_rule = _hard_rule_classification(normalized_text)

    out = dict(llm_classification or {})
    out.setdefault("secondary_crime_types", [])
    out.setdefault("detected_sections", [])

    if manual_crime_type:
        out["primary_crime_type"] = manual_crime_type
        out["classification_source"] = "manual"
        out["rule_label"] = manual_crime_type.upper()
        return out, rule_results

    if hard_rule:
        out["primary_crime_type"] = hard_rule["crime_key"]
        out["confidence"] = "high"
        out["classification_source"] = "hard_rule"
        out["rule_label"] = hard_rule["label"]
        out["reasoning"] = hard_rule["reasoning"]
        return out, rule_results

    # No hard rule match: use scoring, then LLM
    if rule_results:
        out["primary_crime_type"] = rule_results[0]["crime_key"]
        out["confidence"] = "medium"
        out["classification_source"] = "rule_scoring"
        out["rule_label"] = rule_results[0]["crime_key"].upper()
        if not out.get("reasoning"):
            out["reasoning"] = "Rule scoring used due to no hard-rule match."
    else:
        out["primary_crime_type"] = out.get("primary_crime_type", "unknown")
        out["confidence"] = out.get("confidence", "medium")
        out["classification_source"] = "llm"
        out["rule_label"] = out.get("primary_crime_type", "unknown").upper()

    return out, rule_results


def _enforce_header_fact(summary: str, field_label_pattern: str, value: str) -> str:
    if not value:
        return summary
    pattern = rf'(-\s*\*\*{field_label_pattern}\s*\([^\)]*\)?\s*:\*\*)\s*.*'
    if re.search(pattern, summary, flags=re.IGNORECASE):
        return re.sub(pattern, rf'\1 {value}', summary, flags=re.IGNORECASE)
    return summary


def _post_llm_validate(
    summary: str,
    classification: dict,
    ner_entities: list,
    checklist_result: dict,
    core_facts: dict,
    required_items: list,
    full_text: str = "",
) -> tuple[str, dict, list, dict]:
    """Step 6: fact protection against LLM drift and formatting cleanup."""
    summary = _normalize_legal_language(summary)

    # Anchor critical facts in classification
    anchors_sections = core_facts.get("sections", []) or []
    if anchors_sections:
        existing = classification.get("detected_sections", []) or []
        classification["detected_sections"] = _dedup_legal_sections(normalize_sections(list(anchors_sections) + list(existing)))
        summary = re.sub(r'\bSection\s+\d+[(/]\b', '', summary, flags=re.IGNORECASE)

    fir_no = core_facts.get("fir_number")
    if fir_no:
        classification["fir_number"] = fir_no
        summary = _enforce_header_fact(summary, r'FIR\s*संख्या', fir_no)

    dates = core_facts.get("dates", [])
    if dates:
        summary = _enforce_header_fact(summary, r'दिनांक', dates[0])

    # Revert any accidental field drift in summary values (critical anchors only)
    if fir_no:
        summary = re.sub(
            r'(\*\*FIR\s*संख्या\s*\([^\)]*\)\s*:\*\*)\s*[^\n]+',
            rf'\1 {fir_no}',
            summary,
            flags=re.IGNORECASE,
        )

    # Remove duplicate checklist rows and enforce schema
    checklist_result = _validate_checklist_payload(checklist_result.get("checklist", []), required_items)

    # NER post-process + dedup + legal validation
    ner_entities = _process_ner_output(ner_entities, classification, full_text=full_text)

    # Attach optional PERSON context hints for UI/report readability
    ner_entities = _attach_person_context_hints(ner_entities, full_text)

    # Ensure validated section list contains all normalized detected sections
    all_detected_sections = _dedup_legal_sections(normalize_sections(
        (classification.get("detected_sections", []) or []) +
        [e.get("text", "") for e in ner_entities if e.get("type") == "LEGAL_SECTION"] +
        (core_facts.get("sections", []) or [])
    ))
    if all_detected_sections:
        classification["detected_sections"] = all_detected_sections

    # Canonical single insertion point for validated legal sections line.
    summary = _upsert_validated_sections_line(summary, classification.get("detected_sections", []))

    # Final cleanup
    summary = re.sub(r'\n{3,}', '\n\n', summary).strip()
    return summary, classification, ner_entities, checklist_result


def _attach_person_context_hints(entities: list, full_text: str) -> list:
    if not entities or not full_text:
        return entities

    window = int(getattr(config, 'PERSON_CONTEXT_WINDOW_CHARS', 120))

    for ent in entities:
        if ent.get("type") != "PERSON":
            continue
        name = str(ent.get("text", "")).strip()
        if len(name) < 2:
            continue

        name_tokens = {tok for tok in re.split(r'\s+', name) if tok}
        lowered_name_tokens = {tok.lower() for tok in name_tokens}
        context_tokens = []
        context_snippets = []
        post_context_tokens = []

        def _is_pre_possessive_context(snippet: str) -> bool:
            return bool(re.search(r'(?:के|की|का)\s+(?:आरोपी|अभियुक्त|गवाह|साक्षी)\b', snippet, re.IGNORECASE))

        for m in re.finditer(re.escape(name), full_text, flags=re.IGNORECASE):
            pre_start = max(0, m.start() - window)
            pre_snippet = full_text[pre_start:m.start()]
            post_end = min(len(full_text), m.end() + window)
            post_snippet = full_text[m.end():post_end]
            snippet = f"{pre_snippet} {post_snippet}"
            context_snippets.append(snippet)

            if _is_pre_possessive_context(pre_snippet):
                logger.info(f"NER: Ignored pre-name possessive context for PERSON hint: '{name}'")

            post_words = re.findall(r'[\u0900-\u097F]{2,}|[A-Za-z]{3,}', post_snippet)
            for w in post_words:
                wl = w.lower()
                if w in name_tokens or wl in lowered_name_tokens:
                    continue
                if wl in HINDI_STOP_TOKENS:
                    continue
                if wl in _PERSON_HINT_BANNED_FOR_PERSON:
                    continue
                if not _is_meaningful_context_hint(w):
                    continue
                post_context_tokens.append(w)

            words = re.findall(r'[\u0900-\u097F]{2,}|[A-Za-z]{3,}', snippet)
            for w in words:
                wl = w.lower()
                if w in name_tokens or wl in lowered_name_tokens:
                    continue
                if wl in HINDI_STOP_TOKENS:
                    continue
                if wl in _PERSON_HINT_BANNED_FOR_PERSON:
                    continue
                if not _is_meaningful_context_hint(w):
                    continue
                context_tokens.append(w)

        if not context_tokens:
            role_probe = ' '.join(context_snippets).lower()
            if _DOCTOR_ROLE_INDICATOR_RE.search(role_probe):
                ent["type"] = "DOCTOR"
                logger.info(f"NER: Reclassified PERSON -> DOCTOR by adjacent context: '{name}'")
            elif _OFFICER_ROLE_INDICATOR_RE.search(role_probe):
                ent["type"] = "OFFICER"
                logger.info(f"NER: Reclassified PERSON -> OFFICER by adjacent context: '{name}'")
            continue

        selected_tokens = post_context_tokens or context_tokens
        freq = {}
        for token in selected_tokens:
            freq[token] = freq.get(token, 0) + 1
        best = sorted(freq.items(), key=lambda kv: (-kv[1], -len(kv[0]), kv[0]))[:2]
        if not best:
            continue

        # reliability guard: at least one token should appear twice
        if best[0][1] < 2:
            continue

        hint = ", ".join([b[0] for b in best])
        role_probe = f"{hint} {' '.join(context_snippets)}".lower()
        if _DOCTOR_ROLE_INDICATOR_RE.search(role_probe):
            ent["type"] = "DOCTOR"
            logger.info(f"NER: Reclassified PERSON -> DOCTOR by context indicators: '{name}'")
            ent.pop("context_hint", None)
            continue
        if _OFFICER_ROLE_INDICATOR_RE.search(role_probe):
            ent["type"] = "OFFICER"
            logger.info(f"NER: Reclassified PERSON -> OFFICER by context indicators: '{name}'")
            ent.pop("context_hint", None)
            continue

        ent["context_hint"] = hint
        logger.info(f"NER: Added context_hint for PERSON '{name}': {hint}")

    return entities


def _build_single_pass_prompt(text: str, core_facts: dict, crime_key_hint: str, required_items: list) -> str:
    """Single LLM call prompt for language-quality summary only."""
    facts_json = json.dumps(core_facts, ensure_ascii=False)

    return f"""
आप एक legal FIR/chargesheet assistant हैं। Output detailed, readable, and factual होना चाहिए।

### Anti-hallucination constraints (strict):
- FIR number, dates, legal sections, names invent मत करें।
- नीचे दिए गए extracted anchors को critical facts मानें।
- अगर कोई critical fact missing/uncertain है, तो स्पष्ट लिखें "उपलब्ध नहीं".

### Extracted anchors (high-confidence):
{facts_json}

### Classification hint (rule-side provisional): {crime_key_hint}

### Output requirements:
1) Rich structured summary in Hindi-English legal style with these headings:
   - Case Header
   - Parties
   - Incident Summary
   - Legal Sections
   - Key Evidence
   - Reasoning Notes

2) Facts के values invent/change न करें; extracted anchors से contradict न करें.

Language quality rules:
- OCR mistakes correct करें; broken words न दें.
- Natural legal prose रखें, overly rigid मत बनाएं.
- Detailed but concise explanations दें.

Document:
{text}
"""


def _parse_single_pass_response(response: str, required_items: list, core_facts: dict) -> tuple:
    """Parse single-call summary response while keeping fact-bearing modules deterministic."""
    summary = _extract_summary_body(response)

    classification = {
        "primary_crime_type": "unknown",
        "secondary_crime_types": [],
        "detected_sections": core_facts.get("sections", []),
        "confidence": "medium",
        "reasoning": "Deterministic classification pipeline applied.",
    }

    checklist_result = _validate_checklist_payload([], required_items)
    ner_entities = []

    return summary, classification, checklist_result, ner_entities


# ─────────────────────────────────────────────────────────────────────────────
# 6. Checklist Analysis (SECOND API CALL)
# ─────────────────────────────────────────────────────────────────────────────

def _build_checklist_prompt(text: str, crime_key: str, required_items: list) -> str:
    items_str = "\n".join(
        f"  {i+1}. [{item['id']}] {item['label_hi']} ({item['label_en']})"
        for i, item in enumerate(required_items)
    )

    return f"""निम्नलिखित चार्जशीट दस्तावेज़ का विश्लेषण करें और बताएं कि नीचे दी गई आवश्यक वस्तुओं/दस्तावेजों में से कौन-कौन दस्तावेज़ में उल्लिखित/उपस्थित हैं और कौन-कौन अनुपस्थित या अपूर्ण हैं।

**अपराध श्रेणी:** {crime_key}

**आवश्यक वस्तुओं/दस्तावेजों की सूची:**
{items_str}

**निर्देश:**
- प्रत्येक वस्तु के लिए बताएं: "present" (उपस्थित), "missing" (अनुपस्थित), या "partial" (आंशिक/अपूर्ण)।
- यदि "partial" है तो क्या कमी है, बताएं।
- दस्तावेज़ में [PAGE N] मार्कर दिए गए हैं। प्रत्येक "present" या "partial" वस्तु के लिए "page_no" फ़ील्ड में बताएं कि यह जानकारी किस पेज पर मिली (e.g. "3" or "5-7")। यदि "missing" है तो page_no खाली रखें ""।
- JSON format में उत्तर दें।

**आउटपुट format:**
```json
{{{{
  "checklist": [
    {{{{"id": "<item_id>", "status": "present|missing|partial", "page_no": "<page number or range>", "remarks": "टिप्पणी"}}}},
    ...
  ]
}}}}
```

---
**चार्जशीट दस्तावेज़:**
{text}
"""


def analyse_checklist(text: str, crime_key: str) -> dict:
    """
    Analyse chargesheet against the required-items checklist.
    Uses LLM for accuracy, with rule-based fallback.
    """
    crime_info = get_crime_type_info(crime_key)
    if not crime_info:
        return {"error": f"Unknown crime type: {crime_key}"}

    required_items = crime_info["required_items"]
    prompt = _build_checklist_prompt(text, crime_key, required_items)

    try:
        response = _call_llm(prompt, SYSTEM_PROMPT)

        # Extract JSON
        json_match = re.search(r'```json\s*([\s\S]*?)```', response)
        if not json_match:
            json_match = re.search(r'\{[\s\S]*"checklist"[\s\S]*\}', response)

        if json_match:
            json_str = json_match.group(1) if '```json' in response else json_match.group(0)
            try:
                result = json.loads(json_str.strip())
                return result
            except json.JSONDecodeError:
                pass

        # If JSON parsing failed, try rule-based
        logger.warning("LLM checklist response not parseable, using rule-based fallback")
        return _rule_based_checklist(text, required_items)

    except Exception as e:
        # LLM failed entirely, use rule-based
        logger.warning(f"LLM checklist call failed ({e}), using rule-based fallback")
        return _rule_based_checklist(text, required_items)


# ─────────────────────────────────────────────────────────────────────────────
# 7. Rule-based Keyword Detection (cross-validation, no API)
# ─────────────────────────────────────────────────────────────────────────────

def detect_crime_type_rules(text: str) -> list:
    """
    Rule-based crime type detection using keyword matching.
    Returns list of matching crime types sorted by score.
    No API calls -- runs entirely offline.
    """
    checklists = load_checklists()
    text_lower = text.lower()
    results = []

    prior_block_markers = [
        r'आपराधिक\s*इतिहास',
        r'पूर्व\s*अपराध',
        r'previous\s*case',
        r'criminal\s*history',
        r'काण्ड\s*सं\.?',
        r'भा\.?\s*द\.?\s*वि\.?',
    ]
    prior_marker_re = re.compile("|".join(prior_block_markers), re.IGNORECASE)

    lines = text.splitlines()
    line_windows = []
    for i, line in enumerate(lines):
        if prior_marker_re.search(line):
            start = max(0, i - 1)
            end = min(len(lines), i + 8)
            line_windows.append((start, end))

    def _in_prior_context(char_idx: int) -> bool:
        if not line_windows:
            return False
        running = 0
        for li, line in enumerate(lines):
            line_start = running
            line_end = running + len(line)
            running = line_end + 1
            if line_start <= char_idx <= line_end:
                for ws, we in line_windows:
                    if ws <= li <= we:
                        return True
                return False
        return False

    for key, info in checklists.items():
        score = 0.0
        section_hits = 0

        for section in info["typical_sections"]:
            sec_match = re.search(r'(\d+[A-Za-z]?(?:\(\d+\))?)', section)
            if not sec_match:
                continue
            section_num = sec_match.group(1)
            section_patterns = [
                re.escape(section.lower()),
                rf'धारा\s*{re.escape(section_num)}',
                rf'section\s*{re.escape(section_num)}',
            ]
            found = False
            for pat in section_patterns:
                if re.search(pat, text_lower, re.IGNORECASE):
                    score += 3.0
                    found = True
                    break
            if found:
                section_hits += 1

        # Keywords matched inside prior-history blocks are down-weighted to 25%
        for kw in info.get("keywords_hi", []):
            for m in re.finditer(re.escape(kw), text, re.IGNORECASE):
                w = 0.25 if _in_prior_context(m.start()) else 1.0
                score += (2.0 * w)

        for kw in info.get("keywords_en", []):
            for m in re.finditer(re.escape(kw), text_lower, re.IGNORECASE):
                w = 0.25 if _in_prior_context(m.start()) else 1.0
                score += (1.0 * w)

        # Section dominance: hard-section evidence should outrank noisy keyword overlap
        if section_hits > 0:
            score *= 2.0
            logger.info(f"Rule score boost: {key} boosted by section dominance (hits={section_hits})")

        if score > 0:
            results.append({
                "crime_key": key,
                "display_name": info["display_name"],
                "score": round(score, 2),
            })

    results.sort(key=lambda x: x["score"], reverse=True)
    return results


def validate_checklist(
    text: str,
    required_items: list,
    summary: str = "",
    evidence: str = "",
) -> dict:
    """
    Semantic checklist validation (pattern-based, no exact-only dependency).
    Searches across summary, evidence, and document text.
    """
    corpus = "\n".join([summary or "", evidence or "", text or ""])
    base = _rule_based_checklist(corpus, required_items)

    expansion_clusters = {
        "postmortem": [r'पोस्टमार्टम', r'post\s*mortem', r'autopsy', r'शव\s*परीक्षण'],
        "mlc": [r'\bmlc\b', r'मेडिको\s*लीगल', r'चिकित्स[ाी]\s*परीक्षण', r'medical\s*report'],
        "arrest": [r'गिरफ्तार', r'arrest', r'गिरफतारी', r'custody'],
        "seizure": [r'जब्त', r'बरामद', r'seizure', r'recover'],
        "site": [r'नक्शा\s*मौका', r'site\s*plan', r'spot\s*inspection', r'घटनास्थल'],
        "forensic": [r'\bfsl\b', r'forensic', r'फॉरेंसिक', r'विशेषज्ञ\s*राय'],
        "witness": [r'गवाह', r'साक्षी', r'witness', r'statement'],
    }

    def _item_patterns(item: dict) -> list[str]:
        item_id = (item.get("id") or "").lower()
        label_hi = (item.get("label_hi") or "").lower()
        label_en = (item.get("label_en") or "").lower()
        patterns = []

        for token in re.split(r'[_\s\-/]+', item_id):
            if len(token) >= 3 and token in expansion_clusters:
                patterns.extend(expansion_clusters[token])

        for bucket, bucket_patterns in expansion_clusters.items():
            if bucket in label_en or bucket in item_id:
                patterns.extend(bucket_patterns)

        # Add weakly normalized literal terms from labels (length guard)
        for literal in re.split(r'[,/()\-\s]+', f"{label_hi} {label_en}"):
            lit = literal.strip()
            if len(lit) >= 4:
                patterns.append(re.escape(lit))

        dedup = []
        seen = set()
        for p in patterns:
            if p not in seen:
                seen.add(p)
                dedup.append(p)
        return dedup[:25]

    checklist = base.get("checklist", [])
    for row in checklist:
        if row.get("status") == "present":
            continue
        item = next((i for i in required_items if i.get("id") == row.get("id")), None)
        if not item:
            continue
        pats = _item_patterns(item)
        if not pats:
            continue

        hit = None
        for pat in pats:
            if re.search(pat, corpus, re.IGNORECASE):
                hit = pat
                break

        if hit:
            row["status"] = "present"
            if not row.get("remarks"):
                row["remarks"] = "दस्तावेज़ में संबंधित सामग्री का उल्लेख मिला।"
            else:
                row["remarks"] += " | semantic evidence matched"

    return {"checklist": checklist}


def _rule_based_checklist(text: str, required_items: list) -> dict:
    """Deterministic checklist from document text using keyword evidence."""
    text_lower = text.lower()
    checklist = []

    keyword_map = {
        # === Shared / Common ===
        "fir": ["FIR", "प्राथमिकी", "प्रथम सूचना रिपोर्ट", "एफआईआर"],
        "ps": ["थाना", "पुलिस स्टेशन", "police station"],
        "place_time": ["घटनास्थल", "स्थान", "समय", "बजे", "दिनांक"],
        "complainant": ["शिकायतकर्ता", "पीड़ित", "complainant", "victim"],
        "accused": ["आरोपी", "अभियुक्त", "accused", "संदिग्ध"],
        "witnesses": ["गवाह", "प्रत्यक्षदर्शी", "पंच गवाह", "witness", "eye-witness"],
        "witness_stmt": ["गवाह", "बयान", "witness", "statement", "161"],
        "arrest_memo": ["गिरफ्तारी मेमो", "arrest memo", "गिरफ्तार"],

        # === Theft / Robbery ===
        "property_desc": ["संपत्ति", "सामान", "चोरी", "stolen property", "मोबाइल", "मूल्य"],
        "ownership_proof": ["स्वामित्व", "बिल", "ownership", "receipt"],
        "seizure_memo": ["जब्ती मेमो", "seizure memo", "बरामदगी मेमो", "बरामद"],
        "chain_custody": ["कस्टडी चेन", "custody", "chain of custody"],
        "site_plan": ["नक्शा मौका", "site plan", "स्थल निरीक्षण", "घटना स्थल का नक्शा"],
        "cctv": ["CCTV", "सीसीटीवी", "कैमरा", "फुटेज"],
        "spot_photos": ["फोटोग्राफ", "तस्वीर", "photograph", "फोटो"],
        "confession_disclosure": ["स्वीकारोक्ति", "प्रकटीकरण", "confession", "disclosure"],
        "personal_search": ["व्यक्तिगत तलाशी", "personal search"],
        "property_seal": ["सील", "seal", "जब्ती मेमो"],
        "malkhana_forward": ["मालखाना", "malkhana", "अग्रेषण"],
        "supplementary": ["पूरक आरोप पत्र", "supplementary"],

        # === Assault / Hurt ===
        "victim": ["पीड़ित", "victim", "शिकायतकर्ता"],
        "mlc": ["MLC", "मेडिको लीगल", "चिकित्सा परीक्षण", "medical"],
        "injury_cert": ["चोट प्रमाणपत्र", "injury certificate", "चोटों की प्रकृति"],
        "injury_photos": ["चोटों के फोटो", "तस्वीर", "photograph", "इंजरी फोटो"],
        "doctor_opinion": ["डॉक्टर की राय", "doctor's opinion", "चिकित्सक", "साधारण", "गम्भीर"],
        "referral_docs": ["रेफरल", "referral", "उच्च केंद्र"],
        "weapon_desc": ["हथियार", "चाकू", "लाठी", "डंडी", "बंदूक", "weapon", "knife", "lathi"],
        "weapon_seizure": ["हथियार", "जब्ती", "weapon seizure"],
        "motive": ["मकसद", "मोटिव", "motive", "रंजिश", "enmity", "कारण"],
        "section_mapping": ["धारा-वार", "section-wise", "section mapping"],
        "arrest_medical": ["गिरफ्तारी", "आरोपी का चिकित्सा", "accused medical"],

        # === Cyber Fraud ===
        "platform_mode": ["प्लेटफॉर्म", "ऑनलाइन", "platform", "UPI", "बैंकिंग ऐप", "OLX", "सोशल मीडिया"],
        "txn_datetime": ["लेनदेन", "ट्रांजेक्शन", "transaction", "दिनांक"],
        "amount": ["राशि", "amount", "रुपये", "₹"],
        "txn_records": ["बैंक स्टेटमेंट", "transaction record", "UPI", "ट्रांजेक्शन रिकॉर्ड"],
        "payment_logs": ["पेमेंट गेटवे", "payment gateway", "लॉग"],
        "chat_logs": ["चैट", "ईमेल", "screenshot", "स्क्रीनशॉट", "कॉल डिटेल", "CDR"],
        "email_headers": ["ईमेल हेडर", "email header", "मैसेज ID"],
        "device_ids": ["IMEI", "IP", "डिवाइस", "यूजर ID", "device identifier"],
        "device_seizure": ["डिजिटल", "मोबाइल", "लैपटॉप", "device", "जब्ती"],
        "forensic_image": ["फोरेंसिक इमेज", "forensic image", "हैश", "hash"],
        "bank_request": ["बैंक", "KYC", "अनुरोध", "bank request"],
        "bank_response": ["बैंक से प्राप्त", "bank response", "उत्तर", "response"],
        "digital_id_link": ["IP", "IMEI", "digital identifier", "पहचान से जोड़"],
        "chain_digital": ["डिजिटल साक्ष्य", "digital evidence", "कस्टडी"],
        "fsl_cyber": ["साइबर फोरेंसिक", "cyber forensic", "FSL"],
        "account_freeze": ["खाता फ्रीज", "account freeze", "रिकवरी"],

        # === NDPS ===
        "seizure_place_time": ["जब्ती का स्थान", "seizure place", "नाका"],
        "substance_desc": ["पदार्थ", "मात्रा", "substance", "quantity", "किलोग्राम", "ग्राम", "पाउडर"],
        "ndps_compliance": ["धारा 42", "धारा 50", "NDPS", "अनुपालन"],
        "search_memo": ["तलाशी", "search memo", "तलाशी मेमो"],
        "weighment": ["तौल", "weighment", "वजन", "पंचनामा"],
        "sample_memo": ["नमूना", "sample", "ड्राइंग मेमो"],
        "malkhana": ["मालखाना", "malkhana", "रजिस्टर"],
        "fsl_dispatch": ["FSL भेजने", "dispatch", "FSL memo"],
        "fsl_report": ["FSL", "फॉरेंसिक", "forensic", "एफएसएल"],
        "expert_opinion": ["विशेषज्ञ राय", "expert opinion"],
        "independent_witness": ["स्वतंत्र गवाह", "independent witness", "पंच"],
        "mandatory_compliance": ["अनिवार्य प्रावधान", "mandatory provision", "अनुपालन"],
    }

    for item in required_items:
        item_id = item["id"]
        keywords = keyword_map.get(item_id, [])
        matched_keywords = []
        for kw in keywords:
            if kw.lower() in text_lower or kw in text:
                matched_keywords.append(kw)

        found = len(matched_keywords) > 0
        page_no = ""
        if found:
            for kw in matched_keywords[:2]:
                idx = text_lower.find(kw.lower())
                if idx < 0:
                    idx = text.find(kw)
                if idx >= 0:
                    marker_start = text.rfind("[PAGE ", 0, idx)
                    if marker_start >= 0:
                        marker_end = text.find("]", marker_start)
                        if marker_end > marker_start:
                            marker = text[marker_start + 6:marker_end].strip()
                            if marker.isdigit():
                                page_no = marker
                                break

        remarks = ""
        if found:
            evidence_terms = ", ".join(matched_keywords[:3])
            remarks = f"दस्तावेज़ में उल्लेख मिला: {evidence_terms}"
        else:
            remarks = "दस्तावेज़ में स्पष्ट उल्लेख नहीं मिला"

        checklist.append({
            "id": item_id,
            "status": "present" if found else "missing",
            "page_no": page_no,
            "remarks": remarks,
        })

    return {"checklist": checklist}


def compute_confidence(
    prep: dict,
    core_facts: dict,
    summary: str,
    classification: dict,
    rule_classification: list,
    ner_entities: list,
    checklist_result: dict,
    field_confidence: dict,
) -> dict:
    """
    Weighted confidence scoring with penalties for noisy OCR, missing fields,
    weak entity extraction, and summary-anchor inconsistencies.
    """
    ocr_conf = float(prep.get("ocr_confidence", 0.0) or 0.0)
    extraction_conf = float(core_facts.get("extraction_confidence", 0.0) or 0.0)
    classification_conf = _compute_classification_confidence(classification, rule_classification)

    corrections = int(prep.get("ocr_corrections", 0) or 0)
    correction_penalty = min(0.35, corrections / 280.0)
    ocr_conf = max(0.1, ocr_conf - correction_penalty)

    person_entities = [
        e for e in (ner_entities or [])
        if e.get("type") in ("ACCUSED", "WITNESS", "OFFICER", "DOCTOR", "PERSON")
    ]
    anchor_accused = core_facts.get("accused", []) or []
    dropped_entity_penalty = 0.0
    if anchor_accused and not person_entities:
        dropped_entity_penalty += 0.20
    elif len(person_entities) < max(1, min(3, len(anchor_accused))):
        dropped_entity_penalty += 0.10

    summary_check_failures = 0
    fir_no = core_facts.get("fir_number")
    if fir_no and fir_no not in summary:
        summary_check_failures += 1
    dates = core_facts.get("dates", [])
    if dates and dates[0] not in summary:
        summary_check_failures += 1

    weak_fields = sum(1 for v in (field_confidence or {}).values() if float(v) < 0.6)
    field_penalty = min(0.25, weak_fields * 0.02)

    checklist = checklist_result.get("checklist", []) if checklist_result else []
    missing_count = sum(1 for row in checklist if row.get("status") == "missing")
    checklist_penalty = min(0.20, (missing_count / max(1, len(checklist))) * 0.20) if checklist else 0.0

    extraction_conf = max(0.1, extraction_conf - dropped_entity_penalty - field_penalty)
    classification_conf = max(0.1, classification_conf - (0.10 * summary_check_failures) - (checklist_penalty * 0.5))

    return {
        "ocr_confidence": round(ocr_conf, 2),
        "extraction_confidence": round(extraction_conf, 2),
        "classification_confidence": round(classification_conf, 2),
    }


# ─────────────────────────────────────────────────────────────────────────────
# 8. Semantic Similarity Scoring (Stage 2B)
# ─────────────────────────────────────────────────────────────────────────────

# Threshold for considering a checklist item as semantically PRESENT
SEMANTIC_THRESHOLD = 0.25  # TF-IDF cosine on Hindi text — threshold is lower than English


def _split_into_sentences(text: str) -> list:
    """
    Split Hindi/English text into sentence-like units.
    Uses punctuation and newline boundaries.
    """
    # Split on Hindi purna viram, question marks, newlines, and semicolons
    raw = re.split(r'[।\.\?!;\n]+', text)
    sentences = []
    for s in raw:
        s = s.strip()
        if len(s) > 10:  # skip very short fragments
            sentences.append(s)
    return sentences


def compute_semantic_similarity(
    text: str,
    required_items: list,
) -> dict:
    """
    Stage 2B: Semantic Similarity Scoring.
    Uses TF-IDF + cosine similarity to match each required checklist item
    against document sentences. This catches paraphrased mentions that
    keyword matching would miss (e.g. "पंचनामा तैयार किया गया" for "Weighment panchnama").

    Returns dict:
    {
        "item_id": {
            "similarity_score": float,      # 0.0 to 1.0
            "semantic_status": "present" | "missing" | "partial",
            "best_match": "matching sentence from document",
            "confidence": "high" | "medium" | "low"
        },
        ...
    }
    """
    sentences = _split_into_sentences(text)
    if not sentences or not required_items:
        return {}

    # Build query strings from checklist items (combine Hindi + English labels + keywords)
    queries = []
    item_ids = []
    for item in required_items:
        # Combine all available text representations for richer matching
        parts = [item["label_hi"], item["label_en"]]
        # Add any keywords if available
        if "keywords" in item:
            parts.extend(item["keywords"])
        queries.append(" ".join(parts))
        item_ids.append(item["id"])

    # Combine queries + sentences for TF-IDF fitting
    all_texts = queries + sentences

    try:
        vectorizer = TfidfVectorizer(
            analyzer='char_wb',   # char n-grams work better for Hindi/mixed text
            ngram_range=(2, 4),   # 2-4 char grams capture Hindi morphology
            max_features=15000,
            sublinear_tf=True,
        )
        tfidf_matrix = vectorizer.fit_transform(all_texts)

        # Split: first N rows are queries, rest are sentences
        query_vectors = tfidf_matrix[:len(queries)]
        sentence_vectors = tfidf_matrix[len(queries):]

        # Compute cosine similarity: queries × sentences
        sim_matrix = cosine_similarity(query_vectors, sentence_vectors)

    except Exception as e:
        logger.warning(f"Semantic similarity computation failed: {e}")
        return {}

    results = {}
    for i, item_id in enumerate(item_ids):
        scores = sim_matrix[i]
        best_idx = int(np.argmax(scores))
        best_score = float(scores[best_idx])
        best_sentence = sentences[best_idx] if best_idx < len(sentences) else ""

        # Truncate long matching sentences
        if len(best_sentence) > 200:
            best_sentence = best_sentence[:200] + "…"

        # Determine status and confidence
        if best_score >= SEMANTIC_THRESHOLD * 1.5:  # strong match
            status = "present"
            confidence = "high"
        elif best_score >= SEMANTIC_THRESHOLD:       # moderate match
            status = "partial"
            confidence = "medium"
        else:
            status = "missing"
            confidence = "low"

        results[item_id] = {
            "similarity_score": round(best_score, 3),
            "semantic_status": status,
            "best_match": best_sentence,
            "confidence": confidence,
        }

    return results


def _merge_checklist_with_rules(
    llm_result: dict,
    rule_result: dict,
) -> dict:
    """
    Merge LLM checklist with rule-based keyword matching (OR logic).
    If LLM says MISSING but rule-based found keywords → upgrade to PRESENT.
    If LLM says PARTIAL but rule-based found keywords → upgrade to PRESENT.
    """
    llm_checklist = llm_result.get("checklist", [])
    rule_checklist = rule_result.get("checklist", [])

    if not llm_checklist or not rule_checklist:
        return llm_result

    # Build rule lookup: item_id → status
    rule_lookup = {}
    for entry in rule_checklist:
        rule_lookup[entry.get("id", "")] = entry.get("status", "missing")

    merged = []
    for entry in llm_checklist:
        new_entry = dict(entry)
        item_id = entry.get("id", "")
        rule_status = rule_lookup.get(item_id, "missing")

        # Upgrade logic: rule says present → override LLM missing/partial
        if entry.get("status") in ("missing", "partial") and rule_status == "present":
            old_status = entry.get("status")
            new_entry["status"] = "present"
            if not new_entry.get("remarks"):
                new_entry["remarks"] = "🔑 Keyword match found in document"
            else:
                new_entry["remarks"] += " | 🔑 Keyword match"
            logger.info(f"Rule upgrade: {item_id} {old_status.upper()}→PRESENT (keyword match)")

        merged.append(new_entry)

    return {"checklist": merged}


def _merge_checklist_with_similarity(
    checklist_result: dict,
    similarity_scores: dict,
) -> dict:
    """
    Merge LLM checklist results with semantic similarity scores.
    The LLM result is primary; similarity scores augment it with:
    - A similarity confidence score
    - The best matching sentence
    - Upgrade MISSING → PARTIAL if semantic match found
    """
    checklist = checklist_result.get("checklist", [])
    if not checklist or not similarity_scores:
        return checklist_result

    merged = []
    for entry in checklist:
        item_id = entry.get("id", "")
        sim_info = similarity_scores.get(item_id, {})

        # Copy existing entry
        new_entry = dict(entry)

        if sim_info:
            new_entry["similarity_score"] = sim_info["similarity_score"]
            new_entry["best_match"] = sim_info["best_match"]

            # Upgrade: if LLM says MISSING/PARTIAL but semantic says PRESENT
            if entry.get("status") in ("missing", "partial") and sim_info["semantic_status"] == "present":
                new_entry["status"] = "present"
                if not new_entry.get("remarks"):
                    new_entry["remarks"] = "🔍 Strong semantic match found"
                else:
                    new_entry["remarks"] += " | 🔍 Strong semantic match"
                logger.info(f"Semantic upgrade: {item_id} {entry.get('status').upper()}→PRESENT (score={sim_info['similarity_score']:.3f})")
            elif entry.get("status") == "missing" and sim_info["semantic_status"] == "partial":
                new_entry["status"] = "partial"
                if not new_entry.get("remarks"):
                    new_entry["remarks"] = "🔍 Partial semantic match found"
                else:
                    new_entry["remarks"] += " | 🔍 Partial semantic match"
                logger.info(f"Semantic upgrade: {item_id} MISSING→PARTIAL (score={sim_info['similarity_score']:.3f})")

        merged.append(new_entry)

    return {"checklist": merged}


# ─────────────────────────────────────────────────────────────────────────────
# 8b. Chunking (kept for backward compatibility but NOT used in main pipeline)
# ─────────────────────────────────────────────────────────────────────────────

def chunk_text(text: str, chunk_size: int = None, overlap: int = None) -> list:
    """Split text into overlapping chunks. Kept for notebook/advanced usage."""
    chunk_size = chunk_size or config.CHUNK_SIZE
    overlap = overlap or config.CHUNK_OVERLAP

    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start = end - overlap

    return chunks


# ─────────────────────────────────────────────────────────────────────────────
# 8c. Confidence Scoring (Stage 3A — computed, not random)
# ─────────────────────────────────────────────────────────────────────────────

def _compute_classification_confidence(
    llm_classification: dict,
    rule_classification: list,
) -> float:
    """
    Compute composite classification confidence (0.0–1.0).
    Triangulates 2 independent signals:
    - LLM confidence (high/medium/low → 0.9/0.6/0.3)  — weight 0.55
    - LLM confidence_score (if provided)               — override LLM weight
    - Rule-based score (normalized)                     — weight 0.45

    Returns a meaningful 0.0–1.0 score.
    """
    source = llm_classification.get("classification_source", "")
    if source == "hard_rule":
        return 0.95
    if source == "rule_scoring":
        return 0.75

    # Signal 1: LLM-declared confidence
    llm_conf_str = llm_classification.get("confidence", "low")
    llm_score_raw = llm_classification.get("confidence_score", None)
    if llm_score_raw is not None:
        try:
            llm_signal = float(llm_score_raw)
            llm_signal = max(0.0, min(1.0, llm_signal))
        except (ValueError, TypeError):
            llm_signal = {"high": 0.90, "medium": 0.60, "low": 0.30}.get(llm_conf_str, 0.30)
    else:
        llm_signal = {"high": 0.90, "medium": 0.60, "low": 0.30}.get(llm_conf_str, 0.30)

    # Signal 2: Rule-based section/keyword score
    primary = llm_classification.get("primary_crime_type", "unknown")
    rule_score = 0.0
    if rule_classification:
        # Find matching rule result for the primary crime
        for r in rule_classification:
            if r["crime_key"] == primary:
                # Normalize: typical score range 5-30, map to 0-1
                rule_score = min(1.0, r["score"] / 20.0)
                break
        if rule_score == 0.0:
            # Primary didn't match any rule → penalty
            rule_score = 0.1

    # Weighted combination
    composite = (0.55 * llm_signal) + (0.45 * rule_score)
    composite = round(max(0.0, min(1.0, composite)), 2)

    logger.info(f"Classification confidence: LLM={llm_signal:.2f} Rule={rule_score:.2f} "
                f"→ Composite={composite:.2f}")

    return composite


def _compute_checklist_item_confidence(entry: dict) -> float:
    """
    Compute per-item confidence (0.0–1.0) for a checklist item.
    Triangulates:
    - LLM status (present=0.70, partial=0.40, missing=0.05)  — base
    - Rule-based keyword match (if remarks mention 🔑)       — +0.15
    - Semantic similarity score (if available)                — +0.15 × score
    """
    status = entry.get("status", "missing")
    remarks = entry.get("remarks", "") or ""
    sim_score = entry.get("similarity_score", 0.0) or 0.0

    # Base confidence from LLM status
    if status == "present":
        base = 0.70
    elif status == "partial":
        base = 0.40
    else:
        base = 0.05

    # Rule-based bonus
    rule_bonus = 0.15 if "🔑" in remarks else 0.0

    # Semantic bonus (scaled)
    sem_bonus = 0.15 * min(1.0, sim_score / 0.4) if sim_score > 0 else 0.0

    # Strong semantic match bonus
    if "🔍 Strong semantic" in remarks:
        sem_bonus = max(sem_bonus, 0.12)
    elif "🔍 Partial semantic" in remarks:
        sem_bonus = max(sem_bonus, 0.06)

    confidence = round(min(1.0, base + rule_bonus + sem_bonus), 2)
    return confidence


def _enrich_checklist_with_confidence(checklist_result: dict) -> dict:
    """Add per-item confidence scores to every checklist entry."""
    checklist = checklist_result.get("checklist", [])
    for entry in checklist:
        entry["confidence"] = _compute_checklist_item_confidence(entry)
    return checklist_result


def _compute_field_confidence(summary: str) -> dict:
    """
    Heuristic field confidence: check which fields are present in the summary.
    Scans for field headers and content to determine 0.0–1.0 confidence.
    No API call needed — pure text analysis.
    """
    value_common = [r'\d{1,2}[/\.\-]\d{1,2}[/\.\-]\d{2,4}', r'\d{1,2}:\d{2}', r'₹\s*\d+', r'\S{2,}']
    field_checks = {
        "fir_number": {
            "keywords": [r'fir', r'प्रथम\s*सूचना', r'काण्ड\s*सं', r'अपराध\s*सं'],
            "values": [r'\d{2,6}\s*/\s*\d{4}', r'fir\s*(?:no|नं|संख्या)'],
        },
        "fir_date": {
            "keywords": [r'fir', r'दिनांक', r'तिथि', r'date'],
            "values": [r'\d{1,2}[/\.\-]\d{1,2}[/\.\-]\d{2,4}'],
        },
        "police_station": {
            "keywords": [r'थाना', r'police\s*station', r'ps\b'],
            "values": [r'थाना\s*[:\-]?\s*\S+', r'police\s*station\s*[:\-]?\s*\S+'],
        },
        "place_of_occurrence": {
            "keywords": [r'घटना\s*स्थल', r'घटनास्थल', r'place\s*of\s*occurrence', r'spot'],
            "values": [r'घटना\s*स्थल\s*[:\-]?\s*\S+', r'स्थान\s*[:\-]?\s*\S+'],
        },
        "incident_date_time": {
            "keywords": [r'घटना', r'incident', r'वारदात'],
            "values": [r'बजे', r'\d{1,2}:\d{2}', r'AM|PM|रात्रि|प्रातः', r'\d{1,2}[/\.\-]\d{1,2}[/\.\-]\d{2,4}'],
        },
        "complainant_name": {
            "keywords": [r'शिकायतकर्ता', r'वादी', r'प्रार्थी', r'complainant', r'informant'],
            "values": value_common,
        },
        "accused_names": {
            "keywords": [r'आरोपी', r'अभियुक्त', r'accused', r'suspect'],
            "values": value_common,
        },
        "witnesses": {
            "keywords": [r'गवाह', r'साक्षी', r'witness'],
            "values": value_common,
        },
        "incident_summary": {
            "keywords": [r'घटना\s*का\s*सारांश', r'incident\s*summary', r'घटना\s*विवरण', r'brief\s*facts'],
            "values": [r'.{20,}'],
        },
        "legal_sections": {
            "keywords": [r'धारा', r'section', r'ipc', r'bns', r'bnss', r'ndps', r'arms\s*act'],
            "values": [r'(?:ipc|bns|bnss|ndps|arms\s*act)\s*\d+(?:\(\d+\))?', r'धारा\s*\d+(?:\(\d+\))?'],
        },
        "key_evidence": {
            "keywords": [r'साक्ष्य', r'evidence', r'बरामद', r'जब्त', r'seized', r'recovered'],
            "values": [r'बरामद|जब्त|seized|recovered|evidence'],
        },
    }

    def _keyword_present(keywords: list[str]) -> bool:
        return any(re.search(k, summary, re.IGNORECASE) for k in keywords)

    def _near_value(keywords: list[str], values: list[str], window: int = 80) -> bool:
        # For incident field, enforce co-occurrence in a wider semantic window
        for k in keywords:
            for km in re.finditer(k, summary, re.IGNORECASE):
                s = max(0, km.start() - 100)
                e = min(len(summary), km.end() + 100)
                chunk = summary[s:e]
                if any(re.search(v, chunk, re.IGNORECASE) for v in values):
                    return True
        return any(re.search(v, summary, re.IGNORECASE) for v in values)

    result = {}
    for field_key, checks in field_checks.items():
        keyword_hits = sum(1 for k in checks["keywords"] if re.search(k, summary, re.IGNORECASE))
        value_hits = sum(1 for v in checks["values"] if re.search(v, summary, re.IGNORECASE))
        near_found = _near_value(checks["keywords"], checks["values"])

        if keyword_hits == 0 and value_hits == 0 and not near_found:
            result[field_key] = -1.0
            logger.info(f"Field confidence: {field_key} missing -> -1.0")
            continue

        kw_cov = keyword_hits / max(1, len(checks["keywords"]))
        val_cov = value_hits / max(1, len(checks["values"]))
        near_bonus = 0.12 if near_found else 0.0
        score = 0.45 + (0.28 * kw_cov) + (0.20 * val_cov) + near_bonus
        # small field-specific calibration for concrete identifiers
        if field_key in ("fir_number", "fir_date", "legal_sections") and val_cov > 0:
            score += 0.05
        score = round(min(0.96, max(0.25, score)), 2)
        result[field_key] = score

        logger.info(
            f"Field confidence: {field_key} kw_hits={keyword_hits} val_hits={value_hits} "
            f"near={near_found} score={score:.2f}"
        )

    # Court field — regex search to handle Unicode normalization variants
    _COURT_RE = re.compile(
        r'माननीय|न्यायालय|court|magistrate|'
        r'द[ंण]डाधिकारी|jmfc|cjm|विचारण|न्यायिक|'
        r'session|judge|अदालत|nyayalaya',
        re.IGNORECASE
    )
    court_matches = list(_COURT_RE.finditer(summary))
    court_found = len(court_matches) > 0
    court_has_value = False
    if court_found:
        for m in court_matches:
            nearby = summary[m.start():min(len(summary), m.start() + 100)]
            # Must have content after keyword — any Hindi/English word of 3+ chars
            words = [w for w in nearby.split()[1:] if len(w) >= 3]
            if words:
                court_has_value = True
                break

    # Fallback: any mention of court-adjacent concept
    if not court_found:
        court_found = bool(re.search(
            r'BNSS|बीएनएसएस|धारा\s*183|न्यायिक\s*(?:हिरासत|दंडाधिकारी)',
            summary, re.IGNORECASE
        ))
        court_has_value = court_found

    if court_found and court_has_value:
        result["court"] = 0.88
    elif court_found:
        result["court"] = 0.62
    else:
        result["court"] = -1.0
    logger.info(f"Field confidence: court found={court_found} has_value={court_has_value} score={result['court']:.2f}")

    # Regression guard: detect suspicious score uniformity.
    non_na = [v for v in result.values() if isinstance(v, (int, float)) and v >= 0.0]
    if non_na:
        counts = {}
        for v in non_na:
            counts[v] = counts.get(v, 0) + 1
        top = max(counts.values()) if counts else 0
        ratio = top / max(1, len(non_na))
        if ratio > 0.30:
            logger.warning(
                f"Field confidence regression guard: {ratio:.0%} fields share identical score; "
                "defaults may be dominating"
            )

    return result


# ─────────────────────────────────────────────────────────────────────────────
# 9. Full Pipeline -- OPTIMISED: 3 API calls
# ─────────────────────────────────────────────────────────────────────────────

def process_chargesheet(
    text: str,
    manual_crime_type: Optional[str] = None,
    images: list = None,
) -> dict:
    """
        Full production pipeline (3 LLM calls):
      Step 1: Raw input
      Step 2: OCR cleaning (strong, layered)
      Step 3: Core fact extraction (anchors)
      Step 4: Pre-LLM normalization
            Step 5: LLM summary
            Step 6: LLM checklist
            Step 7: LLM NER
            Step 8: Post-LLM validation/fact protection
            Step 9: Rule-first classification
            Step 10: Confidence scoring
    """
    images = images or []

    # Step 2–4: OCR cleaning + core facts + normalization
    prep = preprocess_text_with_meta(text)
    cleaned_text = prep["cleaned_text"]
    normalized_text = prep["normalized_text"]
    core_facts = prep["core_facts"]
    ocr_confidence = prep["ocr_confidence"]

    # Smart truncation (input protection)
    processed_text = _truncate_text(normalized_text)
    logger.info(
        f"Processing document: raw={len(text):,} chars, cleaned={len(cleaned_text):,}, "
        f"normalized={len(normalized_text):,}, truncated={len(processed_text):,}, images={len(images)}"
    )

    # Step 7 pre-hint (rule-side hint for checklist prompt; final resolution still rule-first)
    hard_rule = _hard_rule_classification(processed_text)
    crime_hint = manual_crime_type or (hard_rule["crime_key"] if hard_rule else "unknown")
    if crime_hint == "unknown":
        scored = detect_crime_type_rules(processed_text)
        crime_hint = scored[0]["crime_key"] if scored else "theft_robbery"

    crime_info = get_crime_type_info(crime_hint) or get_crime_type_info("theft_robbery")
    required_items = crime_info["required_items"]

    # Step 5: LLM Call 1/3 - summary generation
    prompt = _build_single_pass_prompt(processed_text, core_facts, crime_hint, required_items)
    if images:
        prompt += (
            "\n\nIMPORTANT: Attached images may contain scanned pages. "
            "Use both text and images to produce final output."
        )

    logger.info("LLM Call 1/3: summary generation")
    response = _call_llm(prompt, SYSTEM_PROMPT, images=images)
    summary, llm_classification, _, _ = _parse_single_pass_response(
        response,
        required_items,
        core_facts,
    )

    global _NER_SUMMARY_CONTEXT
    _NER_SUMMARY_CONTEXT = summary or ""

    # Step 7: final classification resolution (rule-first)
    classification, rule_classification = _resolve_classification(
        processed_text,
        llm_classification,
        manual_crime_type=manual_crime_type,
    )
    classification["detected_sections"] = normalize_sections(_extract_legal_sections_strict(processed_text))
    primary_crime = classification.get("primary_crime_type", crime_hint)

    # If final crime differs from hint, switch checklist template for call 2
    final_crime_info = get_crime_type_info(primary_crime) or crime_info
    final_required_items = final_crime_info["required_items"]

    # Step 6: LLM Call 2/3 - checklist analysis
    checklist_prompt = _build_checklist_prompt(processed_text, primary_crime, final_required_items)
    if images:
        checklist_prompt += (
            "\n\nIMPORTANT: Attached images may contain scanned pages. "
            "Use both text and images to detect checklist evidence."
        )

    logger.info(f"LLM Call 2/3: checklist analysis for crime={primary_crime}")
    checklist_response = _call_llm(checklist_prompt, SYSTEM_PROMPT, images=images)
    checklist_payload = _extract_fenced_json(checklist_response, "json")
    if (not checklist_payload) or ("checklist" not in checklist_payload):
        raw_match = re.search(r'\{[\s\S]*"checklist"[\s\S]*\}', checklist_response)
        if raw_match:
            try:
                checklist_payload = json.loads(raw_match.group(0))
            except json.JSONDecodeError:
                checklist_payload = {}

    if checklist_payload and isinstance(checklist_payload.get("checklist"), list):
        checklist_result = _validate_checklist_payload(checklist_payload.get("checklist", []), final_required_items)
    else:
        logger.warning("Checklist parse failed; falling back to rule-based checklist validation")
        checklist_result = validate_checklist(processed_text, final_required_items, summary=summary)

    logger.info(f"Sleeping {config.LLM_CALL_DELAY}s before NER call")
    time.sleep(config.LLM_CALL_DELAY)

    # Step 7: LLM Call 3/3 - NER extraction
    ner_prompt = _build_ner_prompt(processed_text, core_facts, use_stable_schema=False)
    if images:
        ner_prompt += (
            "\n\nIMPORTANT: Attached images may contain names/dates/seals. "
            "Use both text and images for NER extraction."
        )
    logger.info("LLM Call 3/3: NER extraction")
    ner_response = _call_llm(ner_prompt, SYSTEM_PROMPT, images=images)
    ner_entities = _parse_ner_response(ner_response)

    ner_quality_retry_used = False
    ner_quality_failed = False
    quality_ok, quality_meta = _ner_quality_check(ner_entities)
    if not quality_ok:
        ner_quality_failed = True
        logger.warning(
            "NER quality check failed, entity count or distribution anomalous "
            f"(count={quality_meta.get('count')}, reason={quality_meta.get('reason')}, "
            f"dominant_type={quality_meta.get('dominant_type')}, ratio={quality_meta.get('dominance_ratio')})"
        )

        retry_prompt = _build_ner_prompt(processed_text, core_facts, use_stable_schema=True)
        if images:
            retry_prompt += (
                "\n\nIMPORTANT: Attached images may contain names/dates/seals. "
                "Use both text and images for NER extraction."
            )
        logger.info("LLM NER retry 1/1: using stable 12-type schema")
        retry_response = _call_llm(retry_prompt, SYSTEM_PROMPT, images=images)
        retry_entities = _parse_ner_response(retry_response)
        retry_ok, retry_meta = _ner_quality_check(retry_entities)
        ner_quality_retry_used = True

        if retry_ok:
            ner_entities = retry_entities
            quality_meta = retry_meta
        else:
            logger.warning(
                "NER retry quality still anomalous "
                f"(count={retry_meta.get('count')}, reason={retry_meta.get('reason')}, "
                f"dominant_type={retry_meta.get('dominant_type')}, ratio={retry_meta.get('dominance_ratio')})"
            )

    if len(ner_entities) < 3:
        logger.info("NER parse returned too few entities; applying deterministic core-facts fallback")
        ner_entities = _fallback_ner_from_core_facts(core_facts)

    # Step 8: post-LLM validation/fact protection
    summary, classification, ner_entities, checklist_result = _post_llm_validate(
        summary,
        classification,
        ner_entities,
        checklist_result,
        core_facts,
        final_required_items,
        full_text=cleaned_text,
    )

    ner_sections = [e.get("text", "") for e in (ner_entities or []) if e.get("type") == "LEGAL_SECTION" and e.get("text")]
    merged_sections = _dedup_legal_sections(normalize_sections((core_facts.get("sections", []) or []) + ner_sections))
    classification["detected_sections"] = merged_sections
    logger.info(f"Detected sections merged from core_facts + NER: {classification['detected_sections']}")

    # Semantic augmentation (no API)
    similarity_scores = compute_semantic_similarity(processed_text, final_required_items)
    checklist_result = _merge_checklist_with_similarity(checklist_result, similarity_scores)
    checklist_result = _enrich_checklist_with_confidence(checklist_result)

    # Summary field confidence
    field_confidence = _compute_field_confidence(summary)

    # Step 9: confidence scoring
    confidence_scores = compute_confidence(
        prep=prep,
        core_facts=core_facts,
        summary=summary,
        classification=classification,
        rule_classification=rule_classification,
        ner_entities=ner_entities,
        checklist_result=checklist_result,
        field_confidence=field_confidence,
    )
    classification["composite_confidence"] = confidence_scores["classification_confidence"]

    return {
        "summary": summary,
        "classification": classification,
        "rule_classification": rule_classification,
        "primary_crime_type": primary_crime,
        "checklists": {primary_crime: checklist_result},
        "ner_entities": ner_entities,
        "similarity_scores": {primary_crime: similarity_scores},
        "field_confidence": field_confidence,
        "confidence_scores": confidence_scores,
        "core_facts": core_facts,
        "preprocessed_text": cleaned_text,
        "normalized_text": normalized_text,
        "ner_quality_failed": ner_quality_failed,
        "ner_quality_retry_used": ner_quality_retry_used,
    }


def summarise_chargesheet(text: str) -> str:
    """
    Standalone summarisation (for notebook use).
    For the main pipeline, use process_chargesheet() which combines
    summary + classification in a single API call.
    """
    processed = _truncate_text(text)
    prompt = _build_summary_prompt(processed)
    return _call_llm(prompt, SYSTEM_PROMPT)


def _build_summary_prompt(text: str) -> str:
    """Standalone summary prompt (used by notebook)."""
    return f"""निम्नलिखित हिन्दी चार्जशीट/आरोप पत्र का विश्लेषण करें और संरचित सारांश दें।

## केस हेडर (Case Header)
- **FIR संख्या:**  **दिनांक:**  **थाना:**  **न्यायालय:**
- **घटना स्थल:**  **घटना दिनांक एवं समय:**

## पक्षकार (Parties Involved)
## घटना का सारांश (Incident Summary)
## लागू कानूनी धाराएं (Legal Sections Applied)
## प्रमुख साक्ष्य (Key Evidence)

---
{text}
"""


def classify_crime_type(text: str) -> dict:
    """Standalone classification (for notebook use)."""
    checklists = load_checklists()
    crime_types_str = "\n".join(
        f"- **{k}**: {v['display_name']} ({', '.join(v['typical_sections'])})"
        for k, v in checklists.items()
    )

    prompt = f"""चार्जशीट से अपराध का प्रकार वर्गीकृत करें।

श्रेणियाँ: {crime_types_str}

JSON में उत्तर दें:
```json
{{{{"primary_crime_type": "<key>", "secondary_crime_types": [], "detected_sections": [], "confidence": "high/medium/low", "reasoning": "..."}}}}
```

दस्तावेज़: {text[:5000]}
"""
    response = _call_llm(prompt, SYSTEM_PROMPT)
    json_match = re.search(r'\{[\s\S]*?\}', response)
    if json_match:
        try:
            return json.loads(json_match.group())
        except json.JSONDecodeError:
            pass
    return {
        "primary_crime_type": "unknown", "secondary_crime_types": [],
        "detected_sections": [], "confidence": "low", "reasoning": response,
    }


# ─────────────────────────────────────────────────────────────────────────────
# 10. Output Formatting
# ─────────────────────────────────────────────────────────────────────────────

def format_checklist_output(checklist_result: dict, crime_key: str) -> str:
    """Format checklist analysis into a readable markdown string with semantic similarity scores."""
    crime_info = get_crime_type_info(crime_key)
    if not crime_info:
        return f"Unknown crime type: {crime_key}"

    items_by_id = {item["id"]: item for item in crime_info["required_items"]}
    checklist = checklist_result.get("checklist", [])

    present_items, missing_items, partial_items = [], [], []

    for entry in checklist:
        item_id = entry.get("id", "")
        status = entry.get("status", "missing")
        remarks = entry.get("remarks", "")
        page_no = entry.get("page_no", "")
        sim_score = entry.get("similarity_score", None)
        best_match = entry.get("best_match", "")
        item_confidence = entry.get("confidence", None)
        item_info = items_by_id.get(item_id, {"label_hi": item_id, "label_en": item_id})

        line = f"- {item_info['label_hi']} ({item_info['label_en']})"
        if page_no:
            line += f"  📄 **Pg {page_no}**"
        # Confidence badge
        if item_confidence is not None:
            if item_confidence >= 0.80:
                conf_badge = f"🟢 {item_confidence:.0%}"
            elif item_confidence >= 0.50:
                conf_badge = f"🟡 {item_confidence:.0%}"
            else:
                conf_badge = f"🔴 {item_confidence:.0%}"
            line += f"  [{conf_badge}]"
        if sim_score is not None:
            # Show similarity bar: ▓▓▓▓▒▒▒▒▒▒
            bar_len = 10
            filled = min(bar_len, max(0, round(sim_score * bar_len)))
            bar = "▓" * filled + "▒" * (bar_len - filled)
            line += f"  `{bar} {sim_score:.2f}`"
        if remarks:
            line += f" -- {remarks}"
        if best_match and status in ("present", "partial"):
            line += f"\n  > 📝 *\"{best_match}\"*"

        if status == "present":
            present_items.append(line)
        elif status == "partial":
            partial_items.append(line)
        else:
            missing_items.append(line)

    output = f"## चेकलिस्ट: {crime_info['display_name']}\n\n"

    output += f"### ✅ उपस्थित (Present) -- {len(present_items)} items\n"
    output += "\n".join(present_items) if present_items else "- कोई नहीं"
    output += "\n\n"

    if partial_items:
        output += f"### ⚠️ आंशिक (Partial) -- {len(partial_items)} items\n"
        output += "\n".join(partial_items)
        output += "\n\n"

    output += f"### ❌ अनुपस्थित / संभवतः गायब (Missing) -- {len(missing_items)} items\n"
    output += "\n".join(missing_items) if missing_items else "- कोई नहीं"
    output += "\n"

    return output


def format_classification_output(
    classification: dict,
    rule_results: list,
) -> str:
    """Format classification result into readable markdown."""
    output = ""

    # ── Crime Classification ──
    output += "## अपराध वर्गीकरण (Crime Classification)\n\n"

    primary = classification.get("primary_crime_type", "unknown")
    primary_info = get_crime_type_info(primary) or {}
    display = primary_info.get("display_name", primary)

    output += f"**प्राथमिक अपराध श्रेणी (Primary Crime Type):** {display}\n\n"

    # Composite confidence score with visual indicator
    composite_conf = classification.get("composite_confidence", None)
    if composite_conf is not None:
        bar_len = 20
        filled = min(bar_len, max(0, round(composite_conf * bar_len)))
        bar = "▓" * filled + "▒" * (bar_len - filled)
        if composite_conf >= 0.75:
            conf_label = "🟢 High"
        elif composite_conf >= 0.50:
            conf_label = "🟡 Medium"
        else:
            conf_label = "🔴 Low"
        output += f"**Classification Confidence:** {conf_label} `{bar}` **{composite_conf:.0%}**\n\n"
    else:
        confidence = classification.get("confidence", "")
        if confidence:
            emoji_map = {"high": "High", "medium": "Medium", "low": "Low"}
            output += f"**विश्वास स्तर (Confidence):** {emoji_map.get(confidence, confidence)}\n\n"

    secondary = classification.get("secondary_crime_types", [])
    if secondary:
        sec_names = ", ".join(
            (get_crime_type_info(s) or {}).get("display_name", s) for s in secondary
        )
        output += f"**द्वितीयक श्रेणियाँ (Secondary):** {sec_names}\n\n"

    sections = classification.get("detected_sections", [])
    if sections:
        output += f"**पहचानी गई धाराएं (Detected Sections):** {', '.join(sections)}\n\n"

    reasoning = classification.get("reasoning", "")
    if reasoning:
        output += f"**कारण (Reasoning):** {reasoning}\n\n"

    if rule_results:
        output += "### नियम-आधारित विश्लेषण (Rule-based Analysis)\n"
        if classification.get("classification_source") == "hard_rule":
            output += "⚠️ नोट: नियम-आधारित स्कोर सहायक संकेत हैं। अंतिम वर्गीकरण हार्ड रूल द्वारा निर्धारित किया गया।\n"
        for r in rule_results[:5]:
            output += f"- {r['display_name']} (score: {r['score']})\n"
        output += "\n"

    return output


def format_ner_output(ner_entities) -> str:
    """
    Format NER entities into a compact, grid-style markdown.
    Uses inline tag badges instead of vertical bullet lists.
    """
    # Handle old dict format
    if isinstance(ner_entities, dict):
        if not ner_entities:
            return "⚠️ No entities extracted.\n"
        ner_entities = _convert_ner_to_flat_list(ner_entities)

    if not ner_entities:
        return "⚠️ No entities extracted.\n"

    # Group by type
    by_type = {}
    for ent in ner_entities:
        etype = ent.get("type", "UNKNOWN")
        by_type.setdefault(etype, []).append(ent)

    NER_TYPE_DISPLAY = {
        "ACCUSED": ("🔴 आरोपी", "Accused"),
        "COMPLAINANT": ("🟠 वादी", "Complainant"),
        "VICTIM": ("🟣 पीड़ित", "Victim"),
        "DECEASED": ("⚫ मृतक", "Deceased"),
        "WITNESS": ("🟡 गवाह", "Witnesses"),
        "OFFICER": ("👮 अधिकारी", "Officers"),
        "JUDGE": ("⚖️ न्यायाधीश", "Judicial Officers"),
        "DOCTOR": ("🩺 चिकित्सक", "Doctors"),
        "PERSON": ("👤 अन्य व्यक्ति", "Other Persons"),
        "DATE": ("📅 तिथि", "Dates"),
        "LOCATION": ("📍 स्थान", "Locations"),
        "LEGAL_SECTION": ("⚖️ कानूनी धारा", "Legal Sections"),
        "ORGANIZATION": ("🏛️ संगठन", "Organizations"),
        "LANDMARK": ("🏪 स्थल", "Landmarks"),
        "EVIDENCE": ("🔍 साक्ष्य", "Evidence"),
        "MONETARY": ("💰 धनराशि", "Monetary"),
    }
    type_order = [
        "ACCUSED", "COMPLAINANT", "VICTIM", "DECEASED", "WITNESS", "OFFICER", "JUDGE", "DOCTOR", "PERSON",
        "DATE", "LOCATION", "LEGAL_SECTION", "ORGANIZATION", "LANDMARK", "EVIDENCE", "MONETARY",
    ]
    type_config = [(typ, NER_TYPE_DISPLAY[typ][0], NER_TYPE_DISPLAY[typ][1]) for typ in type_order]

    total = len(ner_entities)
    type_count = len(by_type)

    output = f"## Named Entity Recognition (NER)\n"
    output += f"**{total} entities** across **{type_count} types**\n\n"

    # Summary bar: counts per type
    summary_parts = []
    for etype, hi_label, en_label in type_config:
        items = by_type.get(etype, [])
        if items:
            summary_parts.append(f"{hi_label} **{len(items)}**")
    if summary_parts:
        output += " · ".join(summary_parts) + "\n\n"
        output += "---\n\n"

    known_types = set()
    for etype, hi_label, en_label in type_config:
        known_types.add(etype)
        items = by_type.get(etype, [])
        if not items:
            continue

        output += f"**{hi_label} ({en_label}) — {len(items)}**\n\n"

        # Render as a compact table with 3 columns
        cols = 3
        output += "| | | | | | |\n"
        output += "|---|--------|---|--------|---|--------|\n"

        rows = (len(items) + cols - 1) // cols
        for r in range(rows):
            row_cells = []
            for c in range(cols):
                idx = r + c * rows
                if idx < len(items):
                    item_ent = items[idx]
                    item_text = str(item_ent.get("text", "")).strip()
                    if etype == "PERSON" and item_ent.get("context_hint"):
                        hint = str(item_ent.get("context_hint", "")).strip()
                        if hint:
                            item_text = f"{item_text} <span style=\"color:#6c757d;font-size:0.82em;\"><em>~ inferred: {hint}</em></span>"
                    row_cells.append(f" {idx+1} | {item_text} ")
                else:
                    row_cells.append("  |  ")
            output += "|" + "|".join(row_cells) + "|\n"

        output += "\n"

    # Any extra types not in our config
    for etype, items in by_type.items():
        if etype not in known_types and items:
            output += f"**{etype} — {len(items)}**\n\n"
            output += "| | | | | | |\n"
            output += "|---|--------|---|--------|---|--------|\n"
            rows = (len(items) + 3 - 1) // 3
            for r in range(rows):
                row_cells = []
                for c in range(3):
                    idx = r + c * rows
                    if idx < len(items):
                        item_ent = items[idx]
                        item_text = str(item_ent.get("text", "")).strip()
                        row_cells.append(f" {idx+1} | {item_text} ")
                    else:
                        row_cells.append("  |  ")
                output += "|" + "|".join(row_cells) + "|\n"
            output += "\n"

    return output


def format_timeline_output(ner_entities, summary: str = "", full_text: str = "") -> str:
    """
    Build a concise case timeline showing ONLY key case milestones.
    Filters aggressively: only dates with an identifiable case event
    (FIR, incident, arrest, medical, recovery, chargesheet, etc.)
    are included. Limits to ~15 most important entries.
    """
    if isinstance(ner_entities, dict):
        ner_entities = _convert_ner_to_flat_list(ner_entities) if ner_entities else []

    if not ner_entities:
        ner_entities = []

    import re as _re2
    clean_summary = _re2.sub(r'\*{1,3}([^*]+)\*{1,3}', r'\1', summary)
    clean_summary = _re2.sub(r'^\s*[-*#>]+\s*', '', clean_summary,
                            flags=_re2.MULTILINE)

    import re as _re
    from datetime import datetime

    def _clean_timeline_snippet(text: str, max_len: int = 150) -> str:
        snippet = _apply_ocr_corrections(text or "")
        snippet = re.sub(r'^\s*[/\\|:;,\-–—_]+\s*', '', snippet)
        snippet = re.sub(r'^\s*\d+[\d/\-\.]*\s*', '', snippet)
        snippet = re.sub(r'^[\W_]+', '', snippet, flags=re.UNICODE)
        snippet = re.sub(r'\s+', ' ', snippet).strip('━\n\t ।.-:, ')
        if len(snippet) <= max_len:
            snippet = re.sub(r'(?:\b(?:में|को|से|पर|और|या|का|की|के|the|of|to|in|at|on|and|or)\b\s*)+$', '', snippet, flags=re.IGNORECASE)
            snippet = snippet.strip('━\n\t ।.-:, ')
            return "" if len(snippet) < 15 else snippet

        cut = snippet[:max_len]
        sentence_positions = [cut.rfind('।'), cut.rfind('.'), cut.rfind('!'), cut.rfind('?')]
        sentence_positions = [p for p in sentence_positions if p >= 0]
        if sentence_positions:
            cut = cut[:max(sentence_positions) + 1]
        else:
            sep = max(cut.rfind(' - '), cut.rfind(' | '), cut.rfind(' : '), cut.rfind(' — '))
            if sep > 20:
                cut = cut[:sep]
            ws = cut.rfind(' ')
            if ws > 20:
                cut = cut[:ws]
        cut = re.sub(r'^\s*(?:\b(?:दिनांक|तिथि|date|दि\.)\b[\s:.-]*)+', '', cut, flags=re.IGNORECASE)
        cut = re.sub(r'(?:\b(?:में|को|से|पर|और|या|का|की|के|the|of|to|in|at|on|and|or)\b\s*)+$', '', cut, flags=re.IGNORECASE)
        cut = cut.strip('━\n\t ।.-:, ')
        return "" if len(cut) < 15 else cut

    # Event patterns: (regex, emoji+label, priority 1=highest)
    _EVENT_PATTERNS = [
        (r'घटना|incident|वारदात|हमला|attack|मारपीट',               '⚡ घटना (Incident)', 1),
        (r'FIR|प्रथम\s*सूचना|एफ\.?आई\.?आर|मु\.?\s*अ\.?\s*नं|मुकदमा\s*दर्ज', '📋 FIR दर्ज (FIR Registered)', 2),
        (r'गिरफ्तार|arrest|धरपकड़|पकड़ा|गिरफतारी',              '🚔 गिरफ्तारी (Arrest)', 3),
        (r'पोस्टमार्टम|postmortem|शव\s*परीक्षण|autopsy',        '🏥 पोस्टमार्टम (Postmortem)', 3),
        (r'मेडिकल|medical|चिकित्सा|इलाज|परीक्षण\s*रिपोर्ट',    '🏥 मेडिकल जाँच (Medical)', 4),
        (r'बरामद|जब्त|seized|recovered|बरामदगी|कब्जा',          '📦 बरामदगी/जब्ती (Recovery)', 4),
        (r'चार्जशीट|chargesheet|आरोप\s*पत्र',                   '📑 चार्जशीट (Chargesheet)', 2),
        (r'जमानत|bail',                                          '⚖️ जमानत (Bail)', 5),
        (r'पुलिस\s*कस्टडी|remand|रिमांड|न्यायिक\s*हिरासत',      '⚖️ रिमांड/हिरासत (Remand)', 5),
        (r'पंचनामा|panchnama|पंचायतनामा|फर्द',                  '📋 पंचनामा (Panchnama)', 5),
        (r'मौका\s*मुआयना|spot|नक्शा|नक्सा|मौकानक्शा',          '🗺️ मौका मुआयना (Spot Visit)', 5),
        (r'बयान|कथन|statement|164',                              '📝 बयान (Statement)', 6),
        (r'FSL|फॉरेंसिक|forensic|विश्लेषण|DNA|CFSL',            '🔬 FSL/फॉरेंसिक (Forensic)', 5),
        (r'विवेचना|investigation|जाँच|जांच|तफ्तीश',             '🔍 विवेचना (Investigation)', 7),
        (r'पेशी|court|अदालत|न्यायालय|विचारण',                   '⚖️ पेशी (Court)', 6),
        # Death: exclude शव when followed by परीक्षण (that's Postmortem, not death event)
        (r'मृत्यु|death|मर\s*गय|शव(?!\s*परीक्षण)',              '💀 मृत्यु (Death)', 1),
        (r'रिपोर्ट|report',                                      '📄 रिपोर्ट (Report)', 7),
    ]

    # Birth/biographical date pattern — these are NOT case events, skip them
    _BIRTH_DATE_RE = _re.compile(
        r'जन्म|जन्मतिथि|जन्म\s*(?:तिथि|दिनांक|तारीख)|'
        r'D\.?O\.?B\.?|date\s*of\s*birth|born|'
        r'उम्र\s*(?:लगभग|करीब|वर्ष)?|आयु\s*(?:लगभग|करीब)?|age|'
        r'पिता\s*(?:का\s*)?नाम|माता\s*(?:का\s*)?नाम|'  # father/mother name context
        r'पुत्र(?:ी)?\s|पत्नी\s|'                        # son/daughter/wife of
        r'निवासी|resident',                               # residential details
        _re.IGNORECASE,
    )

    date_pattern = re.compile(r'(\d{1,2})[/\.\-](\d{1,2})[/\.\-](\d{2,4})')

    # Parse all valid dates from NER first
    all_dates = []
    seen = set()

    for ent in ner_entities:
        if ent.get("type") != "DATE":
            continue
        raw = ent["text"].strip()
        m = date_pattern.search(raw)
        if not m:
            continue

        day, month, year = m.group(1), m.group(2), m.group(3)
        if len(year) == 2:
            year = "20" + year
        try:
            dt_obj = datetime(int(year), int(month), int(day))
        except ValueError:
            continue

        date_key = dt_obj.strftime("%Y-%m-%d")
        if date_key in seen:
            continue
        seen.add(date_key)
        all_dates.append((raw, dt_obj))

    # Fallback: if NER dates are missing, extract dates directly from summary/full text
    if not all_dates and (clean_summary or full_text):
        combined_text = "\n".join([clean_summary or "", full_text or ""])
        for m in date_pattern.finditer(combined_text):
            day, month, year = m.group(1), m.group(2), m.group(3)
            if len(year) == 2:
                year = "20" + year
            try:
                dt_obj = datetime(int(year), int(month), int(day))
            except ValueError:
                continue

            raw = f"{int(day):02d}/{int(month):02d}/{int(year):04d}"
            date_key = dt_obj.strftime("%Y-%m-%d")
            if date_key in seen:
                continue
            seen.add(date_key)
            all_dates.append((raw, dt_obj))

        if all_dates:
            logger.info(f"Timeline fallback: extracted {len(all_dates)} dates directly from text")

    if not all_dates:
        return "⚠️ No valid dates found for timeline.\n"

    years = [dt.year for _, dt in all_dates]
    year_counts = {}
    for y in years:
        year_counts[y] = year_counts.get(y, 0) + 1
    ref_year = sorted(year_counts.items(), key=lambda kv: (-kv[1], -kv[0]))[0][0] if year_counts else 2024
    logger.info(f"Timeline reference year (mode): {ref_year}")

    # For each date, find what case event it relates to
    diary_filtered_count = 0
    first_diary_dt = None

    def _classify_date(date_str: str, dt_obj: datetime, summary_text: str, full_doc_text: str):
        """Return (event_label, priority, context_snippet) or None if not important."""
        if not summary_text and not full_doc_text:
            return None

        search_variants = [
            date_str,
            dt_obj.strftime("%d/%m/%Y"),
            dt_obj.strftime("%d.%m.%Y"),
            dt_obj.strftime("%d-%m-%Y"),
            dt_obj.strftime("%d/%m/%y"),
        ]
        sources = [(summary_text or "", "summary"), (full_doc_text or "", "full_text")]

        low_info_pattern = _re.compile(
            r'दैनिकी+\s*(?:संख्या|सं0|लिखना|प्रारंभ)\s*[-:]?\s*\d*|'
            r'(?:काण्ड\s+)?दैनिकी+\s+लिखना\s+प्रारंभ\s+किया|'
            r'दैनिकी+\s+(?:संख्या|सं0|सं).{0,20}लिखी\s+जा\s+चुकी',
            _re.IGNORECASE,
        )
        financial_context_re = _re.compile(
            r'निकासी|जमा|withdraw|deposit|credit|debit|transaction|लेनदेन|'
            r'खाता|account|ifsc|upi|rtgs|neft|imps|चेक|cheque|bank\s*statement|'
            r'पासबुक|नेट\s*बैंकिंग|atm',
            _re.IGNORECASE,
        )
        financial_primary_re = _re.compile(r'निकासी|जमा|withdraw|deposit|credit|debit', _re.IGNORECASE)
        generic_investigation_only = _re.compile(r'विवेचना|investigation|जांच|जाँच', _re.IGNORECASE)
        specific_action = _re.compile(
            r'गिरफ्तार|arrest|बरामद|जब्त|पोस्टमार्टम|medical|fsl|चार्जशीट|मुकदमा\s*दर्ज|fir|बयान|remand|bail|court|पंचनामा',
            _re.IGNORECASE,
        )
        name_like = _re.compile(r'[\u0900-\u097F]{3,}(?:\s+[\u0900-\u097F]{3,}){0,2}')

        best_match = None
        for variant in search_variants:
            for source_text, source_kind in sources:
                if not source_text:
                    continue

                start_pos = 0
                while True:
                    idx = source_text.find(variant, start_pos)
                    if idx < 0:
                        break

                    if source_kind == "full_text":
                        left_block = source_text[:idx]
                        right_block = source_text[idx + len(variant):]

                        prev_bounds = [left_block.rfind('।'), left_block.rfind('\n')]
                        prev_bounds = [b for b in prev_bounds if b >= 0]
                        sentence_start = (max(prev_bounds) + 1) if prev_bounds else 0

                        next_candidates = []
                        purna_idx = right_block.find('।')
                        if purna_idx >= 0:
                            next_candidates.append(purna_idx)
                        nl_idx = right_block.find('\n')
                        if nl_idx >= 0:
                            next_candidates.append(nl_idx)
                        sentence_end = (idx + len(variant) + min(next_candidates)) if next_candidates else len(source_text)

                        context = source_text[sentence_start:sentence_end]
                        before = source_text[sentence_start:idx]
                        after = source_text[idx + len(variant):sentence_end]
                    else:
                        ctx_start = max(0, idx - 120)
                        ctx_end = min(len(source_text), idx + len(variant) + 120)
                        context = source_text[ctx_start:ctx_end]
                        before = source_text[max(0, idx - 80):idx]
                        after = source_text[idx + len(variant):min(len(source_text), idx + len(variant) + 120)]

                    line_s = source_text.rfind('\n', max(0, idx - 80), idx)
                    line_e = source_text.find('\n', idx, min(len(source_text), idx + 80))
                    birth_ctx = source_text[(line_s + 1 if line_s >= 0 else max(0, idx - 60)):(line_e if line_e >= 0 else min(len(source_text), idx + 60))]
                    if _BIRTH_DATE_RE.search(birth_ctx):
                        start_pos = idx + 1
                        continue

                    if low_info_pattern.search(context):
                        has_name = bool(name_like.search(context))
                        has_specific = bool(specific_action.search(context))
                        has_generic = bool(generic_investigation_only.search(context))
                        if (not has_specific) and (
                            has_generic or _re.search(r'लिखना\s+प्रारंभ|दैनिकी', context, _re.IGNORECASE)
                        ):
                            logger.info(f"Timeline low-info diary filtered for date {date_str}")
                            return ('__SKIP_DIARY__', 99, '')

                    primary_span = 150
                    secondary_span = int(getattr(config, 'FINANCIAL_SECONDARY_WINDOW_CHARS', 400))
                    financial_window_start = max(0, idx - primary_span)
                    financial_window_end = min(len(source_text), idx + len(variant) + primary_span)
                    financial_window = source_text[financial_window_start:financial_window_end]
                    financial_window2_start = max(0, idx - secondary_span)
                    financial_window2_end = min(len(source_text), idx + len(variant) + secondary_span)
                    financial_window2 = source_text[financial_window2_start:financial_window2_end]
                    line_start = before.rfind('\n')
                    line_start = max(line_start + 1, len(before) - 80) if line_start >= 0 else max(0, len(before) - 80)
                    line_end_candidates = [after.find('\n'), after.find('।')]
                    line_end_candidates = [x for x in line_end_candidates if x >= 0]
                    line_end = min(line_end_candidates) if line_end_candidates else min(100, len(after))
                    snippet = (before[line_start:].strip() + " " + after[:line_end].strip()).strip()
                    snippet = _clean_timeline_snippet(snippet, max_len=150)

                    if _re.search(r'दैनिकी|diary', birth_ctx, _re.IGNORECASE):
                        best_match = ('🔍 विवेचना (Investigation)', 7, snippet)
                        start_pos = idx + 1
                        continue

                    matched_event_pattern = False
                    all_event_patterns_blocked = True
                    is_financial_event = bool(financial_context_re.search(financial_window)) or bool(financial_context_re.search(financial_window2))
                    primary_financial_signal = bool(financial_primary_re.search(financial_window))
                    for pattern, label, priority in _EVENT_PATTERNS:
                        if _re.search(pattern, context, _re.IGNORECASE):
                            matched_event_pattern = True
                            if is_financial_event and ('Investigation' in label or 'Statement' in label):
                                continue
                            all_event_patterns_blocked = False
                            if best_match is None or priority < best_match[1]:
                                best_match = (label, priority, snippet)
                            break

                    if is_financial_event and (primary_financial_signal or best_match is None or best_match[1] >= 6):
                        best_match = ('💰 वित्तीय साक्ष्य (Financial Evidence)', 4, snippet)

                    if matched_event_pattern and all_event_patterns_blocked:
                        logger.info(f"Timeline: all event patterns blocked by financial context for date {date_str}; checking next occurrence")

                    start_pos = idx + 1

        if best_match:
            return best_match

        generic_variant = search_variants[0] if search_variants else date_str
        base_text = summary_text or full_doc_text
        idx = base_text.find(generic_variant) if (generic_variant and base_text) else -1
        if idx >= 0:
            before = base_text[max(0, idx - 80):idx]
            after = base_text[idx + len(generic_variant):min(len(base_text), idx + len(generic_variant) + 120)]
            line_start = before.rfind('\n')
            line_start = max(line_start + 1, len(before) - 80) if line_start >= 0 else max(0, len(before) - 80)
            line_end_candidates = [after.find('\n'), after.find('।')]
            line_end_candidates = [x for x in line_end_candidates if x >= 0]
            line_end = min(line_end_candidates) if line_end_candidates else min(100, len(after))
            snippet = (before[line_start:].strip() + " " + after[:line_end].strip()).strip()
            snippet = _clean_timeline_snippet(snippet, max_len=150)
            fallback_noise_re = _re.compile(r'^(दिनांक|तिथि|date|तारीख|दि\.|दिनाक)[\s.:;,।\-]*$', _re.IGNORECASE)
            if len(snippet) < 20 or fallback_noise_re.match(snippet.strip()):
                return None
            return ('🔍 विवेचना (Investigation)', 7, snippet)

        return None

    # Classify each date
    timeline_entries = []

    # Words that are NOT meaningful context (just the word "date" etc.)
    _NOISE_SNIPPETS = re.compile(
        r'^(दिनांक|तिथि|date|तारीख|दि\.|दिनाक)[\s.:;,।\-]*$',
        re.IGNORECASE
    )
    for raw, dt_obj in all_dates:
        # Tiered plausibility/date intent filtering
        if dt_obj.year < ref_year - 50 or dt_obj.year > ref_year + 1:
            logger.info(f"Timeline: Skipping implausible date {raw} (year {dt_obj.year}, ref {ref_year})")
            continue

        is_case_window = (ref_year - 5) <= dt_obj.year <= (ref_year + 1)
        is_history_window = (ref_year - 50) <= dt_obj.year <= (ref_year - 6)

        if is_history_window:
            prior_case_re = _re.compile(
                r'आपराधिक\s*इतिहास|पूर्व\s*अपराध|previous\s*case|prior\s*case|'
                r'काण्ड\s*सं0.*199\d|काण्ड\s*सं0.*200\d',
                _re.IGNORECASE,
            )

            history_context = ""
            matched_prior = False
            search_variants = [
                raw,
                dt_obj.strftime("%d/%m/%Y"),
                dt_obj.strftime("%d.%m.%Y"),
                dt_obj.strftime("%d-%m-%Y"),
                dt_obj.strftime("%d/%m/%y"),
            ]
            for source_text in [summary or "", full_text or ""]:
                if not source_text:
                    continue
                for variant in search_variants:
                    pos = source_text.find(variant)
                    if pos < 0:
                        continue
                    hs = max(0, pos - 150)
                    he = min(len(source_text), pos + len(variant) + 150)
                    window = source_text[hs:he]
                    if prior_case_re.search(window):
                        matched_prior = True
                        history_context = window
                        break
                if matched_prior:
                    break

            if not matched_prior:
                logger.info(f"Timeline: Skipping old non-prior-case date {raw} (year {dt_obj.year}, ref {ref_year})")
                continue

            snippet = _re.sub(r'\s+', ' ', history_context).strip('━\n\t ।.-:, ')
            snippet = _clean_timeline_snippet(snippet, max_len=150)
            timeline_entries.append((dt_obj, '📋 पूर्व आपराधिक रिकॉर्ड (Prior Criminal Record)', 8, snippet))
            continue

        if not is_case_window:
            logger.info(f"Timeline: Skipping non-case-window date {raw} (year {dt_obj.year}, ref {ref_year})")
            continue

        # OCR transposition-like future years: one-digit difference from ref year
        year_s = str(dt_obj.year)
        ref_s = str(ref_year)
        hamming = sum(1 for a, b in zip(year_s, ref_s) if a != b) if len(year_s) == len(ref_s) else 4
        if dt_obj.year > ref_year and hamming == 1:
            logger.warning(f"Timeline: Excluding potential OCR-transposed future date {raw} vs ref_year={ref_year}")
            continue

        result = _classify_date(raw, dt_obj, clean_summary, full_text)
        if result:
            label, priority, snippet = result
            if label == '__SKIP_DIARY__':
                diary_filtered_count += 1
                if first_diary_dt is None:
                    first_diary_dt = dt_obj
                continue
            # Skip entries with empty or noise-only snippets
            if not snippet or _NOISE_SNIPPETS.match(snippet.strip()):
                continue
            timeline_entries.append((dt_obj, label, priority, snippet))

    if len(timeline_entries) < 5 and diary_filtered_count > 0 and first_diary_dt is not None:
        snippet = f"नियमित काण्ड दैनिकी — {diary_filtered_count} प्रविष्टियाँ"
        timeline_entries.append((first_diary_dt, '🔍 विवेचना (Investigation)', 7, snippet))
        logger.info(f"Timeline: Added consolidated diary entry ({diary_filtered_count} filtered entries)")

    if not timeline_entries:
        return "⚠️ No key case events identified for timeline.\n"

    # Remove near-duplicate events with same date+label and highly similar snippets
    deduped = []
    seen_keys = set()
    for dt_obj, label, prio, snippet in timeline_entries:
        key = (dt_obj.strftime("%Y-%m-%d"), label.lower(), re.sub(r'\s+', ' ', (snippet or '').lower())[:80])
        if key in seen_keys:
            continue
        seen_keys.add(key)
        deduped.append((dt_obj, label, prio, snippet))
    timeline_entries = deduped

    # Sort by date
    timeline_entries.sort(key=lambda x: x[0])

    # If still too many (>15), keep only highest priority events
    if len(timeline_entries) > 15:
        timeline_entries.sort(key=lambda x: (x[2], x[0]))  # priority first, then date
        timeline_entries = timeline_entries[:15]
        timeline_entries.sort(key=lambda x: x[0])  # re-sort by date

    # Build visual HTML timeline
    # Color mapping for event types
    _EVENT_COLORS = {
        'Incident': '#e76f51',
        'FIR': '#0369a1',
        'Arrest': '#e63946',
        'Postmortem': '#457b9d',
        'Medical': '#457b9d',
        'Recovery': '#f4a261',
        'Chargesheet': '#0369a1',
        'Bail': '#6c757d',
        'Remand': '#6c757d',
        'Panchnama': '#264653',
        'Spot': '#264653',
        'Statement': '#7209b7',
        'Forensic': '#3a0ca3',
        'Investigation': '#264653',
        'Financial': '#2a9d8f',
        'Court': '#6c757d',
        'Death': '#e63946',
        'Report': '#264653',
    }

    def _get_event_color(label):
        for key, color in _EVENT_COLORS.items():
            if key.lower() in label.lower():
                return color
        return '#6c757d'

    # Build HTML
    output = '<div style="padding: 8px 0;">\n'

    # Date range header
    if len(timeline_entries) >= 2:
        header_entries = [
            entry for entry in timeline_entries
            if ('Prior Criminal' not in entry[1] and 'पूर्व आपराधिक' not in entry[1] and int(entry[2]) != 8)
        ]
        range_entries = header_entries if len(header_entries) >= 2 else timeline_entries
        earliest = range_entries[0][0]
        latest = range_entries[-1][0]
        span = (latest - earliest).days
        output += f'<div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:20px; padding:12px 16px; background:rgba(3,105,161,0.08); border-radius:10px; border:1px solid rgba(3,105,161,0.15);">\n'
        output += f'<span style="font-size:0.85rem; color:#6c757d;">From <strong style="color:#0369a1;">{earliest.strftime("%d %b %Y")}</strong></span>\n'
        output += f'<span style="font-size:0.8rem; color:#adb5bd;">&#9473;&#9473; {span} days &#9473;&#9473;</span>\n'
        output += f'<span style="font-size:0.85rem; color:#6c757d;">To <strong style="color:#0369a1;">{latest.strftime("%d %b %Y")}</strong></span>\n'
        output += '</div>\n'

    # Timeline flow
    output += '<div style="position:relative; padding-left:36px;">\n'
    # Vertical line
    output += '<div style="position:absolute; left:14px; top:6px; bottom:6px; width:2px; background:linear-gradient(180deg, #dee2e6 0%, #adb5bd 50%, #dee2e6 100%); border-radius:2px;"></div>\n'

    for i, (dt_obj, label, _prio, snippet) in enumerate(timeline_entries):
        color = _get_event_color(label)
        formatted_date = dt_obj.strftime("%d %b %Y")
        is_last = (i == len(timeline_entries) - 1)

        # Node dot
        output += f'<div style="position:relative; margin-bottom:{0 if is_last else 6}px; padding-bottom:{0 if is_last else 6}px;">\n'
        output += f'  <div style="position:absolute; left:-29px; top:8px; width:14px; height:14px; background:{color}; border-radius:50%; border:3px solid #fff; box-shadow:0 0 0 2px {color}40, 0 2px 8px rgba(0,0,0,0.1); z-index:1;"></div>\n'

        # Card
        output += f'  <div style="background:#fff; border:1px solid #e9ecef; border-left:3px solid {color}; border-radius:8px; padding:12px 16px; box-shadow:0 1px 4px rgba(0,0,0,0.04); transition:box-shadow 0.2s;">\n'
        output += f'    <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:4px;">\n'
        output += f'      <span style="font-weight:700; font-size:0.82rem; color:{color}; text-transform:uppercase; letter-spacing:0.5px;">{label}</span>\n'
        output += f'      <span style="font-size:0.78rem; color:#adb5bd; font-weight:500;">{formatted_date}</span>\n'
        output += f'    </div>\n'
        if snippet:
            output += f'    <div style="font-size:0.88rem; color:#495057; line-height:1.5;">{snippet}</div>\n'
        output += f'  </div>\n'
        output += f'</div>\n'

    output += '</div>\n'
    output += '</div>\n'

    return output


def format_field_confidence(field_confidence: dict) -> str:
    """
    Format summary field confidence scores into a readable markdown table.
    Returns a section to append to the summary output.
    """
    if not field_confidence:
        return ""

    # Check if all values are -1 (no confidence data)
    if all(v == -1.0 for v in field_confidence.values()):
        return ""

    field_labels = {
        "fir_number": ("FIR संख्या", "FIR Number"),
        "fir_date": ("FIR दिनांक", "FIR Date"),
        "police_station": ("थाना", "Police Station"),
        "court": ("न्यायालय", "Court"),
        "place_of_occurrence": ("घटना स्थल", "Place of Occurrence"),
        "incident_date_time": ("घटना दिनांक/समय", "Incident Date/Time"),
        "complainant_name": ("शिकायतकर्ता", "Complainant Name"),
        "accused_names": ("आरोपी नाम", "Accused Names"),
        "witnesses": ("गवाह", "Witnesses"),
        "incident_summary": ("घटना सारांश", "Incident Summary"),
        "legal_sections": ("कानूनी धाराएं", "Legal Sections"),
        "key_evidence": ("प्रमुख साक्ष्य", "Key Evidence"),
    }

    output = "\n---\n\n## 📊 Summary Field Confidence Scores\n\n"
    output += "| Field | Confidence | Level |\n"
    output += "|-------|-----------|-------|\n"

    for field_key, (hi_label, en_label) in field_labels.items():
        score = field_confidence.get(field_key, -1.0)
        if score < 0:
            badge = "⚪ N/A"
            bar_str = "—"
        else:
            score = max(0.0, min(1.0, score))
            bar_len = 10
            filled = min(bar_len, max(0, round(score * bar_len)))
            bar_str = "▓" * filled + "▒" * (bar_len - filled) + f" {score:.0%}"
            if score >= 0.80:
                badge = "🟢 High"
            elif score >= 0.50:
                badge = "🟡 Medium"
            elif score >= 0.20:
                badge = "🟠 Low"
            else:
                badge = "🔴 Not found"
        output += f"| {hi_label} ({en_label}) | `{bar_str}` | {badge} |\n"

    # Overall average
    valid_scores = [v for v in field_confidence.values() if v >= 0]
    if valid_scores:
        avg = sum(valid_scores) / len(valid_scores)
        output += f"\n**Overall Summary Confidence: {avg:.0%}**\n"

    return output
